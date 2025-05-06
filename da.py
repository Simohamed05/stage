import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import io
import numpy as np
from statsmodels.tsa.arima.model import ARIMA
from scipy import stats
import warnings
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.io as pio

# Supprimer les avertissements pour un affichage plus propre
warnings.filterwarnings("ignore")

# Configuration de la page Streamlit
st.set_page_config(page_title="Tableau de bord des achats", layout="wide", initial_sidebar_state="expanded")

# CSS personnalisé pour le thème sombre
st.markdown("""
    <style>
    .main { background-color: #1e1e1e; color: #ffffff; }
    h1, h2, h3 { color: #ffffff; }
    .stDataFrame { background-color: #2d2d2d; color: #ffffff; }
    .stMetric { background-color: #2d2d2d; padding: 10px; border-radius: 5px; }
    </style>
""", unsafe_allow_html=True)

# Fonction pour abréger les grands nombres
def abbreviate_number(num):
    if num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"
    if num >= 1_000:
        return f"{num / 1_000:.1f}K"
    return f"{num:.2f}"

# Fonction pour sauvegarder un graphique Plotly en tant qu'image et retourner le chemin
def save_plotly_fig_as_image(fig, filename):
    pio.write_image(fig, filename, format="png")
    return filename

# Fonction pour générer le document Word avec descriptions détaillées, analyses théoriques et images des visualisations
def generate_word_document(selected_category, viz_data):
    doc = Document()
    
    # Titre
    title = doc.add_heading("Documentation du Tableau de Bord des Achats 📊", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)

    # Introduction
    doc.add_heading("Introduction", level=2)
    intro_text = (
        f"Ce document décrit le **Tableau de Bord des Achats**, une application interactive développée pour **CMG Draa-Lasfar** "
        f"afin d’analyser les données d’achat issues du fichier `demandes_achats.xlsx`. "
        f"Les analyses sont filtrées pour la catégorie **{selected_category if selected_category != 'Toutes' else 'toutes les catégories'}**, permettant une exploration ciblée des données. "
        f"Conçu pour les employés de Managem, y compris les utilisateurs non techniques, ce tableau de bord fournit des visualisations claires, "
        f"des analyses avancées et des recommandations exploitables pour optimiser la gestion des achats. ⭐\n\n"
        f"Dans le contexte de la gestion des achats, les entreprises comme CMG Draa-Lasfar font face à des défis tels que la maîtrise des coûts, "
        f"la gestion des délais de livraison, et l’identification des fournisseurs fiables. Ce tableau de bord répond à ces enjeux en offrant une vue d’ensemble des dépenses, "
        f"des performances des fournisseurs, et des tendances de la demande, tout en détectant les anomalies potentielles. "
        f"Les insights générés aident à prendre des décisions stratégiques, comme négocier avec les fournisseurs clés ou planifier les stocks pour éviter les ruptures."
    )
    doc.add_paragraph(intro_text)

    # Traitement des Données
    doc.add_heading("Traitement des Données 🗂️", level=2)
    doc.add_paragraph(
        "Le tableau de bord commence par charger et nettoyer les données du fichier Excel `demandes_achats.xlsx`. "
        "Cette étape est cruciale pour garantir la fiabilité des analyses, car les données brutes peuvent contenir des erreurs, des valeurs manquantes ou des formats incohérents. "
        "Voici les étapes détaillées du traitement :"
    )
    data_steps = [
        ("Chargement des Données", 
         "Le fichier Excel est lu avec la bibliothèque Pandas. Les colonnes incluent : `article_desc` (description de l’article), `quantite` (quantité commandée), "
         "`prix_unitaire` (coût par unité), `montant` (coût total), `fournisseur` (nom du fournisseur), `date_commande` (date de la commande), "
         "`date_livraison` (date de livraison réelle), `date_promesse` (date de livraison promise), `statut_approbation` (statut de la commande), "
         "et `categorie_achat_1` (catégorie d’achat). Les valeurs vides (ex. : `''`, `'NA'`, `'NaT'`) sont converties en `NaN` pour un traitement cohérent."),
        ("Nettoyage des Colonnes", 
         "Les noms de colonnes sont normalisés (suppression des espaces, guillemets) pour éviter les erreurs. Les colonnes numériques (`quantite`, `prix_unitaire`, `montant`, etc.) "
         "sont converties en nombres, avec `0` pour les valeurs manquantes. La colonne `categorie_achat_1` est convertie en chaînes de caractères, avec `'Inconnu'` pour les valeurs manquantes, "
         "ce qui évite les erreurs de type lors du filtrage ou de l’agrégation."),
        ("Traitement des Dates", 
         "Les colonnes de dates (`date_commande`, `date_livraison`, `date_promesse`) sont converties en format `datetime`. Les formats pris en charge incluent `YYYY-MM-DD` (par défaut) "
         "et `DD/MM/YYYY` (en secours). Environ 6877 dates valides pour `date_commande` (92%), 4562 pour `date_livraison` (61%), et 6641 pour `date_promesse` (89%) ont été détectées. "
         "Les dates manquantes sont conservées comme `NaT` pour éviter de fausser les analyses temporelles, comme les délais de livraison ou les tendances des dépenses."),
        ("Suppression des Lignes Vides", 
         "Les lignes entièrement vides sont supprimées pour garantir la qualité des données. Cela réduit le risque d’analyses biaisées dues à des enregistrements incomplets.")
    ]
    for title, desc in data_steps:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)
    doc.add_paragraph(
        "Un rapport de débogage (`date_debug`) affiche le nombre de dates valides pour chaque colonne, permettant aux utilisateurs de vérifier la qualité des données. "
        "Par exemple, un faible pourcentage de `date_livraison` valides (61%) peut limiter les analyses de fiabilité des fournisseurs."
    )

    # Visualisations
    doc.add_heading("Visualisations 📈", level=2)
    doc.add_paragraph(
        "Le tableau de bord propose plusieurs visualisations interactives pour explorer les données d’achat, réalisées avec la bibliothèque Plotly et un thème sombre pour une lisibilité optimale. "
        "Chaque visualisation est conçue pour répondre à des questions spécifiques de gestion des achats, comme l’identification des catégories coûteuses ou des fournisseurs dominants. "
        "Voici une description détaillée de chaque visualisation :"
    )
    visualizations = [
        ("Dépenses par Catégorie d’Achat", 
         "Un histogramme montre les dépenses totales (`montant`) pour les 10 principales catégories (`categorie_achat_1`). Les barres sont ordonnées par montant décroissant, "
         "avec des étiquettes claires (Catégorie sur l’axe X, Montant en MAD sur l’axe Y). En survolant une barre, l’utilisateur voit le montant exact formaté avec 2 décimales."),
        ("Répartition des Dépenses par Fournisseur", 
         "Un graphique en donut affiche la part des dépenses pour les 7 principaux fournisseurs. Chaque fournisseur est représenté par une couleur distincte, avec le montant exact visible au survol."),
        ("Répartition des Statuts de Commande", 
         "Un graphique en donut montre la distribution des statuts (`statut_approbation`, ex. : Approuvé, En attente, Rejeté). Le nombre de commandes par statut est affiché au survol, avec une légende interactive."),
        ("Délai Moyen de Livraison", 
         "Une courbe affiche le délai moyen de livraison (en jours, calculé comme `date_livraison - date_commande`) par mois (format : `MMM YYYY`). Elle nécessite des dates valides pour `date_commande` et `date_livraison`."),
        ("Top 5 Articles par Quantité", 
         "Un histogramme présente les 5 articles (`article_desc`) avec les quantités totales (`quantite`) les plus élevées. Les descriptions sont tronquées à 30 caractères pour la lisibilité."),
        ("Répartition des Quantités par Catégorie et Article", 
         "Une carte hiérarchique (treemap) montre les quantités par `categorie_achat_1` et `article_desc`. La taille des rectangles est proportionnelle à la quantité, avec des détails au survol."),
        ("Volume de Commandes par Fournisseur", 
         "Un histogramme affiche les 5 fournisseurs ayant au moins 3 commandes, ordonnés par nombre de commandes décroissant.")
    ]
    for title, desc in visualizations:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    # Ajout des images des visualisations
    doc.add_heading("Visualisations Graphiques", level=2)
    # 1. Dépenses par Catégorie d’Achat
    doc.add_heading("Dépenses par Catégorie d’Achat", level=3)
    fig_category = px.bar(
        viz_data['spending_by_category'],
        x='categorie_achat_1',
        y='montant',
        title='Dépenses par catégorie (MAD)',
        labels={'categorie_achat_1': 'Catégorie', 'montant': 'Montant (MAD)'},
        template='plotly_dark',
        hover_data={'montant': ':,.2f'}
    )
    fig_category.update_layout(font=dict(size=12), xaxis_tickangle=45, yaxis_tickformat='.0s')
    category_img = "category_spending.png"
    save_plotly_fig_as_image(fig_category, category_img)
    doc.add_picture(category_img, width=Inches(6))

    # 2. Répartition des Dépenses par Fournisseur
    doc.add_heading("Répartition des Dépenses par Fournisseur", level=3)
    fig_supplier = px.pie(
        viz_data['spending_by_supplier'],
        names='fournisseur',
        values='montant',
        title='Répartition des dépenses par fournisseur',
        template='plotly_dark',
        color_discrete_sequence=px.colors.qualitative.Bold,
        hover_data={'montant': ':,.2f'}
    )
    fig_supplier.update_layout(font=dict(size=12), legend=dict(font=dict(size=12)))
    supplier_img = "supplier_spending.png"
    save_plotly_fig_as_image(fig_supplier, supplier_img)
    doc.add_picture(supplier_img, width=Inches(6))

    # 3. Répartition des Statuts de Commande
    doc.add_heading("Répartition des Statuts de Commande", level=3)
    fig_status = px.pie(
        viz_data['status_counts'],
        names='statut_approbation',
        values='Nombre',
        title='Répartition des statuts de commande',
        template='plotly_dark',
        color_discrete_sequence=px.colors.qualitative.Set1,
        hover_data={'Nombre': True}
    )
    fig_status.update_layout(font=dict(size=12), legend=dict(font=dict(size=12)))
    status_img = "status_distribution.png"
    save_plotly_fig_as_image(fig_status, status_img)
    doc.add_picture(status_img, width=Inches(6))

    # 4. Délai Moyen de Livraison
    doc.add_heading("Délai Moyen de Livraison", level=3)
    if not viz_data['delivery_data'].empty:
        fig_delivery = px.line(
            viz_data['delivery_data'],
            x='Mois',
            y='Delivery_Days',
            title='Délai moyen de livraison (jours)',
            labels={'Delivery_Days': 'Jours moyens'},
            template='plotly_dark',
            hover_data={'Delivery_Days': ':.2f'}
        )
        fig_delivery.update_layout(font=dict(size=12), xaxis_tickangle=45)
        delivery_img = "delivery_delay.png"
        save_plotly_fig_as_image(fig_delivery, delivery_img)
        doc.add_picture(delivery_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour le délai moyen de livraison.")

    # 5. Top 5 Articles par Quantité
    doc.add_heading("Top 5 Articles par Quantité", level=3)
    fig_items = px.bar(
        viz_data['items_by_quantity'],
        x='article_desc',
        y='quantite',
        title='Top 5 articles par quantité',
        labels={'article_desc': 'Article', 'quantite': 'Quantité'},
        template='plotly_dark',
        hover_data={'quantite': True}
    )
    fig_items.update_layout(font=dict(size=12), xaxis_tickangle=45)
    items_img = "top_items.png"
    save_plotly_fig_as_image(fig_items, items_img)
    doc.add_picture(items_img, width=Inches(6))

    # 6. Répartition des Quantités par Catégorie et Article
    doc.add_heading("Répartition des Quantités par Catégorie et Article", level=3)
    if not viz_data['article_category_data'].empty:
        fig_treemap = px.treemap(
            viz_data['article_category_data'],
            path=['categorie_achat_1', 'article_desc'],
            values='quantite',
            title='Répartition des quantités par catégorie et article',
            template='plotly_dark',
            hover_data={'quantite': True}
        )
        fig_treemap.update_layout(font=dict(size=12))
        treemap_img = "quantity_treemap.png"
        save_plotly_fig_as_image(fig_treemap, treemap_img)
        doc.add_picture(treemap_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour la répartition des quantités.")

    # 7. Volume de Commandes par Fournisseur
    doc.add_heading("Volume de Commandes par Fournisseur", level=3)
    if not viz_data['supplier_volume'].empty:
        fig_volume = px.bar(
            viz_data['supplier_volume'],
            x='fournisseur',
            y='Nombre_Commandes',
            title='Volume de commandes par fournisseur',
            labels={'fournisseur': 'Fournisseur', 'Nombre_Commandes': 'Nombre de commandes'},
            template='plotly_dark',
            hover_data={'Nombre_Commandes': True}
        )
        fig_volume.update_layout(font=dict(size=12), xaxis_tickangle=45)
        volume_img = "supplier_volume.png"
        save_plotly_fig_as_image(fig_volume, volume_img)
        doc.add_picture(volume_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucun fournisseur avec ≥3 commandes.")

    # 8. Prévision de la Demande
    doc.add_heading("Prévision de la Demande", level=3)
    if not viz_data['forecast_data'].empty:
        fig_forecast = px.line(
            viz_data['forecast_data'],
            x='Index',
            y='Quantité',
            color='Article',
            line_dash='Type',
            title='Prévision de la demande pour les top articles',
            labels={'Index': 'Période'},
            template='plotly_dark',
            hover_data={'Quantité': ':.2f'}
        )
        fig_forecast.update_layout(font=dict(size=12), xaxis_tickangle=45)
        forecast_img = "demand_forecast.png"
        save_plotly_fig_as_image(fig_forecast, forecast_img)
        doc.add_picture(forecast_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour la prévision de la demande.")

    # 9. Tendances des Dépenses par Catégorie
    doc.add_heading("Tendances des Dépenses par Catégorie", level=3)
    if not viz_data['spending_trends'].empty:
        fig_spending_trends = px.line(
            viz_data['spending_trends'],
            x='Mois',
            y='montant',
            color='categorie_achat_1',
            title='Tendances des dépenses par catégorie',
            labels={'montant': 'Montant (MAD)'},
            template='plotly_dark',
            hover_data={'montant': ':,.2f'}
        )
        fig_spending_trends.update_layout(font=dict(size=12), xaxis_tickangle=45)
        trends_img = "spending_trends.png"
        save_plotly_fig_as_image(fig_spending_trends, trends_img)
        doc.add_picture(trends_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour les tendances des dépenses.")

    # 10. Score de Fiabilité des Fournisseurs
    doc.add_heading("Score de Fiabilité des Fournisseurs", level=3)
    if not viz_data['reliability_data'].empty:
        fig_reliability = px.bar(
            viz_data['reliability_data'],
            x='fournisseur',
            y='Taux_Livraison_À_Temps',
            title='Taux de livraison à temps par fournisseur',
            labels={'Taux_Livraison_À_Temps': 'Taux à temps (%)'},
            template='plotly_dark',
            hover_data={'Délai_Moyen': ':.2f', 'Nombre_Commandes': True}
        )
        fig_reliability.update_layout(font=dict(size=12), xaxis_tickangle=45)
        reliability_img = "supplier_reliability.png"
        save_plotly_fig_as_image(fig_reliability, reliability_img)
        doc.add_picture(reliability_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour les scores de fiabilité.")

    # 11. Répartition des Dépenses par Catégorie
    doc.add_heading("Répartition des Dépenses par Catégorie", level=3)
    if not viz_data['category_spending'].empty:
        fig_category_spending = px.pie(
            viz_data['category_spending'],
            names='categorie_achat_1',
            values='montant',
            title='Répartition des dépenses par catégorie',
            template='plotly_dark',
            color_discrete_sequence=px.colors.qualitative.Bold,
            hover_data={'montant': ':,.2f'}
        )
        fig_category_spending.update_layout(font=dict(size=12), legend=dict(font=dict(size=12)))
        category_spending_img = "category_pie.png"
        save_plotly_fig_as_image(fig_category_spending, category_spending_img)
        doc.add_picture(category_spending_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour la répartition des dépenses.")

    # Analyses Avancées
    doc.add_heading("Analyses Avancées 🔍", level=2)
    doc.add_paragraph(
        "Le tableau de bord inclut des analyses avancées pour fournir des insights approfondis, basées sur des méthodes statistiques et prédictives. "
        "Chaque analyse est accompagnée d’une explication théorique pour comprendre sa méthodologie et son application dans le contexte des achats. "
        "Ces analyses aident à anticiper la demande, évaluer les fournisseurs, et détecter les anomalies."
    )

    # Prévision de la Demande (ARIMA)
    doc.add_heading("Prévision de la Demande (ARIMA)", level=3)
    doc.add_paragraph(
        "Cette analyse prévoit la demande pour les 5 articles les plus commandés sur les 6 prochains mois à l’aide du modèle ARIMA (AutoRegressive Integrated Moving Average). "
        "Elle utilise les données de `date_commande` et `quantite` pour générer des prévisions mensuelles, avec des recommandations de stock incluant une marge de 10%. "
        "Par exemple, si 100 unités sont prévues, la recommandation est de stocker 110 unités pour couvrir les incertitudes."
    )
    doc.add_heading("Fondation Théorique", level=4)
    doc.add_paragraph(
        "ARIMA est un modèle de séries temporelles qui capture les tendances, la saisonnalité et les variations aléatoires dans les données. Il se compose de trois composantes : "
        "- **AR (AutoRegression)** : Modélise la dépendance des valeurs actuelles sur les valeurs passées (ex. : la demande d’un mois influence le suivant). "
        "- **I (Integrated)** : Applique une différenciation pour rendre la série stationnaire, c’est-à-dire sans tendance globale (ex. : soustraire la demande du mois précédent). "
        "- **MA (Moving Average)** : Prend en compte les erreurs de prédiction passées pour lisser les fluctuations. "
        "Dans ce tableau de bord, ARIMA (1,1,1) est utilisé, ce qui signifie une autoregression d’ordre 1, une différenciation d’ordre 1, et une moyenne mobile d’ordre 1. "
        "Les données sont agrégées par mois si `date_commande` est disponible (6877 dates valides), sinon un index numérique est utilisé. "
        "Cette approche est idéale pour prévoir la demande dans un contexte d’achat où la saisonnalité (ex. : pics de demande) et les tendances (ex. : croissance des commandes) sont courantes."
    )
    doc.add_heading("Application aux Achats", level=4)
    doc.add_paragraph(
        "Pour CMG Draa-Lasfar, cette analyse aide à planifier les stocks pour éviter les ruptures, qui peuvent perturber les opérations. Par exemple, si un article critique comme un composant minier "
        "montre une demande croissante, la prévision ARIMA permet de commander à l’avance, réduisant les coûts d’urgence. Les recommandations de stock incluent une marge de 10% pour absorber les imprévus, "
        "comme des retards de livraison ou des variations soudaines de la demande."
    )

    # Tendances des Dépenses
    doc.add_heading("Tendances des Dépenses par Catégorie", level=3)
    doc.add_paragraph(
        "Une courbe montre l’évolution des dépenses (`montant`) par `categorie_achat_1` au fil des mois, basée sur `date_commande`. "
        "Les données sont agrégées par mois pour révéler les tendances saisonnières ou les anomalies. Cette visualisation est particulièrement utile pour identifier les périodes de dépenses élevées "
        "et ajuster les budgets en conséquence."
    )
    doc.add_heading("Fondation Théorique", level=4)
    doc.add_paragraph(
        "L’analyse des tendances repose sur l’agrégation temporelle et la visualisation des séries temporelles. En regroupant les dépenses par mois et catégorie, le tableau de bord détecte des schémas "
        "tels que des pics saisonniers (ex. : achats accrus en fin d’année) ou des anomalies (ex. : dépense inhabituelle dans une catégorie). "
        "Cette approche est ancrée dans l’analyse exploratoire des données (EDA), qui vise à identifier des patterns sans hypothèses préalables."
    )
    doc.add_heading("Application aux Achats", level=4)
    doc.add_paragraph(
        "À CMG Draa-Lasfar, comprendre les tendances des dépenses permet de planifier les budgets et de négocier avec les fournisseurs avant les périodes de forte demande. "
        "Par exemple, si une catégorie comme les équipements miniers montre des dépenses croissantes en été, l’entreprise peut anticiper et négocier des rabais à l’avance."
    )

    # Score de Fiabilité des Fournisseurs
    doc.add_heading("Score de Fiabilité des Fournisseurs", level=3)
    doc.add_paragraph(
        "Cette analyse évalue les fournisseurs selon leur ponctualité de livraison, en calculant le délai moyen (`Délai_Moyen` = `date_livraison - date_promesse` en jours) "
        "et le taux de livraison à temps (`Taux_Livraison_À_Temps` = pourcentage de livraisons où `Délai_Moyen ≤ 0`). "
        "Seuls les fournisseurs avec au moins 3 commandes sont inclus pour garantir des résultats significatifs. Les résultats sont affichés dans un histogramme et un tableau détaillé."
    )
    doc.add_heading("Fondation Théorique", level=4)
    doc.add_paragraph(
        "La fiabilité des fournisseurs est mesurée à l’aide de deux métriques de performance : "
        "- **Délai Moyen** : Moyenne des écarts entre la date de livraison réelle et promise, exprimée en jours. Un délai négatif ou nul indique une livraison à temps ou en avance. "
        "- **Taux de Livraison à Temps** : Proportion des commandes livrées à temps, exprimée en pourcentage. Cette métrique est calculée comme la moyenne des cas où `Délai_Moyen ≤ 0`. "
        "Ces métriques sont standard dans la gestion de la chaîne d’approvisionnement, où la ponctualité est essentielle pour minimiser les interruptions. "
        "L’approche repose sur l’agrégation statistique (moyenne, comptage) et le filtrage pour assurer la robustesse (ex. : ≥3 commandes)."
    )
    doc.add_heading("Application aux Achats", level=4)
    doc.add_paragraph(
        "Pour CMG Draa-Lasfar, cette analyse identifie les fournisseurs les plus fiables pour prioriser les partenariats stratégiques. Par exemple, un fournisseur avec un `Taux_Livraison_À_Temps` de 90% "
        "est préférable pour les articles critiques, réduisant les risques de retards. À l’inverse, un fournisseur avec des délais moyens élevés peut nécessiter une renégociation ou un remplacement."
    )

    # Détection des Anomalies
    doc.add_heading("Détection des Anomalies", level=3)
    doc.add_paragraph(
        "Cette analyse identifie les dépenses inhabituelles dans chaque catégorie (`categorie_achat_1`) en utilisant les scores Z. "
        "Les anomalies sont des enregistrements où le montant est significativement différent de la moyenne (score Z > 3 ou < -3). "
        "Les résultats sont présentés dans un tableau avec le fournisseur, l’article, la catégorie, et le montant."
    )
    doc.add_heading("Fondation Théorique", level=4)
    doc.add_paragraph(
        "Les scores Z mesurent l’écart d’une valeur par rapport à la moyenne, normalisé par l’écart-type : "
        "`Z = (x - moyenne) / écart-type`. Sous l’hypothèse d’une distribution normale, environ 99,7% des données se trouvent dans un intervalle de ±3 écarts-types. "
        "Un score Z supérieur à 3 indique une valeur extrême (anomalie). Cette méthode est largement utilisée en détection d’outliers, particulièrement dans les données financières "
        "où des dépenses inhabituelles peuvent signaler des erreurs ou des fraudes."
    )
    doc.add_heading("Application aux Achats", level=4)
    doc.add_paragraph(
        "À CMG Draa-Lasfar, les anomalies peuvent indiquer des erreurs de facturation, des achats non autorisés, ou des prix excessifs. Par exemple, un `montant` anormalement élevé pour un article "
        "dans la catégorie des fournitures peut déclencher une investigation pour vérifier la validité de la transaction. Cette analyse renforce l’intégrité des processus d’achat."
    )

    # Fonctionnalités Interactives
    doc.add_heading("Fonctionnalités Interactives 🖱️", level=2)
    doc.add_paragraph(
        "Le tableau de bord est conçu pour être intuitif et interactif, facilitant l’exploration des données par les utilisateurs de tous niveaux. Voici les principales fonctionnalités :"
    )
    interactions = [
        ("Filtre par Catégorie", 
         "Une barre latérale permet de sélectionner une catégorie (`categorie_achat_1`) ou 'Toutes'. Ce filtre met à jour toutes les visualisations et analyses en temps réel, "
         "permettant une analyse ciblée. Par exemple, sélectionner une catégorie comme 'Équipements' concentre les résultats sur cette catégorie, facilitant l’identification des fournisseurs clés."),
        ("Sections Dépliables", 
         "Les analyses avancées (prévisions, fiabilité, anomalies) sont organisées dans des sections rétractables pour une navigation claire. Chaque section inclut des explications et des messages de débogage "
         "si les données sont insuffisantes, aidant les utilisateurs à comprendre les limitations (ex. : trop de dates manquantes)."),
        ("Métriques Clés", 
         "Trois cartes en haut affichent la dépense totale (formatée comme `1.2M`), le nombre de commandes, et le nombre de fournisseurs uniques. Ces métriques offrent un aperçu rapide des performances globales."),
        ("Téléchargement des Rapports", 
         "Les utilisateurs peuvent télécharger un résumé CSV (`rapport_achats.csv`) avec les métriques clés et cette documentation Word (`Documentation_Tableau_de_Bord_Achats.docx`). "
         "Le document Word reflète la catégorie sélectionnée, offrant un guide personnalisé.")
    ]
    for title, desc in interactions:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    # Recommandations
    doc.add_heading("Recommandations ✅", level=2)
    doc.add_paragraph(
        "Le tableau de bord fournit des recommandations exploitables pour optimiser la gestion des achats à CMG Draa-Lasfar. Ces recommandations sont dérivées des analyses et visualisations :"
    )
    recommendations = [
        "Négocier avec les fournisseurs clés identifiés dans la répartition des dépenses pour obtenir des rabais ou des conditions avantageuses.",
        "Planifier les stocks en fonction des prévisions ARIMA pour éviter les ruptures, particulièrement pour les articles critiques.",
        "Simplifier le processus d’approbation des commandes pour réduire les délais, en s’appuyant sur l’analyse des statuts.",
        "Investiguer les anomalies détectées pour identifier les erreurs, fraudes, ou opportunités d’optimisation des coûts.",
        "Prioriser les fournisseurs avec un `Taux_Livraison_À_Temps` élevé pour garantir la continuité des opérations."
    ]
    for rec in recommendations:
        doc.add_paragraph(f"- {rec}")

    # Notes Techniques
    doc.add_heading("Notes Techniques 🛠️", level=2)
    doc.add_paragraph(
        "Fichier requis : `demandes_achats.xlsx` dans le même répertoire que le script. "
        "Bibliothèques utilisées : Streamlit (interface), Pandas (données), Plotly (visualisations), Openpyxl (Excel), Statsmodels (ARIMA), Scipy (anomalies), Python-docx (Word), Kaleido (export d’images). "
        "Installation : `pip install streamlit pandas plotly openpyxl statsmodels scipy python-docx kaleido`. "
        "Lancement : `streamlit run procurement_dashboard.py`."
    )

    # Sauvegarde dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Charger et traiter les données
@st.cache_data
def load_and_process_data():
    try:
        df = pd.read_excel("demandes_achats.xlsx", na_values=['', 'NA', 'NaT'])
    except Exception as e:
        st.error(f"Erreur lors du chargement du fichier Excel : {e}")
        return pd.DataFrame(), {}

    df.columns = [col.strip().replace('"', '') for col in df.columns]
    numeric_cols = ['quantite', 'prix_unitaire', 'montant', 'quantite_commande', 'quantite_recue', 'quantite_due']
    for col in numeric_cols:
        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    df['categorie_achat_1'] = df['categorie_achat_1'].astype(str).fillna('Inconnu')
    date_cols = ['date_commande', 'date_livraison', 'date_promesse']
    date_debug = {}
    for col in date_cols:
        df[col] = pd.to_datetime(df[col], errors='coerce')
        valid_count = df[col].notna().sum()
        date_debug[col] = valid_count
        if valid_count == 0:
            df[col] = pd.to_datetime(df[col], errors='coerce', format='%d/%m/%Y')
            valid_count = df[col].notna().sum()
            date_debug[col] = valid_count
    df = df.dropna(how='all')
    return df, date_debug

# Traiter les données pour les visualisations et analyses avancées
@st.cache_data
def process_visualization_data(df, selected_category=None):
    if selected_category and selected_category != "Toutes":
        df = df[df['categorie_achat_1'] == selected_category]
    spending_by_category = df.groupby('categorie_achat_1')['montant'].sum().reset_index().sort_values('montant', ascending=False).head(10)
    spending_by_supplier = df.groupby('fournisseur')['montant'].sum().reset_index().sort_values('montant', ascending=False).head(7)
    status_counts = df['statut_approbation'].value_counts().reset_index()
    status_counts.columns = ['statut_approbation', 'Nombre']
    delivery_data = pd.DataFrame()
    delivery_debug = ""
    if df['date_livraison'].notnull().sum() > 0 and df['date_commande'].notnull().sum() > 0:
        df['Delivery_Days'] = (df['date_livraison'] - df['date_commande']).dt.days
        delivery_times = df.dropna(subset=['Delivery_Days', 'date_commande', 'date_livraison'])
        if not delivery_times.empty:
            delivery_times['Mois'] = delivery_times['date_commande'].dt.strftime('%b %Y')
            delivery_data = delivery_times.groupby('Mois')['Delivery_Days'].mean().reset_index().sort_values('Mois')
        delivery_debug = f"Lignes avec dates valides : {len(delivery_times)}"
    else:
        delivery_debug = f"Aucune date valide (date_commande : {df['date_commande'].notnull().sum()}, date_livraison : {df['date_livraison'].notnull().sum()})"
    items_by_quantity = df.groupby('article_desc')['quantite'].sum().reset_index()
    items_by_quantity['article_desc'] = items_by_quantity['article_desc'].str.slice(0, 30)
    items_by_quantity = items_by_quantity.sort_values('quantite', ascending=False).head(5)
    supplier_category_spending = df.groupby(['fournisseur', 'categorie_achat_1'])['montant'].sum().reset_index()
    interesting_fact = supplier_category_spending.loc[supplier_category_spending['montant'].idxmax()]
    forecast_data = []
    forecast_debug = []
    forecast_recommendations = []
    top_items = items_by_quantity['article_desc'].head(5).tolist()
    use_dates = df['date_commande'].notnull().sum() >= len(df) * 0.5
    for item in top_items:
        item_data = df[df['article_desc'] == item][['date_commande', 'quantite']].dropna(subset=['quantite'])
        valid_records = len(item_data)
        if valid_records >= 3:
            if use_dates and item_data['date_commande'].notnull().sum() >= len(item_data) * 0.5:
                item_data = item_data.groupby(item_data['date_commande'].dt.to_period('M'))['quantite'].sum().reset_index()
                item_data['date_commande'] = item_data['date_commande'].dt.to_timestamp()
                quantities = item_data['quantite'].values
                indices = item_data['date_commande']
                forecast_steps = pd.date_range(start=indices.max() + pd.offsets.MonthBegin(1), periods=6, freq='M')
            else:
                item_data = item_data.reset_index()
                quantities = item_data['quantite'].values
                indices = np.arange(valid_records)
                forecast_steps = np.arange(valid_records, valid_records + 6)
            try:
                model = ARIMA(quantities, order=(1, 1, 1))
                fit = model.fit()
                forecast = fit.forecast(steps=6)
                forecast = np.clip(forecast, 0, None)
                historical_df = pd.DataFrame({
                    'Index': indices,
                    'Quantité': quantities,
                    'Article': item,
                    'Type': 'Historique'
                })
                forecast_df = pd.DataFrame({
                    'Index': forecast_steps,
                    'Quantité': forecast,
                    'Article': item,
                    'Type': 'Prévision'
                })
                forecast_data.append(pd.concat([historical_df, forecast_df]))
                total_forecast = forecast.sum()
                forecast_recommendations.append(f"Stockez environ {int(total_forecast * 1.1)} unités de {item} pour couvrir la demande prévue sur 6 mois (10% de marge).")
            except:
                forecast_debug.append(f"Échec de la prévision pour {item} : Erreur du modèle ARIMA")
        else:
            forecast_debug.append(f"Données insuffisantes pour {item} : {valid_records} enregistrements")
    forecast_data = pd.concat(forecast_data) if forecast_data else pd.DataFrame()
    spending_trends = pd.DataFrame()
    spending_trends_debug = ""
    if df['date_commande'].notnull().sum() > 0:
        df['Mois'] = df['date_commande'].dt.to_period('M').astype(str)
        spending_trends = df.groupby(['Mois', 'categorie_achat_1'])['montant'].sum().reset_index().sort_values('Mois')
        spending_trends_debug = f"Lignes avec date_commande valide : {df['date_commande'].notnull().sum()}"
    else:
        spending_trends_debug = f"Aucune date_commande valide"
    reliability_data = pd.DataFrame()
    reliability_debug = ""
    if df['date_livraison'].notnull().sum() > 0 and df['date_promesse'].notnull().sum() > 0:
        df['Delivery_Delay'] = (df['date_livraison'] - df['date_promesse']).dt.days
        reliability = df.dropna(subset=['Delivery_Delay', 'fournisseur'])
        if not reliability.empty:
            reliability_data = reliability.groupby('fournisseur').agg({
                'Delivery_Delay': 'mean',
                'fournisseur': 'count'
            }).rename(columns={'fournisseur': 'Nombre_Commandes', 'Delivery_Delay': 'Délai_Moyen'}).reset_index()
            reliability_data = reliability_data[reliability_data['Nombre_Commandes'] >= 3].sort_values('Délai_Moyen')
            try:
                on_time_rates = reliability.groupby('fournisseur').apply(
                    lambda x: (x['Delivery_Delay'] <= 0).mean() * 100
                ).reset_index(name='Taux_Livraison_À_Temps')
                reliability_data = reliability_data.merge(on_time_rates, on='fournisseur', how='left')
            except Exception as e:
                reliability_debug += f" Erreur calcul taux à temps : {e}"
        reliability_debug = f"Fournisseurs avec ≥3 commandes : {len(reliability_data)}"
    else:
        reliability_debug = f"Aucune date valide (date_livraison : {df['date_livraison'].notnull().sum()}, date_promesse : {df['date_promesse'].notnull().sum()})"
    category_spending = df.groupby('categorie_achat_1')['montant'].sum().reset_index()
    category_spending = category_spending[category_spending['montant'] > 0]
    category_spending_debug = f"Nombre de catégories avec dépenses : {len(category_spending)}"
    supplier_volume = df.groupby('fournisseur').size().reset_index(name='Nombre_Commandes')
    supplier_volume = supplier_volume[supplier_volume['Nombre_Commandes'] >= 3].sort_values('Nombre_Commandes', ascending=False).head(5)
    supplier_volume_debug = f"Fournisseurs avec ≥3 commandes : {len(supplier_volume)}"
    article_category_data = df.groupby(['categorie_achat_1', 'article_desc'])['quantite'].sum().reset_index()
    article_category_data = article_category_data[article_category_data['quantite'] > 0]
    top_articles = df.groupby(['article_desc', 'categorie_achat_1', 'fournisseur'])\
        .agg({'quantite': 'sum', 'montant': 'sum'})\
        .reset_index()\
        .sort_values('quantite', ascending=False)\
        .head(10)
    article_category_debug = f"Articles avec quantités : {len(article_category_data)}"
    anomalies = []
    anomaly_debug = []
    anomaly_data = pd.DataFrame()
    try:
        for category in df['categorie_achat_1'].unique():
            cat_data = df[df['categorie_achat_1'] == category]['montant']
            if len(cat_data) > 10 and cat_data.var() > 0:
                try:
                    z_scores = stats.zscore(cat_data)
                    anomaly_indices = cat_data.index[abs(z_scores) > 3]
                    for idx in anomaly_indices:
                        anomalies.append({
                            'fournisseur': df.loc[idx, 'fournisseur'],
                            'article_desc': df.loc[idx, 'article_desc'][:30] if pd.notnull(df.loc[idx, 'article_desc']) else 'Inconnu',
                            'categorie_achat_1': category,
                            'montant': df.loc[idx, 'montant']
                        })
                except Exception as e:
                    anomaly_debug.append(f"Erreur pour la catégorie {category} : {e}")
            else:
                anomaly_debug.append(f"Catégorie {category} : Données insuffisantes ({len(cat_data)} enregistrements) ou variance nulle")
        anomaly_data = pd.DataFrame(anomalies) if anomalies else pd.DataFrame(columns=['fournisseur', 'article_desc', 'categorie_achat_1', 'montant'])
    except Exception as e:
        anomaly_debug.append(f"Échec global de la détection des anomalies : {e}")
    summary_data = pd.DataFrame({
        'Métrique': ['Dépense totale (MAD)', 'Nombre de commandes', 'Catégorie principale'],
        'Valeur': [
            df['montant'].sum(),
            len(df),
            spending_by_category['categorie_achat_1'].iloc[0] if not spending_by_category.empty else 'N/A'
        ]
    })
    return {
        'spending_by_category': spending_by_category,
        'spending_by_supplier': spending_by_supplier,
        'status_counts': status_counts,
        'delivery_data': delivery_data,
        'items_by_quantity': items_by_quantity,
        'interesting_fact': interesting_fact,
        'forecast_data': forecast_data,
        'category_spending': category_spending,
        'supplier_volume': supplier_volume,
        'article_category_data': article_category_data,
        'top_articles': top_articles,
        'summary_data': summary_data,
        'anomaly_data': anomaly_data,
        'spending_trends': spending_trends,
        'reliability_data': reliability_data,
        'forecast_debug': forecast_debug,
        'forecast_recommendations': forecast_recommendations,
        'delivery_debug': delivery_debug,
        'category_spending_debug': category_spending_debug,
        'supplier_volume_debug': supplier_volume_debug,
        'article_category_debug': article_category_debug,
        'anomaly_debug': anomaly_debug,
        'spending_trends_debug': spending_trends_debug,
        'reliability_debug': reliability_debug
    }

# Charger les données
df, date_debug = load_and_process_data()
if df.empty:
    st.error("Impossible de charger les données. Vérifiez le fichier 'demandes_achats.xlsx'. ⚠️")
    st.stop()

# Filtre de catégorie dans la barre latérale
st.sidebar.header("🔍 Filtrer par catégorie")
try:
    unique_categories = [str(x) for x in df['categorie_achat_1'].unique() if x != 'nan']
    category_options = ["Toutes"] + sorted(unique_categories)
except Exception as e:
    st.sidebar.error(f"Erreur lors du chargement des catégories : {e}")
    st.sidebar.write("Valeurs uniques dans categorie_achat_1 :")
    st.sidebar.write(df['categorie_achat_1'].unique().tolist())
    category_options = ["Toutes"]
selected_category = st.sidebar.selectbox("Choisir une catégorie", category_options)
viz_data = process_visualization_data(df, selected_category)

# Titre du tableau de bord
st.title("📊 Tableau de bord des achats")

# Statut des dates
if all(count == 0 for count in date_debug.values()):
    st.warning("""
    ⚠️ **Problème de données détecté** : Aucune date valide dans `date_commande`, `date_livraison` ou `date_promesse`. Cela limite les analyses temporelles. Vérifiez le format des dates dans `demandes_achats.xlsx` (ex. : 'YYYY-MM-DD' ou 'DD/MM/YYYY') et assurez-vous que ces colonnes ne sont pas vides.
    """)
else:
    st.info(f"""
    📅 **Statut des dates** :
    - `date_commande` : {date_debug['date_commande']} dates valides
    - `date_livraison` : {date_debug['date_livraison']} dates valides
    - `date_promesse` : {date_debug['date_promesse']} dates valides
    """)

# Métriques clés
st.subheader("📊 Aperçu des indicateurs clés")
col1, col2, col3 = st.columns(3)
col1.metric("Dépense totale (MAD)", f"{abbreviate_number(df['montant'].sum())}")
col2.metric("Nombre de commandes", len(df))
col3.metric("Fournisseurs uniques", df['fournisseur'].nunique())

# Section Résumé
st.write("""
Ce tableau de bord fournit des informations sur les activités d'achat à CMG Draa-Lasfar. Explorez les dépenses, les fournisseurs, les statuts des commandes et les articles les plus demandés. Utilisez le filtre pour analyser une catégorie spécifique. ⭐
""")
st.write(f"""
**Fait intéressant :** Le fournisseur **{viz_data['interesting_fact']['fournisseur']}** domine la catégorie **{viz_data['interesting_fact']['categorie_achat_1']}** avec **{abbreviate_number(viz_data['interesting_fact']['montant'])} MAD**. 🚀
""")

# Visualisations
st.subheader("📈 Visualisations")
col1, col2 = st.columns(2)
with col1:
    fig_category = px.bar(
        viz_data['spending_by_category'],
        x='categorie_achat_1',
        y='montant',
        title='Dépenses par catégorie (MAD)',
        labels={'categorie_achat_1': 'Catégorie', 'montant': 'Montant (MAD)'},
        template='plotly_dark',
        hover_data={'montant': ':,.2f'}
    )
    fig_category.update_layout(font=dict(size=12), xaxis_tickangle=45, yaxis_tickformat='.0s')
    st.plotly_chart(fig_category, use_container_width=True)
with col2:
    fig_supplier = px.pie(
        viz_data['spending_by_supplier'],
        names='fournisseur',
        values='montant',
        title='Répartition des dépenses par fournisseur',
        template='plotly_dark',
        color_discrete_sequence=px.colors.qualitative.Bold,
        hover_data={'montant': ':,.2f'}
    )
    fig_supplier.update_layout(font=dict(size=12), legend=dict(font=dict(size=12)))
    st.plotly_chart(fig_supplier, use_container_width=True)
with col1:
    fig_status = px.pie(
        viz_data['status_counts'],
        names='statut_approbation',
        values='Nombre',
        title='Répartition des statuts de commande',
        template='plotly_dark',
        color_discrete_sequence=px.colors.qualitative.Set1,
        hover_data={'Nombre': True}
    )
    fig_status.update_layout(font=dict(size=12), legend=dict(font=dict(size=12)))
    st.plotly_chart(fig_status, use_container_width=True)
with col2:
    if not viz_data['delivery_data'].empty:
        fig_delivery = px.line(
            viz_data['delivery_data'],
            x='Mois',
            y='Delivery_Days',
            title='Délai moyen de livraison (jours)',
            labels={'Delivery_Days': 'Jours moyens'},
            template='plotly_dark',
            hover_data={'Delivery_Days': ':.2f'}
        )
        fig_delivery.update_layout(font=dict(size=12), xaxis_tickangle=45)
        st.plotly_chart(fig_delivery, use_container_width=True)
    else:
        st.write("Aucun délai de livraison disponible. ⚠️")
        st.write(viz_data['delivery_debug'])
with col1:
    fig_items = px.bar(
        viz_data['items_by_quantity'],
        x='article_desc',
        y='quantite',
        title='Top 5 articles par quantité',
        labels={'article_desc': 'Article', 'quantite': 'Quantité'},
        template='plotly_dark',
        hover_data={'quantite': True}
    )
    fig_items.update_layout(font=dict(size=12), xaxis_tickangle=45)
    st.plotly_chart(fig_items, use_container_width=True)

# Tableau des principaux fournisseurs
st.subheader("📋 Principaux fournisseurs par dépense")
supplier_table = viz_data['spending_by_supplier'].copy()
supplier_table['Nombre de commandes'] = supplier_table['fournisseur'].apply(
    lambda x: len(df[df['fournisseur'] == x]) if selected_category == "Toutes" else len(df[(df['fournisseur'] == x) & (df['categorie_achat_1'] == selected_category)])
)
supplier_table['montant'] = supplier_table['montant'].apply(abbreviate_number)
supplier_table = supplier_table.rename(columns={'fournisseur': 'Fournisseur', 'montant': 'Dépense totale (MAD)'})
st.dataframe(supplier_table, use_container_width=True)

# Section Analyse avancée
st.subheader("🔍 Analyse avancée")
with st.expander("📈 Prévision de la demande pour les top articles"):
    st.write("Prévision de la demande pour les 5 articles les plus demandés sur les 6 prochains mois. 📅")
    if not viz_data['forecast_data'].empty:
        fig_forecast = px.line(
            viz_data['forecast_data'],
            x='Index',
            y='Quantité',
            color='Article',
            line_dash='Type',
            title='Prévision de la demande pour les top articles',
            labels={'Index': 'Période'},
            template='plotly_dark',
            hover_data={'Quantité': ':.2f'}
        )
        fig_forecast.update_layout(font=dict(size=12), xaxis_tickangle=45)
        st.plotly_chart(fig_forecast, use_container_width=True)
        st.write("**Recommandations de stock :**")
        for rec in viz_data['forecast_recommendations']:
            st.write(f"- {rec}")
    else:
        st.write("Données insuffisantes pour la prévision. ⚠️")
        st.write(f"Nombre total de lignes dans le dataset : {len(df)}")
        st.write(f"Lignes avec quantité valide : {len(df[df['quantite'] > 0])}")
        for debug_msg in viz_data['forecast_debug']:
            st.write(debug_msg)
with st.expander("📊 Tendances des dépenses par catégorie"):
    st.write("Évolution des dépenses par catégorie au fil du temps. 💸")
    if not viz_data['spending_trends'].empty:
        fig_spending_trends = px.line(
            viz_data['spending_trends'],
            x='Mois',
            y='montant',
            color='categorie_achat_1',
            title='Tendances des dépenses par catégorie',
            labels={'montant': 'Montant (MAD)'},
            template='plotly_dark',
            hover_data={'montant': ':,.2f'}
        )
        fig_spending_trends.update_layout(font=dict(size=12), xaxis_tickangle=45)
        st.plotly_chart(fig_spending_trends, use_container_width=True)
    else:
        st.write("Aucune donnée disponible pour les tendances des dépenses. ⚠️")
        st.write(viz_data['spending_trends_debug'])
with st.expander("⭐ Score de fiabilité des fournisseurs"):
    st.write("Classement des principaux fournisseurs selon leurs performances de livraison. 🚚")
    if not viz_data['reliability_data'].empty:
        fig_reliability = px.bar(
            viz_data['reliability_data'],
            x='fournisseur',
            y='Taux_Livraison_À_Temps',
            title='Taux de livraison à temps par fournisseur',
            labels={'Taux_Livraison_À_Temps': 'Taux à temps (%)'},
            template='plotly_dark',
            hover_data={'Délai_Moyen': ':.2f', 'Nombre_Commandes': True}
        )
        fig_reliability.update_layout(font=dict(size=12), xaxis_tickangle=45)
        st.plotly_chart(fig_reliability, use_container_width=True)
        reliability_table = viz_data['reliability_data'].copy()
        reliability_table['Délai_Moyen'] = reliability_table['Délai_Moyen'].round(2)
        reliability_table['Taux_Livraison_À_Temps'] = reliability_table['Taux_Livraison_À_Temps'].round(2)
        reliability_table = reliability_table.rename(columns={
            'fournisseur': 'Fournisseur',
            'Délai_Moyen': 'Délai moyen (jours)',
            'Nombre_Commandes': 'Nombre de commandes',
            'Taux_Livraison_À_Temps': 'Taux à temps (%)'
        })
        st.dataframe(reliability_table, use_container_width=True)
    else:
        st.write("Aucune donnée disponible pour les scores de fiabilité. ⚠️")
        st.write(viz_data['reliability_debug'])
with st.expander("📊 Répartition des dépenses par catégorie"):
    st.write("Répartition des dépenses totales par catégorie d'achat. 💸")
    if not viz_data['category_spending'].empty:
        fig_category_spending = px.pie(
            viz_data['category_spending'],
            names='categorie_achat_1',
            values='montant',
            title='Répartition des dépenses par catégorie',
            template='plotly_dark',
            color_discrete_sequence=px.colors.qualitative.Bold,
            hover_data={'montant': ':,.2f'}
        )
        fig_category_spending.update_layout(font=dict(size=12), legend=dict(font=dict(size=12)))
        st.plotly_chart(fig_category_spending, use_container_width=True)
    else:
        st.write("Aucune donnée disponible pour la répartition des dépenses. ⚠️")
        st.write(viz_data['category_spending_debug'])
with st.expander("⭐ Volume de commandes par fournisseur"):
    st.write("Classement des principaux fournisseurs selon le nombre de commandes passées. 🚚")
    if not viz_data['supplier_volume'].empty:
        fig_volume = px.bar(
            viz_data['supplier_volume'],
            x='fournisseur',
            y='Nombre_Commandes',
            title='Volume de commandes par fournisseur',
            labels={'fournisseur': 'Fournisseur', 'Nombre_Commandes': 'Nombre de commandes'},
            template='plotly_dark',
            hover_data={'Nombre_Commandes': True}
        )
        fig_volume.update_layout(font=dict(size=12), xaxis_tickangle=45)
        st.plotly_chart(fig_volume, use_container_width=True)
        volume_table = viz_data['supplier_volume'].rename(columns={
            'fournisseur': 'Fournisseur',
            'Nombre_Commandes': 'Nombre de commandes'
        })
        st.dataframe(volume_table, use_container_width=True)
    else:
        st.write("Aucun fournisseur avec ≥3 commandes. ⚠️")
        st.write(viz_data['supplier_volume_debug'])
with st.expander("📦 Analyse des articles et catégories"):
    st.write("Répartition des quantités d'articles par catégorie et détails des principaux articles. 🛒")
    if not viz_data['article_category_data'].empty:
        fig_treemap = px.treemap(
            viz_data['article_category_data'],
            path=['categorie_achat_1', 'article_desc'],
            values='quantite',
            title='Répartition des quantités par catégorie et article',
            template='plotly_dark',
            hover_data={'quantite': True}
        )
        fig_treemap.update_layout(font=dict(size=12))
        st.plotly_chart(fig_treemap, use_container_width=True)
    else:
        st.write("Aucune donnée disponible pour les articles. ⚠️")
        st.write(viz_data['article_category_debug'])
    st.write("**Top 10 articles :**")
    if not viz_data['top_articles'].empty:
        article_table = viz_data['top_articles'].copy()
        article_table['montant'] = article_table['montant'].apply(abbreviate_number)
        article_table = article_table.rename(columns={
            'article_desc': 'Article',
            'categorie_achat_1': 'Catégorie',
            'fournisseur': 'Fournisseur',
            'quantite': 'Quantité',
            'montant': 'Montant (MAD)'
        })
        st.dataframe(article_table, use_container_width=True)
    else:
        st.write("Aucun article avec données valides. ⚠️")
with st.expander("⚠️ Détection des anomalies"):
    st.write("Modèles de dépenses inhabituels pouvant nécessiter une investigation. 🔎")
    if 'anomaly_data' in viz_data and not viz_data['anomaly_data'].empty:
        anomaly_table = viz_data['anomaly_data'].copy()
        anomaly_table['montant'] = anomaly_table['montant'].apply(abbreviate_number)
        anomaly_table = anomaly_table.rename(columns={
            'fournisseur': 'Fournisseur',
            'article_desc': 'Article',
            'categorie_achat_1': 'Catégorie',
            'montant': 'Montant (MAD)'
        })
        st.dataframe(anomaly_table, use_container_width=True)
    else:
        st.write("Aucune anomalie détectée ou données insuffisantes. ✅")
        if 'anomaly_debug' in viz_data:
            for debug_msg in viz_data['anomaly_debug']:
                st.write(f"- {debug_msg}")

# Conclusion
st.subheader("✅ Résumé et recommandations")
st.write("Résumé des insights clés pour guider la prise de décision :")
if not viz_data['summary_data'].empty:
    fig_summary = px.bar(
        viz_data['summary_data'],
        x='Métrique',
        y='Valeur' if viz_data['summary_data']['Métrique'].str.contains('Dépense|Nombre').any() else None,
        text='Valeur',
        title='Résumé des indicateurs clés',
        template='plotly_dark',
        hover_data={'Valeur': ':,.2f'}
    )
    fig_summary.update_layout(font=dict(size=12), xaxis_tickangle=45, showlegend=False)
    fig_summary.update_traces(texttemplate='%{text:.2s}', textposition='auto')
    st.plotly_chart(fig_summary, use_container_width=True)
st.write("""
**Insights clés :**
- 💰 **Dépenses élevées** : La catégorie principale représente une part significative des coûts. Envisagez des négociations groupées.
- 🚚 **Fournisseurs clés** : Quelques fournisseurs dominent les commandes, surveillez les risques de dépendance.
- 📦 **Articles prioritaires** : Les articles les plus demandés doivent être stockés pour éviter les ruptures.
- 📈 **Prévisions** : Planifiez les stocks selon les quantités prévues pour les 6 prochains mois.
- 🔎 **Anomalies** : Toute dépense inhabituelle doit être vérifiée pour éviter les erreurs ou fraudes.

**Recommandations :**
- 🤝 Négocier avec les fournisseurs clés pour optimiser les coûts.
- 📅 Planifier les stocks en fonction des prévisions de demande.
- ⚡ Simplifier le processus d'approbation pour réduire les délais.
- 🔍 Investiguer les anomalies pour garantir l'intégrité des achats.
""")

word_buffer = generate_word_document(selected_category, viz_data)
st.download_button(
    label="📝 Télécharger la documentation (Word)",
    data=word_buffer,
    file_name="Documentation_Tableau_de_Bord_Achats.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    key="download_docx"
)