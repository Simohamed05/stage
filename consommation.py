import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import io
import numpy as np
from statsmodels.tsa.arima.model import ARIMA
from scipy import stats
import warnings
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.io as pio
import uuid

# Suppress warnings for cleaner output 🚨
warnings.filterwarnings("ignore")

# Streamlit page configuration 🌟
st.set_page_config(page_title="Tableau de Bord Consommation 🌍", layout="wide", initial_sidebar_state="expanded")

# Custom dark theme CSS 🎨
st.markdown("""
    <style>
    .main { background-color: #1e1e1e; color: #ffffff; }
    h1, h2, h3 { color: #ffffff; }
    .stDataFrame { background-color: #2d2d2d; color: #ffffff; }
    .stMetric { background-color: #2d2d2d; padding: 10px; border-radius: 5px; }
    .stButton>button { background-color: #4CAF50; color: white; border-radius: 5px; }
    .stSelectbox, .stSlider { margin-bottom: 10px; }
    .stExpander { background-color: #2d2d2d; border-radius: 5px; }
    </style>
""", unsafe_allow_html=True)

# Utility functions 🛠️
def abbreviate_number(num):
    if pd.isna(num) or num == 0:
        return "0"
    if num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"
    if num >= 1_000:
        return f"{num / 1_000:.1f}K"
    return f"{num:.2f}"

def convert_df_to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def save_plotly_fig_as_image(fig, filename):
    pio.write_image(fig, filename, format="png")
    return filename

def sanitize_text(text):
    if pd.isna(text):
        return "N/A"
    return str(text).replace("\n", " ").replace("\r", " ").strip()

# Function to generate Word report 📝
def generate_word_document(selected_org, selected_cat, date_range, viz_data, filtered_data, total_quantity, total_cost, unique_articles, top_article):
    doc = Document()
    
    # Title
    title = doc.add_heading("Documentation du Tableau de Bord de Consommation 📊✨", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)

    # Introduction
    doc.add_heading("Introduction 🌟", level=2)
    intro_text = (
        f"Ce document décrit le **Tableau de Bord de Consommation**, une application interactive développée pour analyser les données de consommation issues du fichier `consommation.xlsx`. "
        f"Les analyses sont filtrées pour l'organisation **{selected_org if selected_org != 'Tous' else 'toutes les organisations'}**, "
        f"la catégorie **{selected_cat if selected_cat != 'Tous' else 'toutes les catégories'}**, et la période du **{date_range[0]} au {date_range[1]}**. "
        f"Conçu pour les utilisateurs de tous niveaux, ce tableau de bord fournit des visualisations interactives, des analyses avancées, et des recommandations exploitables pour optimiser la gestion de la consommation. 🚀\n\n"
        f"Dans le contexte de la gestion de la consommation, les entreprises font face à des défis tels que la maîtrise des coûts, l’optimisation des stocks, et la détection des anomalies. "
        f"Ce tableau de bord répond à ces enjeux en offrant une vue d’ensemble des tendances de consommation, des catégories dominantes, et des articles les plus consommés, "
        f"tout en identifiant les pics inhabituels pour une meilleure planification. 🌍"
    )
    doc.add_paragraph(intro_text)

    # Data Processing
    doc.add_heading("Traitement des Données 🗂️", level=2)
    doc.add_paragraph(
        "Le tableau de bord commence par charger et nettoyer les données du fichier Excel `consommation.xlsx`. "
        "Cette étape garantit la fiabilité des analyses en corrigeant les erreurs, valeurs manquantes, et formats incohérents. 🧹"
    )
    data_steps = [
        ("Chargement des Données 📥", 
         "Le fichier Excel est lu avec Pandas. Les colonnes incluent : `Date` (date de consommation), `Org_Log` (organisation), `Desc_Cat` (catégorie), "
         "`Article` (nom de l’article), `Qte` (quantité consommée), et `Montant` (coût total). Les lignes vides sont supprimées."),
        ("Nettoyage des Données 🧼", 
         "La colonne `Date` est convertie en format `datetime` (format attendu : `YYYYMMDD`). Les colonnes numériques (`Qte`, `Montant`) sont converties en nombres, "
         "avec `0` pour les valeurs manquantes. Les colonnes `Org_Log`, `Desc_Cat`, et `Article` sont converties en chaînes de caractères, avec `'Inconnu'` pour les valeurs manquantes."),
        ("Calculs Dérivés 🧮", 
         "Le coût moyen par unité est calculé comme `Montant / Qte` (lorsque `Qte > 0`). Les données sont agrégées par article pour éviter les doublons dans certaines visualisations.")
    ]
    for title, desc in data_steps:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    # Global Metrics
    doc.add_heading("Métriques Globales 📏", level=2)
    doc.add_paragraph(f"Quantité Totale: {abbreviate_number(total_quantity)} unités 🛒")
    doc.add_paragraph(f"Coût Total: {abbreviate_number(total_cost)} MAD 💸")
    doc.add_paragraph(f"Coût Moyen par Jour: {abbreviate_number(total_cost / filtered_data['Date'].nunique())} MAD/j 📅")
    doc.add_paragraph(f"Articles Uniques: {unique_articles} 🏷️")
    if top_article is not None:
        doc.add_paragraph(f"Article le Plus Coûteux: {sanitize_text(top_article['Article'])} ({abbreviate_number(top_article['Montant'])} MAD) 💰")

    # Visualizations
    doc.add_heading("Visualisations 📈🚀", level=2)
    doc.add_paragraph(
        "Les visualisations interactives, réalisées avec Plotly et un thème sombre, permettent d’explorer les données de consommation. "
        "Chaque graphique répond à des questions spécifiques, comme l’identification des catégories coûteuses ou des tendances temporelles. 🎨"
    )
    visualizations = [
        ("Tendance Temporelle 📅", 
         "Une courbe montre l’évolution de la quantité ou du coût total par jour, avec une option de lissage (moyenne mobile)."),
        ("Répartition par Catégorie 🥧", 
         "Un graphique en donut affiche la part des coûts ou quantités par catégorie, avec les pourcentages et montants exacts au survol."),
        ("Top Articles 🏆", 
         "Un histogramme présente les 5 articles les plus consommés par quantité ou coût, avec des étiquettes claires."),
        ("Répartition des Quantités par Organisation 📊", 
         "Un graphique en boîte montre la variabilité des quantités consommées par organisation, mettant en évidence les médianes et valeurs aberrantes.")
    ]
    for title, desc in visualizations:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    # Embed Visualizations
    doc.add_heading("Visualisations Graphiques 🖼️", level=2)
    
    # 1. Tendance Temporelle
    doc.add_heading("Tendance Temporelle 📅", level=3)
    fig_trend = px.line(
        viz_data['trend_data'],
        x='Date',
        y='Valeur',
        color='Métrique',
        title=f"Tendance de {viz_data['trend_metric']} au Fil du Temps",
        labels={'Valeur': viz_data['trend_metric']},
        template='plotly_dark',
        hover_data={'Valeur': ':,.2f'}
    )
    fig_trend.update_layout(font=dict(size=12), xaxis_tickangle=45)
    trend_img = "trend.png"
    save_plotly_fig_as_image(fig_trend, trend_img)
    doc.add_picture(trend_img, width=Inches(6))

    # 2. Répartition par Catégorie
    doc.add_heading("Répartition par Catégorie 🥧", level=3)
    if not viz_data['cat_data'].empty and isinstance(viz_data['cat_data'], pd.DataFrame) and 'Desc_Cat' in viz_data['cat_data'].columns and 'Valeur' in viz_data['cat_data'].columns:
        fig_cat = px.pie(
            viz_data['cat_data'],
            names='Desc_Cat',
            values='Valeur',
            title=f"Répartition des {viz_data['cat_metric']} par Catégorie",
            template='plotly_dark',
            color_discrete_sequence=px.colors.qualitative.Bold,
            hover_data={'Valeur': ':,.2f'}
        )
        fig_cat.update_traces(textinfo="percent+label")
        cat_img = "category.png"
        save_plotly_fig_as_image(fig_cat, cat_img)
        doc.add_picture(cat_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour la répartition par catégorie. 😢")

    # 3. Top Articles
    doc.add_heading("Top Articles 🏆", level=3)
    fig_top = px.bar(
        viz_data['top_articles'],
        x='Article',
        y='Valeur',
        title=f"Top 5 Articles par {viz_data['art_metric']}",
        labels={'Valeur': viz_data['art_metric']},
        template='plotly_dark',
        hover_data={'Valeur': ':,.2f'},
        text_auto=True
    )
    fig_top.update_layout(font=dict(size=12), xaxis_tickangle=45)
    top_img = "top_articles.png"
    save_plotly_fig_as_image(fig_top, top_img)
    doc.add_picture(top_img, width=Inches(6))

    # 4. Répartition des Quantités par Organisation
    doc.add_heading("Répartition des Quantités par Organisation 📊", level=3)
    fig_box = px.box(
        viz_data['box_data'],
        x='Org_Log',
        y='Qte',
        title="Répartition des Quantités par Organisation",
        template='plotly_dark',
        hover_data={'Qte': ':,.2f'}
    )
    fig_box.update_layout(font=dict(size=12), xaxis_tickangle=45, showlegend=False)
    box_img = "box.png"
    save_plotly_fig_as_image(fig_box, box_img)
    doc.add_picture(box_img, width=Inches(6))

    # Advanced Analyses
    doc.add_heading("Analyses Avancées 🔍🚀", level=2)
    doc.add_paragraph(
        "Des analyses statistiques et prédictives fournissent des insights approfondis pour anticiper la consommation et détecter les anomalies. 🧠"
    )

    # Demand Forecasting
    doc.add_heading("Prévision de la Consommation 📈", level=3)
    doc.add_paragraph(
        "Prévision de la consommation pour les 5 articles les plus consommés sur les 6 prochains mois à l’aide du modèle ARIMA. 🔮"
    )
    if not viz_data['forecast_data'].empty:
        fig_forecast = px.line(
            viz_data['forecast_data'],
            x='Index',
            y='Quantité',
            color='Article',
            line_dash='Type',
            title='Prévision de la Consommation pour les Top Articles',
            labels={'Index': 'Période'},
            template='plotly_dark',
            hover_data={'Quantité': ':,.2f'}
        )
        fig_forecast.update_layout(font=dict(size=12), xaxis_tickangle=45)
        forecast_img = "forecast.png"
        save_plotly_fig_as_image(fig_forecast, forecast_img)
        doc.add_picture(forecast_img, width=Inches(6))
    else:
        doc.add_paragraph("Aucune donnée disponible pour la prévision de la consommation. 😢")

    # Anomaly Detection
    doc.add_heading("Détection des Anomalies ⚠️", level=3)
    doc.add_paragraph(
        "Identification des consommations inhabituelles par catégorie à l’aide des scores Z (Z > 3 ou < -3). 🕵️‍♂️"
    )
    if not viz_data['anomaly_data'].empty:
        table = doc.add_table(rows=1, cols=len(viz_data['anomaly_data'].columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(viz_data['anomaly_data'].columns):
            hdr_cells[i].text = col
        for _, row in viz_data['anomaly_data'].head(10).iterrows():
            row_cells = table.add_row().cells
            for col_idx, value in enumerate(row):
                row_cells[col_idx].text = sanitize_text(value)
    else:
        doc.add_paragraph("Aucune anomalie détectée. ✅")

    # Save document to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Load and process data 📂
@st.cache_data
def load_and_process_data():
    try:
        df = pd.read_excel("consommation.xlsx", na_values=['', 'NA', 'NaT'])
    except Exception as e:
        st.error(f"Erreur lors du chargement du fichier Excel : {e} 😢")
        return pd.DataFrame()
    
    # Clean data 🧹
    df = df.dropna(how='all')
    df["Date"] = pd.to_datetime(df["Date"], format="%Y%m%d", errors='coerce')
    df["Qte"] = pd.to_numeric(df["Qte"], errors='coerce').fillna(0)
    df["Montant"] = pd.to_numeric(df["Montant"], errors='coerce').fillna(0)
    df["Org_Log"] = df["Org_Log"].astype(str).fillna('Inconnu')
    df["Desc_Cat"] = df["Desc_Cat"].astype(str).fillna('Inconnu')
    df["Article"] = df["Article"].astype(str).fillna('Inconnu')
    df["unit_cost"] = df.apply(lambda row: row["Montant"] / row["Qte"] if row["Qte"] > 0 else 0, axis=1)
    
    # Aggregated data for visualizations 📊
    df_aggregated = df.groupby("Article").agg({
        "Qte": "sum",
        "Montant": "sum",
        "Org_Log": "first",
        "Desc_Cat": "first",
        "Date": "first"
    }).reset_index()
    df_aggregated["unit_cost"] = df_aggregated.apply(
        lambda row: row["Montant"] / row["Qte"] if row["Qte"] > 0 else 0, axis=1
    )
    
    return df, df_aggregated

# Process data for visualizations and analyses 📈
@st.cache_data
def process_visualization_data(df, df_aggregated, selected_org, selected_cat, date_range, trend_metric, cat_metric, art_metric):
    filtered_df = df.copy()
    filtered_df_aggregated = df_aggregated.copy()
    
    # Apply filters 🔍
    filtered_df = filtered_df[(filtered_df["Date"].dt.date >= date_range[0]) & (filtered_df["Date"].dt.date <= date_range[1])]
    if selected_org != "Tous":
        filtered_df = filtered_df[filtered_df["Org_Log"] == selected_org]
        filtered_df_aggregated = filtered_df_aggregated[filtered_df_aggregated["Org_Log"] == selected_org]
    if selected_cat != "Tous":
        filtered_df = filtered_df[filtered_df["Desc_Cat"] == selected_cat]
        filtered_df_aggregated = filtered_df_aggregated[filtered_df_aggregated["Desc_Cat"] == selected_cat]
    
    # Visualization data 📊
    trend_data = filtered_df.groupby(filtered_df["Date"].dt.date)[trend_metric].sum().reset_index()
    trend_data.columns = ['Date', 'Valeur']
    trend_data['Métrique'] = trend_metric
    
    # Handle category data with robust validation 🛡️
    if not filtered_df.empty and cat_metric in filtered_df.columns and 'Desc_Cat' in filtered_df.columns:
        try:
            # Perform groupby and ensure DataFrame output
            cat_data = filtered_df.groupby("Desc_Cat")[cat_metric].sum().reset_index()
            # Rename columns explicitly
            cat_data.columns = ['Desc_Cat', 'Valeur']
            # Verify that cat_data is a DataFrame
            if not isinstance(cat_data, pd.DataFrame):
                cat_data = pd.DataFrame({'Desc_Cat': [], 'Valeur': []})
            elif cat_data.empty:
                cat_data = pd.DataFrame(columns=['Desc_Cat', 'Valeur'])
        except Exception as e:
            st.warning(f"Erreur lors du traitement des données de catégorie : {e} 😢")
            cat_data = pd.DataFrame(columns=['Desc_Cat', 'Valeur'])
    else:
        cat_data = pd.DataFrame(columns=['Desc_Cat', 'Valeur'])
    
    top_articles = filtered_df_aggregated.nlargest(5, art_metric)[["Article", art_metric]]
    top_articles.columns = ['Article', 'Valeur']
    box_data = filtered_df[["Org_Log", "Qte"]]
    
    # Forecasting 🔮
    forecast_data = []
    forecast_recommendations = []
    top_items = filtered_df_aggregated.nlargest(5, "Qte")["Article"].tolist()
    for item in top_items:
        item_data = filtered_df[filtered_df["Article"] == item][["Date", "Qte"]].dropna(subset=['Qte'])
        if len(item_data) >= 3:
            item_data = item_data.groupby(item_data["Date"].dt.to_period('M'))["Qte"].sum().reset_index()
            item_data["Date"] = item_data["Date"].dt.to_timestamp()
            quantities = item_data["Qte"].values
            indices = item_data["Date"]
            forecast_steps = pd.date_range(start=indices.max() + pd.offsets.MonthBegin(1), periods=6, freq='M')
            try:
                model = ARIMA(quantities, order=(1, 1, 1))
                fit = model.fit()
                forecast = fit.forecast(steps=6)
                forecast = np.clip(forecast, 0, None)
                historical_df = pd.DataFrame({
                    'Index': indices,
                    'Quantité': quantities,
                    'Article': item,
                    'Type': 'Historique'
                })
                forecast_df = pd.DataFrame({
                    'Index': forecast_steps,
                    'Quantité': forecast,
                    'Article': item,
                    'Type': 'Prévision'
                })
                forecast_data.append(pd.concat([historical_df, forecast_df]))
                total_forecast = forecast.sum()
                forecast_recommendations.append(f"Stockez environ {int(total_forecast * 1.1)} unités de {item} pour couvrir la consommation prévue sur 6 mois (10% de marge). 📦")
            except:
                pass
    forecast_data = pd.concat(forecast_data) if forecast_data else pd.DataFrame()
    
    # Anomaly detection 🕵️‍♂️
    anomalies = []
    for cat in filtered_df["Desc_Cat"].unique():
        cat_data = filtered_df[filtered_df["Desc_Cat"] == cat]["Montant"]
        if len(cat_data) > 10 and cat_data.var() > 0:
            z_scores = stats.zscore(cat_data)
            anomaly_indices = cat_data.index[abs(z_scores) > 3]
            for idx in anomaly_indices:
                anomalies.append({
                    'Article': filtered_df.loc[idx, 'Article'],
                    'Desc_Cat': cat,
                    'Montant': filtered_df.loc[idx, 'Montant']
                })
    anomaly_data = pd.DataFrame(anomalies) if anomalies else pd.DataFrame(columns=['Article', 'Desc_Cat', 'Montant'])
    
    return {
        'trend_data': trend_data,
        'trend_metric': trend_metric,
        'cat_data': cat_data,
        'cat_metric': cat_metric,
        'top_articles': top_articles,
        'art_metric': art_metric,
        'box_data': box_data,
        'forecast_data': forecast_data,
        'forecast_recommendations': forecast_recommendations,
        'anomaly_data': anomaly_data
    }

# Load data 📂
df, df_aggregated = load_and_process_data()
if df.empty:
    st.error("Impossible de charger les données. Vérifiez le fichier 'consommation.xlsx'. ⚠️😢")
    st.stop()

# Sidebar filters 🔍
st.sidebar.header("🔍 Filtres & Paramètres 🛠️")
org_options = ["Tous"] + sorted(df["Org_Log"].unique().tolist())
selected_org = st.sidebar.selectbox("Organisation 🏢", org_options)
cat_options = ["Tous"] + sorted(df["Desc_Cat"].unique().tolist())
selected_cat = st.sidebar.selectbox("Catégorie 🗂️", cat_options)
date_range = st.sidebar.date_input("Période 📅", 
                                  [df["Date"].min().date(), df["Date"].max().date()],
                                  min_value=df["Date"].min().date(),
                                  max_value=df["Date"].max().date())

# Advanced options ⚙️
with st.sidebar.expander("⚙️ Options avancées 🔧"):
    trend_metric = st.radio("Métrique Tendance 📈", ["Qte", "Montant"])
    cat_metric = st.radio("Métrique Catégorie 🥧", ["Qte", "Montant"], key="cat_met")
    art_metric = st.radio("Métrique Article 🏆", ["Qte", "Montant"], key="art_met")
    smoothing = st.slider("Lissage (jours) ⏳", 0, 30, 7)

# Process visualization data 📊
viz_data = process_visualization_data(df, df_aggregated, selected_org, selected_cat, date_range, trend_metric, cat_metric, art_metric)

# Dashboard title 🌟
st.title("📊 Tableau de Bord de Consommation 🌍")
st.write("Explorez les données de consommation avec des visualisations interactives et des analyses avancées pour optimiser la gestion des ressources. 🚀✨")

# KPI Metrics 📏
st.header("📌 Métriques Globales 🌟")
total_quantity = df_aggregated["Qte"].sum()
total_cost = df_aggregated["Montant"].sum()
unique_articles = df_aggregated["Article"].nunique()
top_article = df_aggregated.loc[df_aggregated["Montant"].idxmax()] if not df_aggregated.empty else None

col1, col2, col3, col4 = st.columns(4)
col1.metric("Quantité Totale 🛒", f"{abbreviate_number(total_quantity)} unités")
col2.metric("Coût Total 💸", f"{abbreviate_number(total_cost)} MAD")
col3.metric("Coût Moyen par Jour 📅", f"{abbreviate_number(total_cost / df['Date'].nunique())} MAD/j")
col4.metric("Articles Uniques 🏷️", unique_articles)

# Interesting Fact 🔍
st.header("🔍 Aperçu Intéressant 🌟")
if top_article is not None:
    st.info(
        f"L’article le plus coûteux est **{sanitize_text(top_article['Article'])}** "
        f"avec un coût total de **{abbreviate_number(top_article['Montant'])} MAD**. 💰✨"
    )
else:
    st.warning("Aucune donnée disponible pour l’aperçu intéressant. 😢")

# Export Data 💾
st.header("💾 Exporter les Données 📥")
st.download_button(
    label="Télécharger les Données en CSV 📄",
    data=convert_df_to_csv(df_aggregated),
    file_name="donnees_consommation.csv",
    mime="text/csv",
)

# Visualizations 📈
st.header("📈 Visualisations 🎨")
col1, col2 = st.columns(2)

with col1:
    st.subheader(f"Tendance de {trend_metric} au Fil du Temps 📅")
    st.markdown("Cette courbe montre l’évolution quotidienne de la consommation, avec une option de lissage. 📈")
    fig_trend = px.line(
        viz_data['trend_data'],
        x='Date',
        y='Valeur',
        color='Métrique',
        title=f"Tendance de {trend_metric} au Fil du Temps",
        labels={'Valeur': trend_metric},
        template='plotly_dark',
        hover_data={'Valeur': ':,.2f'}
    )
    if smoothing > 0:
        smoothed = viz_data['trend_data']['Valeur'].rolling(window=smoothing).mean()
        fig_trend.add_scatter(
            x=viz_data['trend_data']['Date'],
            y=smoothed,
            mode='lines',
            name='Lissé',
            line=dict(dash='dash', color='blue')
        )
    fig_trend.update_layout(font=dict(size=12), xaxis_tickangle=45)
    st.plotly_chart(fig_trend, use_container_width=True)

with col2:
    st.subheader(f"Répartition des {cat_metric} par Catégorie 🥧")
    st.markdown("Ce graphique en donut montre la part de chaque catégorie dans la consommation totale. 🎨")
    if not viz_data['cat_data'].empty and isinstance(viz_data['cat_data'], pd.DataFrame) and 'Desc_Cat' in viz_data['cat_data'].columns and 'Valeur' in viz_data['cat_data'].columns:
        fig_cat = px.pie(
            viz_data['cat_data'],
            names='Desc_Cat',
            values='Valeur',
            title=f"Répartition des {cat_metric} par Catégorie",
            template='plotly_dark',
            color_discrete_sequence=px.colors.qualitative.Bold,
            hover_data={'Valeur': ':,.2f'}
        )
        fig_cat.update_traces(textinfo="percent+label")
        fig_cat.update_layout(font=dict(size=12))
        st.plotly_chart(fig_cat, use_container_width=True)
    else:
        st.warning("Aucune donnée disponible pour la répartition par catégorie. 😢⚠️")

with col1:
    st.subheader(f"Top 5 Articles par {art_metric} 🏆")
    st.markdown("Ce graphique présente les articles les plus consommés par quantité ou coût. 🌟")
    fig_top = px.bar(
        viz_data['top_articles'],
        x='Article',
        y='Valeur',
        title=f"Top 5 Articles par {art_metric}",
        labels={'Valeur': art_metric},
        template='plotly_dark',
        hover_data={'Valeur': ':,.2f'},
        text_auto=True
    )
    fig_top.update_layout(font=dict(size=12), xaxis_tickangle=45)
    st.plotly_chart(fig_top, use_container_width=True)

with col2:
    st.subheader("Répartition des Quantités par Organisation 📊")
    st.markdown("Ce graphique en boîte montre la variabilité des quantités consommées par organisation. 📈")
    fig_box = px.box(
        viz_data['box_data'],
        x='Org_Log',
        y='Qte',
        title="Répartition des Quantités par Organisation",
        template='plotly_dark',
        hover_data={'Qte': ':,.2f'}
    )
    fig_box.update_layout(font=dict(size=12), xaxis_tickangle=45, showlegend=False)
    st.plotly_chart(fig_box, use_container_width=True)

# Advanced Analyses 🔍
st.header("🔍 Analyses Avancées 🚀")
with st.expander("📈 Prévision de la Consommation 🔮"):
    st.markdown("Prévision de la consommation pour les 5 articles les plus consommés sur les 6 prochains mois. 🌟")
    if not viz_data['forecast_data'].empty:
        fig_forecast = px.line(
            viz_data['forecast_data'],
            x='Index',
            y='Quantité',
            color='Article',
            line_dash='Type',
            title='Prévision de la Consommation pour les Top Articles',
            labels={'Index': 'Période'},
            template='plotly_dark',
            hover_data={'Quantité': ':,.2f'}
        )
        fig_forecast.update_layout(font=dict(size=12), xaxis_tickangle=45)
        st.plotly_chart(fig_forecast, use_container_width=True)
        st.markdown("**Recommandations de stock :** 📦")
        for rec in viz_data['forecast_recommendations']:
            st.markdown(f"- {rec}")
    else:
        st.warning("Données insuffisantes pour la prévision. ⚠️😢")

with st.expander("⚠️ Détection des Anomalies 🕵️‍♂️"):
    st.markdown("Identification des consommations inhabituelles par catégorie (scores Z > 3 ou < -3). 🔍")
    if not viz_data['anomaly_data'].empty:
        st.dataframe(viz_data['anomaly_data'], use_container_width=True)
    else:
        st.warning("Aucune anomalie détectée. ✅🎉")

# Data Table 📋
st.header("📋 Données de Consommation 🗂️")
items_per_page = 10
total_pages = (len(df_aggregated) + items_per_page - 1) // items_per_page
page = st.slider("Page 📄", 1, max(1, total_pages), 1)
start_idx = (page - 1) * items_per_page
end_idx = start_idx + items_per_page
paged_df = df_aggregated.iloc[start_idx:end_idx]
st.dataframe(
    paged_df[["Article", "Org_Log", "Desc_Cat", "Qte", "Montant", "unit_cost"]].style.format(
        {"Qte": "{:.2f}", "Montant": "{:.2f}", "unit_cost": "{:.2f}"}
    ),
    use_container_width=True
)
st.write(f"Page {page} sur {total_pages} 📖")

# Word Report Download 📥
st.header("📥 Téléchargement du Rapport 📝")
word_buffer = generate_word_document(selected_org, selected_cat, date_range, viz_data, df, total_quantity, total_cost, unique_articles, top_article)
st.download_button(
    label="Télécharger le Rapport en Word 📄✨",
    data=word_buffer,
    file_name="rapport_consommation.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)