import streamlit as st
import pandas as pd
import plotly.express as px
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import plotly.io as pio
from PIL import Image
import openai
import httpx

# Initialize OpenAI client with provided API key and no proxies
try:
    client = openai.OpenAI(
        api_key="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",
        http_client=httpx.Client(proxies={})  # Explicitly disable proxies
    )
except Exception as e:
    st.error(f"Erreur lors de l'initialisation du client OpenAI : {str(e)}")
    st.stop()

# Cache expensive computations
@st.cache_data
def compute_monthly_costs(data):
    monthly_data = data.groupby('Mois')['Montant'].sum().reset_index()
    month_order = ['Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin', 
                   'Juillet', 'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre']
    monthly_data['Mois'] = pd.Categorical(monthly_data['Mois'], categories=month_order, ordered=True)
    return monthly_data.sort_values('Mois')

@st.cache_data
def compute_category_breakdown(data):
    return data.groupby('Desc_Cat')['Montant'].sum().reset_index()

# Function to save Plotly figure as image
def save_plotly_fig_as_image(fig):
    img_bytes = pio.to_image(fig, format='png', width=800, height=350)
    img = Image.open(BytesIO(img_bytes))
    img_buffer = BytesIO()
    img.save(img_buffer, format='PNG')
    return img_buffer.getvalue()

# Function to generate Word document
def generate_word_report(engin_data, selected, figs, descriptions, metrics, predictions, budget_threshold):
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Rapport d\'Analyse pour R1600-{selected.split("-")[-1]}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Date du rapport : {pd.Timestamp.now().strftime("%d/%m/%Y")}')
    
    # Table of Contents
    doc.add_heading('Tableau des Matières', level=2)
    doc.add_paragraph('1. Introduction\n2. Résumé des Analyses\n3. Visualisations\n4. Indicateurs Clés\n5. Prévisions\n6. Recommandations', style='List Bullet')
    
    # Introduction
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'Ce rapport présente une analyse détaillée des coûts associés à l\'engin R1600, '
        'incluant les visualisations, les indicateurs clés, les prévisions et les recommandations '
        'pour optimiser la gestion des dépenses.'
    )
    
    # Summary
    doc.add_heading('2. Résumé des Analyses', level=2)
    total_cost = engin_data['Montant'].sum()
    top_category = engin_data.groupby('Desc_Cat')['Montant'].sum().idxmax()
    trend = predictions.get('trend', 'Stable')
    doc.add_paragraph(
        f'Coût total : {total_cost:,.0f} MAD\n'
        f'Catégorie principale : {top_category}\n'
        f'Tendance récente : {trend}\n'
        f'Nombre d\'interventions : {len(engin_data)}\n'
        f'Coût médian : {engin_data["Montant"].median():,.0f} MAD'
    )
    
    # Visualizations
    doc.add_heading('3. Visualisations', level=2)
    for i, (fig, desc, title) in enumerate(zip(figs, descriptions, [
        'Évolution des Dépenses avec Projection', 'Distribution des Coûts', 
        'Répartition par Catégorie', 'Coût Mensuel'
    ])):
        doc.add_heading(f'3.{i+1} {title}', level=3)
        img_stream = BytesIO(save_plotly_fig_as_image(fig))
        doc.add_picture(img_stream, width=Inches(6))
        doc.add_paragraph(desc)
    
    # Key Metrics
    doc.add_heading('4. Indicateurs Clés', level=2)
    table = doc.add_table(rows=len(metrics) + 1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrique'
    hdr_cells[1].text = 'Valeur'
    for i, (metric, value) in enumerate(metrics.items()):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)
    
    # Predictions
    doc.add_heading('5. Prévisions', level=2)
    doc.add_paragraph(
        f'Estimation moyenne : {predictions["avg"]:,.0f} MAD/mois\n'
        f'Intervalle de confiance (95%) : {predictions["ci_lower"]:,.0f} - {predictions["ci_upper"]:,.0f} MAD\n'
        f'Fiabilité : {predictions["reliability"]:.0f}%\n'
        f'Tendance : {predictions["trend"]}'
    )
    
    # Budget Threshold Analysis
    if budget_threshold > 0:
        high_costs = engin_data[engin_data['Montant'] > budget_threshold]
        doc.add_heading('Analyse des Coûts Excédant le Seuil', level=3)
        doc.add_paragraph(
            f'Seuil défini : {budget_threshold:,.0f} MAD\n'
            f'Nombre d\'interventions dépassant le seuil : {len(high_costs)}\n'
            f'Coût total des interventions dépassant le seuil : {high_costs["Montant"].sum():,.0f} MAD'
        )
    
    # Recommendations
    doc.add_heading('6. Recommandations', level=2)
    doc.add_paragraph(
        'Sur la base des analyses, voici les recommandations :\n'
        '- Prioriser la maintenance préventive pour réduire les coûts dans la catégorie principale.\n'
        f'- Examiner les interventions coûteuses dans {top_category} pour identifier des alternatives économiques.\n'
        '- Mettre en place un suivi mensuel pour détecter les tendances haussières tôt.\n'
        '- Négocier avec les fournisseurs pour les pièces fréquemment utilisées.'
    )
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to generate chatbot response using OpenAI
def generate_response(prompt, context_data):
    try:
        # Prepare context with data summary
        data_summary = (
            f"Les données concernent les engins R1600. Coût total: {context_data['Montant'].sum():,.0f} MAD, "
            f"nombre d'interventions: {len(context_data)}, "
            f"catégorie principale: {context_data.groupby('Desc_Cat')['Montant'].sum().idxmax()}. "
            "Demandez des détails sur les coûts, catégories, ou engins spécifiques."
        )
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Vous êtes un assistant analysant des données sur les engins R1600. Fournissez des réponses précises et concises basées sur les données fournies."},
                {"role": "user", "content": f"{data_summary}\nQuestion: {prompt}"}
            ],
            max_tokens=200
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Erreur: {str(e)}"

# Configuration de la page
st.set_page_config(layout="wide", page_title="Analyse des Engins R1600")

# Custom CSS for floating chatbot icon and container
st.markdown("""
<style>
.chatbot-icon {
    position: fixed;
    bottom: 20px;
    right: 20px;
    background-color: #F28C38;
    color: white;
    border-radius: 50%;
    width: 60px;
    height: 60px;
    display: flex;
    align-items: center;
    justify-content: center;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    cursor: pointer;
    font-size: 24px;
    z-index: 1000;
}
.chatbot-container {
    position: fixed;
    bottom: 90px;
    right: 20px;
    width: 350px;
    max-height: 500px;
    background-color: #fff;
    border-radius: 10px;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    display: none;
    flex-direction: column;
    z-index: 1000;
    overflow: hidden;
}
.chatbot-container.open {
    display: flex;
}
.chatbot-header {
    background-color: #F28C38;
    color: white;
    padding: 10px;
    font-weight: bold;
}
.chatbot-messages {
    flex-grow: 1;
    overflow-y: auto;
    padding: 10px;
    background-color: #f5f5f5;
}
.chatbot-input {
    border-top: 1px solid #ddd;
    padding: 10px;
    background-color: white;
}
.chatbot-input input {
    width: 100%;
    border: 1px solid #ddd;
    border-radius: 5px;
    padding: 8px;
    font-size: 14px;
}
.stChatMessage {
    margin-bottom: 10px;
}
.stChatMessage.user {
    text-align: right;
}
.stChatMessage.user > div {
    background-color: #F28C38;
    color: white;
    display: inline-block;
    padding: 8px 12px;
    border-radius: 10px;
    max-width: 80%;
}
.stChatMessage.assistant > div {
    background-color: #e0e0e0;
    display: inline-block;
    padding: 8px 12px;
    border-radius: 10px;
    max-width: 80%;
}
</style>
""", unsafe_allow_html=True)

# Initialize session state for chatbot
if "chat_open" not in st.session_state:
    st.session_state.chat_open = False
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = [
        {"role": "assistant", "content": "Bonjour ! Posez vos questions sur les engins R1600 ou les données du tableau de bord."}
    ]

# Charger les données
@st.cache_data
def load_data():
    try:
        df = pd.read_excel('engins2.xlsx', sheet_name='BASE DE DONNEE')
        
        # Nettoyage des données
        df['Desc_CA'] = df['Desc_CA'].str.replace('', '').str.strip()
        df['Desc_Cat'] = df['Desc_Cat'].str.strip()
        df['Montant'] = pd.to_numeric(df['Montant'], errors='coerce')
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
        
        # Liste des engins spécifiques à analyser
        engins_cibles = [
            'CHARGEUSE CATERPILLAR 10T   R1600 Nｰ14 DS',
            'CHARGEUSE CATERPILLAR R 1600H Nｰ15 DS',
            'Chargeuse CATERPILLAR 10T  R1600 Nｰ16',
            'Chargeuse CATERPILLAR 10T R1600 Nｰ17',
            'Chargeuse  CAT    R1600 10T  Nｰ18',                
            'CHARGEUSE CATERPILLAR R 1600 Nｰ20',
            'CHARGEUSE CATERPILLAR R 1600 Nｰ21',
            'CHARGEUSE CATERPILLAR R1600 Nｰ22',
            'CHARGEUSE CATERPILLAR R1600 Nｰ23'
        ]
        
        # Filtrer uniquement les engins cibles
        df = df[df['Desc_CA'].isin(engins_cibles)].copy()
        
        # Extraire le numéro de l'engin et créer un nom standardisé
        df['Numéro_Engin'] = df['Desc_CA'].str.extract(r'Nｰ(\d+)')
        df['Engin_Formaté'] = 'R1600-' + df['Numéro_Engin']
        
        # Mois en français
        months_fr = {
            'January': 'Janvier', 'February': 'Février', 'March': 'Mars',
            'April': 'Avril', 'May': 'Mai', 'June': 'Juin',
            'July': 'Juillet', 'August': 'Août', 'September': 'Septembre',
            'October': 'Octobre', 'November': 'Novembre', 'December': 'Décembre'
        }
        df['Mois'] = df['Date'].dt.month_name().map(months_fr)
        
        return df
    except Exception as e:
        st.error(f"Erreur lors du chargement du fichier Excel : {str(e)}")
        return None

# Chargement des données
df = load_data()

# Vérifier si les données sont chargées
if df is None or df.empty:
    st.stop()

# =============================================
# SECTION 1 : EN-TÊTE AVEC KPI PRINCIPAUX
# =============================================
st.markdown("""
<div style='background-color:#e3f2fd; padding:20px; border-radius:10px; border-left:5px solid #1976d2; margin-bottom:20px;'>
    <h1 style='color:#F28C38; text-align:center; margin-top:0;'>📊 Analyse des Engins R1600</h1>
    <p style='color:#424242; text-align:center;'>Suivi des coûts et des interventions pour optimiser la gestion des chargeuses</p>
</div>
""", unsafe_allow_html=True)

kpi_container = st.container()
with kpi_container:
    kpi1, kpi2, kpi3, kpi4 = st.columns(4)
    with kpi1:
        st.markdown(f"""
        <div style='background-color:#fff8e1; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Coût total</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{df['Montant'].sum():,.0f} MAD</p>
            <p style='color:#424242; font-size:12px;'>Somme totale des dépenses</p>
        </div>
        """, unsafe_allow_html=True)
    with kpi2:
        top_engine = df.groupby('Engin_Formaté')['Montant'].sum().idxmax()
        st.markdown(f"""
        <div style='background-color:#e8f5e9; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Engin le plus coûteux</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{top_engine.split('-')[-1]}</p>
            <p style='color:#424242; font-size:12px;'>Numéro de l'engin</p>
        </div>
        """, unsafe_allow_html=True)
    with kpi3:
        top_category = df.groupby('Desc_Cat')['Montant'].sum().idxmax()
        st.markdown(f"""
        <div style='background-color:#f3e5f5; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Catégorie principale</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{top_category}</p>
            <p style='color:#424242; font-size:12px;'>Dépense dominante</p>
        </div>
        """, unsafe_allow_html=True)
    with kpi4:
        avg_cost = df['Montant'].mean()
        st.markdown(f"""
        <div style='background-color:#ffebee; padding:15px; border-radius:10px; text-align:center;'>                                                                                                      
            <h6 style='color:#F28C38; margin-bottom:10px;'>Coût moyen</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{avg_cost:,.0f} MAD</p>
            <p style='color:#424242; font-size:12px;'>Par intervention</p>
        </div>
        """, unsafe_allow_html=True)

# =============================================
# SECTION 2 : FILTRES (SIDEBAR)
# =============================================
with st.sidebar:
    st.subheader("Filtres")
    
    selected_engin = st.radio(
        'Sélectionner un engin',
        ['Tous'] + sorted(df['Engin_Formaté'].unique().tolist()),
        help="Filtrer par engin spécifique"
    )
    
    # Sélecteur de plage de dates
    st.subheader("Plage de dates")
    
    default_start = df['Date'].min().date()
    default_end = df['Date'].max().date()
    date_range = st.date_input(
        "Période",
        value=(default_start, default_end),
        min_value=default_start,
        max_value=default_end,
        help="Choisissez une plage de dates pour filtrer les interventions"
    )
    
    # Statistiques filtrées
    filtered_data = df.copy()
    if selected_engin != 'Tous':
        filtered_data = filtered_data[filtered_data['Engin_Formaté'] == selected_engin]
    if len(date_range) == 2:  # Vérifie qu'une plage complète est sélectionnée
        start_date, end_date = date_range
        filtered_data = filtered_data[(filtered_data['Date'].dt.date >= start_date) & 
                                    (filtered_data['Date'].dt.date <= end_date)]
    
    # Vérifier si 'Montant' existe dans filtered_data
    if 'Montant' not in filtered_data.columns:
        st.error("Erreur : La colonne 'Montant' est introuvable dans les données filtrées. Colonnes disponibles : " + str(filtered_data.columns.tolist()))
    elif filtered_data.empty:
        st.warning("Aucune donnée disponible après filtrage. Veuillez ajuster les filtres.")
    else:
        total_cost = filtered_data['Montant'].sum()
        avg_cost = filtered_data['Montant'].mean()
        num_interventions = len(filtered_data)
        st.subheader("Statistiques")
        st.metric("Consommation totale", f"{total_cost:,.0f} MAD")
        st.metric("Nombre d'interventions", num_interventions)
        st.metric("Coût moyen", f"{avg_cost:,.0f} MAD")

# Appliquer les filtres
filtered_data = df.copy()
if selected_engin != 'Tous':
    filtered_data = filtered_data[filtered_data['Engin_Formaté'] == selected_engin]
if len(date_range) == 2:
    start_date, end_date = date_range
    filtered_data = filtered_data[(filtered_data['Date'].dt.date >= start_date) & 
                                (filtered_data['Date'].dt.date <= end_date)]

# Vérifier si filtered_data est vide ou si 'Montant' est absent
if filtered_data.empty:
    st.warning("Aucune donnée disponible après filtrage. Veuillez ajuster les filtres.")
    st.stop()
if 'Montant' not in filtered_data.columns:
    st.error("Erreur : La colonne 'Montant' est introuvable après filtrage. Colonnes disponibles : " + str(filtered_data.columns.tolist()))
    st.stop()

# =============================================
# SECTION 3 : ONGLETS PRINCIPAUX
# =============================================
tab1, tab2, tab3, tab4 = st.tabs(["📋 Données", "📊 Analyse", "🔄 Comparaisons", "💡 Recommandations"])

# Onglet 1 : Données
with tab1:
    st.markdown("""
    <div style='background-color:#e8f5e9; padding:20px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>Données détaillées des consommations</h2>
        <p style='color:#424242;'>Visualisez et exportez les interventions enregistrées</p>
    </div>
    """, unsafe_allow_html=True)
    
    height = 600 if st.checkbox("Afficher toutes les données (défilement)") else 300
    
    st.dataframe(
        filtered_data[['Date', 'Engin_Formaté', 'Desc_Cat', 'Montant', 'Mois']]
        .sort_values(['Engin_Formaté', 'Date'])
        .style.format({
            'Montant': '{:,.0f} MAD',
            'Date': lambda x: x.strftime('%d/%m/%Y') if not pd.isnull(x) else ''
        }),
        height=height,
        use_container_width=True
    )
    
    st.markdown("""
    <div style='background-color:#f3e5f5; padding:20px; border-radius:10px; border-left:5px solid #8e24aa; margin-top:20px;'>
        <h3 style='color:#F28C38; margin-top:0;'>Exporter les données</h3>
        <p style='color:#424242;'>Téléchargez les données filtrées au format souhaité</p>
    </div>
    """, unsafe_allow_html=True)
    
    export_col1, export_col2 = st.columns([1, 3])
    with export_col1:
        export_format = st.radio("Format", ['CSV', 'Excel'])
    with export_col2:
        if export_format == 'CSV':
            csv = filtered_data.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Télécharger CSV",
                data=csv,
                file_name='consommation_r1600.csv',
                mime='text/csv'
            )
        else:
            output = BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                filtered_data.to_excel(writer, index=False)
            st.download_button(
                label="📥 Télécharger Excel",
                data=output.getvalue(),
                file_name='consommation_r1600.xlsx',
                mime='application/vnd.ms-excel'
            )

# Onglet 2 : Analyse
with tab2:
    # Select engin for analysis
    if selected_engin == 'Tous':
        selected = st.selectbox('Choisir un engin à analyser', sorted(df['Engin_Formaté'].unique()), key='engin_select')
        engin_data = df[df['Engin_Formaté'] == selected]
    else:
        engin_data = df[df['Engin_Formaté'] == selected_engin]
    
    # Main layout: Visualizations on left, Metrics and Predictions on right
    col1, col2 = st.columns([7, 3])

    # Visualizations Section
    with col1:
        st.markdown("""
        <div style='background-color:#fff8e1; padding:15px; border-radius:10px; border-left:5px solid #ffa000; margin-bottom:20px;'>
            <h3 style='color:#F28C38; margin-top:0;'>📈 Visualisations des Coûts</h3>
        </div>
        """, unsafe_allow_html=True)

        # Store figures and descriptions for Word export
        figs = []
        descriptions = []

        # Graph 1: Evolution and Projection
        st.markdown("#### Évolution des Dépenses")
        fig1 = px.line(
            engin_data.groupby('Date')['Montant'].sum().reset_index(),
            x='Date', y='Montant',
            title='Évolution des Dépenses avec Projection',
            height=350,
            template='plotly_white'
        )
        fig1.update_traces(line=dict(color='#F28C38'), hovertemplate='%{x|%d/%m/%Y}<br>%{y:,.0f} MAD')
        
        if len(engin_data) >= 3:
            dates = engin_data.groupby('Date')['Montant'].sum().index
            x = np.arange(len(dates))
            y = engin_data.groupby('Date')['Montant'].sum().values
            coeff = np.polyfit(x, y, 1)
            future_dates = [dates[-1] + pd.DateOffset(months=i) for i in range(1, 4)]
            projection = np.polyval(coeff, [x[-1]+1, x[-1]+2, x[-1]+3])
            
            fig1.add_scatter(
                x=future_dates, y=projection, mode='lines+markers',
                name='Projection (linéaire)', line=dict(color='red', dash='dot'),
                hovertemplate='%{x|%d/%m/%Y}<br>%{y:,.0f} MAD'
            )
            fig1.add_scatter(
                x=future_dates, y=projection * 1.3, mode='lines',
                name='Fourchette haute (+30%)', line=dict(color='orange', dash='dash'),
                hovertemplate='%{x|%d/%m/%Y}<br>%{y:,.0f} MAD'
            )
            fig1.add_scatter(
                x=future_dates, y=projection * 0.7, mode='lines',
                name='Fourchette basse (-30%)', line=dict(color='green', dash='dash'),
                hovertemplate='%{x|%d/%m/%Y}<br>%{y:,.0f} MAD'
            )
        
        fig1.update_layout(
            xaxis_title="Date", yaxis_title="Montant (MAD)",
            legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5),
            margin=dict(t=50, b=50)
        )
        st.plotly_chart(fig1, use_container_width=True)
        desc1 = (
            "Cette courbe montre l'évolution des dépenses avec une projection linéaire sur 3 mois. "
            "Les fourchettes haute et basse indiquent une variabilité potentielle de ±30%, permettant "
            "d'anticiper les tendances futures."
        )
        st.markdown(f"<p style='color:#424242; font-size:14px;'>{desc1}</p>", unsafe_allow_html=True)
        figs.append(fig1)
        descriptions.append(desc1)

        # Graph 2: Cost Distribution
        st.markdown("#### Distribution des Coûts")
        fig2 = px.histogram(
            engin_data, x='Montant',
            title='Distribution des Coûts',
            height=350,
            template='plotly_white',
            nbins=20
        )
        fig2.update_traces(marker=dict(color='#1976d2'), hovertemplate='Coût: %{x:,.0f} MAD<br>Fréquence: %{y}')
        fig2.update_layout(
            xaxis_title="Montant (MAD)", yaxis_title="Fréquence",
            margin=dict(t=50, b=50)
        )
        st.plotly_chart(fig2, use_container_width=True)
        desc2 = (
            "Cet histogramme montre la répartition des coûts par intervention, mettant en évidence "
            "les montants les plus fréquents et les valeurs aberrantes (outliers). Cela aide à identifier "
            "les interventions coûteuses pour optimiser la gestion budgétaire."
        )
        st.markdown(f"<p style='color:#424242; font-size:14px;'>{desc2}</p>", unsafe_allow_html=True)
        figs.append(fig2)
        descriptions.append(desc2)

        # Graph 3: Category Breakdown
        st.markdown("#### Répartition par Catégorie")
        fig3 = px.pie(
            compute_category_breakdown(engin_data),
            values='Montant', names='Desc_Cat',
            title='Répartition par Catégorie',
            height=350,
            template='plotly_white'
        )
        fig3.update_traces(textinfo='percent+label', hovertemplate='%{label}: %{value:,.0f} MAD (%{percent})')
        fig3.update_layout(margin=dict(t=50, b=50))
        st.plotly_chart(fig3, use_container_width=True)
        desc3 = (
            "Ce graphique montre la répartition des dépenses par catégorie, permettant d'identifier "
            "les principales sources de coûts et de prioriser les efforts de réduction des dépenses."
        )
        st.markdown(f"<p style='color:#424242; font-size:14px;'>{desc3}</p>", unsafe_allow_html=True)
        figs.append(fig3)
        descriptions.append(desc3)

        # Graph 4: Monthly Costs
        st.markdown("#### Coûts Mensuels")
        monthly_data = compute_monthly_costs(engin_data)
        fig4 = px.bar(
            monthly_data,
            x='Mois', y='Montant',
            title='Coût Mensuel',
            height=350,
            template='plotly_white'
        )
        fig4.update_traces(marker=dict(color='#388e3c'), hovertemplate='%{x}<br>%{y:,.0f} MAD')
        fig4.update_layout(
            xaxis_title="Mois", yaxis_title="Montant (MAD)",
            margin=dict(t=50, b=50)
        )
        st.plotly_chart(fig4, use_container_width=True)
        desc4 = (
            "Ce graphique à barres montre les dépenses mensuelles, mettant en évidence les variations "
            "saisonnières ou les pics de coûts, utiles pour planifier les budgets mensuels."
        )
        st.markdown(f"<p style='color:#424242; font-size:14px;'>{desc4}</p>", unsafe_allow_html=True)
        figs.append(fig4)
        descriptions.append(desc4)

    # Metrics and Predictions Section
    with col2:
        # Key Metrics
        st.markdown("""
        <div style='background-color:#f3e5f5; padding:15px; border-radius:10px; border-left:5px solid #8e24aa; margin-bottom:20px;'>
            <h3 style='color:#F28C38; margin-top:0;'>📊 Indicateurs Clés</h3>
        </div>
        """, unsafe_allow_html=True)

        if 'Montant' not in engin_data.columns:
            st.error("Erreur : La colonne 'Montant' est introuvable. Colonnes disponibles : " + str(engin_data.columns.tolist()))
        elif engin_data.empty:
            st.warning("Aucune donnée disponible pour cet engin. Veuillez sélectionner un autre engin.")
        else:
            last_month = engin_data.groupby('Mois')['Montant'].sum().iloc[-1] if not engin_data.groupby('Mois')['Montant'].sum().empty else 0
            avg_3m = engin_data.groupby('Mois')['Montant'].sum().tail(3).mean() if len(engin_data.groupby('Mois')['Montant'].sum()) >= 3 else 0
            max_cost = engin_data['Montant'].max()
            total_cost = engin_data['Montant'].sum()
            num_interventions = len(engin_data)
            median_cost = engin_data['Montant'].median()
            cost_variance = engin_data['Montant'].var() if len(engin_data) > 1 else 0

            metrics = {
                'Dernier mois': f"{last_month:,.0f} MAD",
                'Moyenne 3 mois': f"{avg_3m:,.0f} MAD",
                'Coût maximal': f"{max_cost:,.0f} MAD",
                'Coût total': f"{total_cost:,.0f} MAD",
                'Interventions': num_interventions,
                'Coût médian': f"{median_cost:,.0f} MAD",
                'Variance des coûts': f"{cost_variance:,.0f} MAD²"
            }
            st.markdown(f"""
            <div style='color:#ffffff; font-size:14px;'>
                <p><strong>Dernier mois :</strong> {last_month:,.0f} MAD</p>
                <p><strong>Moyenne 3 mois :</strong> {avg_3m:,.0f} MAD</p>
                <p><strong>Coût maximal :</strong> {max_cost:,.0f} MAD</p>
                <p><strong>Coût total :</strong> {total_cost:,.0f} MAD</p>
                <p><strong>Interventions :</strong> {num_interventions}</p>
                <p><strong>Coût médian :</strong> {median_cost:,.0f} MAD</p>
                <p><strong>Variance des coûts :</strong> {cost_variance:,.0f} MAD²</p>
            </div>
            """, unsafe_allow_html=True)

        # Predictions
        st.markdown("""
        <div style='background-color:#e8f5e9; padding:15px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:20px;'>
            <h3 style='color:#F28C38; margin-top:0;'>🔮 Prévisions des Coûts</h3>
            <p style='color:#424242;'>Estimations basées sur les données historiques</p>
        </div>
        """, unsafe_allow_html=True)

        predictions = {}
        if len(engin_data) >= 3:
            last_3 = engin_data.groupby('Mois')['Montant'].sum().tail(3)
            avg = last_3.mean()
            std = last_3.std() if len(last_3) > 1 else 0
            num_months = len(engin_data.groupby('Mois'))
            reliability = min(90, 50 + 5 * num_months - 10 * (std / avg if avg > 0 else 0))
            reliability = max(50, reliability)
            ci_lower = avg - 1.96 * std / np.sqrt(len(last_3)) if std > 0 else avg * 0.7
            ci_upper = avg + 1.96 * std / np.sqrt(len(last_3)) if std > 0 else avg * 1.3
            # Determine trend based on last 3 months
            trend = 'Stable'
            if len(last_3) >= 2:
                trend_values = last_3.values
                if trend_values[-1] > trend_values[-2] * 1.1:
                    trend = 'Hausse'
                elif trend_values[-1] < trend_values[-2] * 0.9:
                    trend = 'Baisse'

            predictions = {
                'avg': avg,
                'ci_lower': ci_lower,
                'ci_upper': ci_upper,
                'reliability': reliability,
                'trend': trend
            }

            st.markdown(f"""
            <div style='color:#ffffff; font-size:14px;'>
                <p><strong>Estimation moyenne :</strong> {avg:,.0f} MAD/mois</p>
                <p><strong>Intervalle de confiance (95%) :</strong> {ci_lower:,.0f} - {ci_upper:,.0f} MAD</p>
                <p><strong>Fiabilité des prévisions :</strong> {reliability:.0f}%</p>
                <p><strong>Tendance récente :</strong> {trend}</p>
            </div>
            """, unsafe_allow_html=True)
            st.progress(int(reliability), "Fiabilité des prévisions")
            st.caption("Basé sur les 3 derniers mois et la variabilité des données")

            # Export predictions
            if st.button("💾 Exporter prévisions", key="export_predictions"):
                future_dates = pd.date_range(
                    start=engin_data['Date'].max() + pd.DateOffset(months=1),
                    periods=3, freq='M'
                )
                projections = pd.DataFrame({
                    'Mois': future_dates.strftime('%Y-%m'),
                    'Estimation Moyenne (MAD)': [avg]*3,
                    'Intervalle Bas (MAD)': [ci_lower]*3,
                    'Intervalle Haut (MAD)': [ci_upper]*3,
                    'Fiabilité (%)': [reliability]*3
                })
                csv = projections.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="Télécharger CSV",
                    data=csv,
                    file_name=f'previsions_{selected}.csv',
                    mime='text/csv'
                )
        else:
            st.markdown("""
            <div style='background-color:#ffebee; padding:10px; border-radius:5px;'>
                <p style='color:#d32f2f;'>⚠ Données insuffisantes (minimum 3 mois requis)</p>
            </div>
            """, unsafe_allow_html=True)

    # Export Section
    st.markdown("""
    <div style='background-color:#e3f2fd; padding:15px; border-radius:10px; border-left:5px solid #1976d2; margin-top:20px;'>
        <h3 style='color:#F28C38; margin-top:0;'>📄 Exporter les Analyses</h3>
    </div>
    """, unsafe_allow_html=True)

    if not engin_data.empty:
        # CSV Export
        csv = engin_data.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥 Télécharger CSV",
            data=csv,
            file_name=f'analyse_{selected}.csv',
            mime='text/csv'
        )
        # Word Export
        budget_threshold = 10000
        if st.button("📝 Exporter Rapport Word"):
            word_buffer = generate_word_report(
                engin_data, selected, figs, descriptions, metrics, predictions, budget_threshold
            )
            st.download_button(
                label="Télécharger Rapport Word",
                data=word_buffer,
                file_name=f'rapport_analyse_{selected}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

# Onglet 3 : Comparaisons
with tab3:
    st.markdown("""
    <div style='background-color:#e3f2fd; padding:20px; border-radius:10px; border-left:5px solid #1976d2; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>Comparaisons entre engins</h2>
        <p style='color:#424242;'>Analyse des coûts par engin et catégorie</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.plotly_chart(
        px.imshow(
            df.pivot_table(index='Engin_Formaté', columns='Desc_Cat', values='Montant', aggfunc='sum'),
            labels=dict(x="Catégorie", y="Engin", color="Coût"),
            title='Heatmap des coûts par engin et catégorie',
            aspect="auto",
            height=400
        ),
        use_container_width=True
    )
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        <div style='background-color:#fff8e1; padding:15px; border-radius:10px; border-left:5px solid #ffa000; margin-bottom:10px;'>
            <h3 style='color:#F28C38; margin-top:0;'>Coût total par engin</h3>
        </div>
        """, unsafe_allow_html=True)
        st.plotly_chart(
            px.bar(
                df.groupby('Engin_Formaté')['Montant'].sum().reset_index().sort_values('Montant'),
                x='Montant', y='Engin_Formaté',
                title='',
                height=400
            ),
            use_container_width=True
        )
    with col2:
        st.markdown("""
        <div style='background-color:#e8f5e9; padding:15px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:10px;'>
            <h3 style='color:#F28C38; margin-top:0;'>Distribution des coûts</h3>
        </div>
        """, unsafe_allow_html=True)
        st.plotly_chart(
            px.box(
                df, x='Engin_Formaté', y='Montant',
                title='',
                height=400
            ),
            use_container_width=True
        )

# Onglet 4 : Recommandations
with tab4:
    st.markdown("""
    <div style='background-color:#e3f2fd; padding:20px; border-radius:10px; border-left:5px solid #1976d2; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>🚀 Plan d'Action Simplifié</h2>
        <p style='color:#424242;'>Recommandations pratiques pour réduire les coûts</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        top_category = df.groupby('Desc_Cat')['Montant'].sum().idxmax()
        st.markdown(f"""
        <div style='background-color:#e3f2fd; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #1976d2;'>
            <h3 style='color:#F28C38; margin-top:0;'>🔍 Top 3 des Dépenses à Surveiller</h3>
            <ol style='color:#424242;'>
                <li style='margin-bottom:10px;'><b>{top_category}</b> <span style='color:#d32f2f; font-weight:bold;'>{df[df['Desc_Cat']==top_category]['Montant'].sum():,.0f} MAD</span></li>
                <li style='margin-bottom:10px;'><b>{df.groupby('Desc_Cat')['Montant'].sum().nlargest(2).index[1]}</b> <span style='color:#d32f2f; font-weight:bold;'>{df.groupby('Desc_Cat')['Montant'].sum().nlargest(2).values[1]:,.0f} MAD</span></li>
                <li><b>{df.groupby('Desc_Cat')['Montant'].sum().nlargest(3).index[2]}</b> <span style='color:#d32f2f; font-weight:bold;'>{df.groupby('Desc_Cat')['Montant'].sum().nlargest(3).values[2]:,.0f} MAD</span></li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        problem_engines = df.groupby('Engin_Formaté')['Montant'].sum().nlargest(3)
        st.markdown(f"""
        <div style='background-color:#fff8e1; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #ffa000;'>
            <h3 style='color:#F28C38; margin-top:0;'>🚜 Engins Prioritaires</h3>
            <div style='display:flex; align-items:center; margin-bottom:10px; color:#424242;'>
                <div style='background-color:#d32f2f; width:20px; height:20px; border-radius:50%; margin-right:10px;'></div>
                <div>Engin R1600-{problem_engines.index[0].split('-')[-1]} <span style='color:#d32f2f; font-weight:bold;'>- {problem_engines.values[0]:,.0f} MAD</span></div>
            </div>
            <div style='display:flex; align-items:center; margin-bottom:10px; color:#424242;'>
                <div style='background-color:#ffa000; width:20px; height:20px; border-radius:50%; margin-right:10px;'></div>
                <div>Engin R1600-{problem_engines.index[1].split('-')[-1]} <span style='color:#d32f2f; font-weight:bold;'>- {problem_engines.values[1]:,.0f} MAD</span></div>
            </div>
            <div style='display:flex; align-items:center; color:#424242;'>
                <div style='background-color:#fbc02d; width:20px; height:20px; border-radius:50%; margin-right:10px;'></div>
                <div>Engin R1600-{problem_engines.index[2].split('-')[-1]} <span style='color:#d32f2f; font-weight:bold;'>- {problem_engines.values[2]:,.0f} MAD</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("""
    <div style='background-color:#e8f5e9; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #388e3c;'>
        <h3 style='color:#F28C38; margin-top:0;'>📅 Plan d'Action sur 3 Mois</h3>
        <div style='border-left:4px solid #388e3c; padding-left:20px;'>
            <div style='margin-bottom:15px; color:#424242;'>
                <h4 style='margin-bottom:5px; color:#2e7d32;'>Mois 1 : Audit Initial</h4>
                <p>• Vérifier les 3 engins les plus coûteux<br>• Analyser les pièces les plus remplacées</p>
            </div>
            <div style='margin-bottom:15px; color:#424242;'>
                <h4 style='margin-bottom:5px; color:#2e7d32;'>Mois 2 : Négociations</h4>
                <p>• Contacter fournisseurs pour remises volume<br>• Standardiser les pièces communes</p>
            </div>
            <div style='color:#424242;'>
                <h4 style='margin-bottom:5px; color:#2e7d32;'>Mois 3 : Optimisation</h4>
                <p>• Mettre en place maintenance préventive<br>• Former les opérateurs aux bonnes pratiques</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div style='background-color:#f3e5f5; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #8e24aa;'>
        <h3 style='color:#F28C38; margin-top:0;'>✅ Checklist des Actions Clés</h3>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Identifier les 5 pièces les plus changées</div>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Comparer les coûts avec les standards du marché (en MAD)</div>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Réaliser un diagnostic des engins prioritaires</div>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Organiser une réunion avec les fournisseurs</div>
        <div style='color:#424242;'><input type='checkbox' style='margin-right:10px;'> Mettre en place un suivi mensuel des coûts (MAD)</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown(f"""
    <div style='background-color:#ffebee; padding:20px; border-radius:10px; border-left:5px solid #d32f2f;'>
        <h3 style='color:#F28C38; text-align:center; margin-top:0;'>💡 3 Astuces pour Réduire les Coûts (MAD)</h3>
        <div style='display:flex; justify-content:space-between; text-align:center; margin-top:20px;'>
            <div style='width:30%;'>
                <div style='background-color:#ffcdd2; padding:15px; border-radius:10px; height:120px;'>
                    <h4 style='color:#c62828;'>🛠️ Maintenance</h4>
                    <p style='color:#424242;'>Économie potentielle :<br><b>{df['Montant'].mean()*0.25:,.0f} MAD/mois</b></p>
                </div>
            </div>
            <div style='width:30%;'>
                <div style='background-color:#c8e6c9; padding:15px; border-radius:10px; height:120px;'>
                    <h4 style='color:#2e7d32;'>🔄 Pièces Standard</h4>
                    <p style='color:#424242;'>Économie potentielle :<br><b>{df['Montant'].mean()*0.40:,.0f} MAD/mois</b></p>
                </div>
            </div>
            <div style='width:30%;'>
                <div style='background-color:#bbdefb; padding:15px; border-radius:10px; height:120px;'>
                    <h4 style='color:#1565c0;'>📊 Suivi Rigoureux</h4>
                    <p style='color:#424242;'>Économie potentielle :<br><b>{df['Montant'].mean()*0.15:,.0f} MAD/mois</b></p>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Chatbot Functionality
def toggle_chat():
    st.session_state.chat_open = not st.session_state.chat_open

# Floating chatbot icon (using st.button instead of JavaScript)
if st.button("💬", key="chat_toggle", help="Ouvrir/Fermer le chatbot"):
    toggle_chat()

# Chatbot container
chat_class = "chatbot-container open" if st.session_state.chat_open else "chatbot-container"
with st.container():
    st.markdown(f'<div class="{chat_class}">', unsafe_allow_html=True)
    st.markdown('<div class="chatbot-header">Chatbot R1600</div>', unsafe_allow_html=True)
    
    # Messages container
    with st.container():
        st.markdown('<div class="chatbot-messages">', unsafe_allow_html=True)
        for message in st.session_state.chat_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Input container
    with st.container():
        st.markdown('<div class="chatbot-input">', unsafe_allow_html=True)
        if prompt := st.chat_input("Posez votre question...", key="chat_input"):
            # Append user message
            st.session_state.chat_messages.append({"role": "user", "content": prompt})
            # Generate and append assistant response
            response = generate_response(prompt, df)
            st.session_state.chat_messages.append({"role": "assistant", "content": response})
            # Rerun to update the chat display
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
