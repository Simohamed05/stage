import streamlit as st
import pandas as pd
import plotly.express as px
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import plotly.io as pio
import json
import bcrypt
import io
import os
import zipfile
import sys
import locale

# Configuration de la page
st.set_page_config(page_title="Tableau de bord de la consommation des équipements miniers", layout="wide")
st.markdown("""
    <style>
    .stApp { 
        background-color: #f5f7fa;
        background-image: none;
    }
    .main-container {
        background-color: white;
        padding: 25px; 
        border-radius: 12px; 
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        margin-bottom: 20px;
    }
    h1, h2, h3, h4, h5, h6 { 
        color: #2c3e50; 
        font-family: 'Segoe UI', Arial, sans-serif;
    }
    h1 {
        border-bottom: 2px solid #3498db;
        padding-bottom: 10px;
    }
    .metric-card {
        background-color: white;
        border-left: 4px solid #3498db; 
        padding: 18px; 
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        margin-bottom: 15px;
    }
    .metric-title {
        color: #7f8c8d;
        font-size: 18px;
        margin-bottom: 5px;
    }
    .metric-value {
        color: #2c3e50;
        font-size: 30px;
        font-weight: bold;
    }
    .stButton>button {
        background-color: #3498db; 
        color: white; 
        border-radius: 8px; 
        border: none;
        padding: 8px 18px;
        font-weight: 500;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        background-color: #2980b9; 
        color: white; 
        box-shadow: 0 4px 8px rgba(41,128,185,0.2);
        transform: translateY(-1px);
    }
    .css-1d391kg {
        background-color: white;
        box-shadow: 2px 0 15px rgba(0,0,0,0.05);
    }
    .sidebar .sidebar-content {
        background-color: white;
    }
    .stTabs [role="tablist"] button {
        color: #7f8c8d;
        font-weight: 500;
        padding: 8px 16px;
    }
    .stTabs [role="tablist"] button[aria-selected="true"] {
        color: #3498db;
        border-bottom: 3px solid #3498db;
        background-color: rgba(52,152,219,0.1);
    }
    .stDataFrame {
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    .stTextInput>div>div>input, 
    .stSelectbox>div>div>select,
    .stDateInput>div>div>input,
    .stMultiSelect>div>div>select {
        border: 1px solid #dfe6e9;
        border-radius: 8px;
        padding: 10px 12px;
    }
    .primary-color {
        color: #3498db;
    }
    .secondary-color {
        color: #2c3e50;
    }
    .accent-color {
        color: #e74c3c;
    }
    .header-container {
        background: linear-gradient(135deg, #3498db 0%, #2c3e50 100%);
        padding: 25px;
        border-radius: 10px;
        margin-bottom: 25px;
        color: white;
    }
    .analysis-card {
        background: white;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        margin-bottom: 20px;
        border-top: 4px solid #3498db;
    }
    </style>
""", unsafe_allow_html=True)

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def load_users():
    file_path = resource_path("users.json")
    if os.path.exists(file_path):
        with open(file_path, "r") as f:
            return json.load(f)
    return {}

def save_users(users):
    file_path = resource_path("users.json")
    with open(file_path, "w") as f:
        json.dump(users, f)

def hash_password(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def check_password(password, hashed):
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def load_data(uploaded_files=None):
    try:
        if uploaded_files is None or not uploaded_files:
            return pd.DataFrame()

        dfs = []
        required_columns = ['Date', 'CATEGORIE', 'Desc_Cat', 'Desc_CA', 'Montant']
        max_file_size = 200 * 1024 * 1024  # 200 Mo en octets

        for uploaded_file in uploaded_files:
            st.write(f"Traitement du fichier : {uploaded_file.name}, Taille : {uploaded_file.size / 1024 / 1024:.2f} Mo, Type : {'ZIP' if uploaded_file.name.endswith('.zip') else 'Excel'}")
            if uploaded_file.size > max_file_size:
                st.warning(f"Le fichier {uploaded_file.name} dépasse la limite de 200 Mo et sera ignoré.")
                continue
            try:
                uploaded_file.seek(0)
            except Exception as e:
                st.warning(f"Erreur lors de la réinitialisation du pointeur pour {uploaded_file.name} : {str(e)}. Ce fichier sera ignoré.")
                continue

            if uploaded_file.name.endswith('.zip'):
                try:
                    file_bytes = uploaded_file.read()
                    if not file_bytes:
                        st.warning(f"Le fichier ZIP {uploaded_file.name} est vide et sera ignoré.")
                        continue
                    file_stream = io.BytesIO(file_bytes)
                    with zipfile.ZipFile(file_stream, 'r') as z:
                        for filename in z.namelist():
                            if filename.endswith('.xlsx'):
                                with z.open(filename) as f:
                                    try:
                                        df = pd.read_excel(f)
                                        if not all(col in df.columns for col in required_columns):
                                            st.warning(f"Le fichier {filename} dans le ZIP {uploaded_file.name} ne contient pas toutes les colonnes requises : {', '.join(required_columns)}. Il sera ignoré.")
                                            continue
                                        df['CATEGORIE'] = df['CATEGORIE'].astype(str).replace('nan', 'Unknown')
                                        if pd.api.types.is_numeric_dtype(df['Date']):
                                            df['Date'] = pd.to_datetime(df['Date'], origin='1899-12-30', unit='D')
                                        elif not pd.api.types.is_datetime64_any_dtype(df['Date']):
                                            df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
                                        df['Montant'] = df['Montant'].astype(str).str.replace(r'[^\d.,]', '', regex=True)
                                        df['Montant'] = df['Montant'].str.replace(',', '.', regex=False)
                                        df['Montant'] = pd.to_numeric(df['Montant'], errors='coerce')
                                        initial_rows = df.shape[0]
                                        df = df.dropna(subset=['Montant'])
                                        dropped_rows = initial_rows - df.shape[0]
                                        df['Mois'] = df['Date'].dt.month_name()
                                        months_fr = {
                                            'January': 'Janvier', 'February': 'Février', 'March': 'Mars',
                                            'April': 'Avril', 'May': 'Mai', 'June': 'Juin',
                                            'July': 'Juillet', 'August': 'Août', 'September': 'Septembre',
                                            'October': 'Octobre', 'November': 'Novembre', 'December': 'Décembre'
                                        }
                                        df['Mois'] = df['Mois'].map(months_fr)
                                        dfs.append(df)
                                    except Exception as e:
                                        st.warning(f"Erreur lors du chargement du fichier {filename} dans le ZIP {uploaded_file.name} : {str(e)}")
                                        continue
                except zipfile.BadZipFile:
                    st.warning(f"Le fichier {uploaded_file.name} n'est pas un fichier ZIP valide et sera ignoré.")
                    continue
                except Exception as e:
                    st.warning(f"Erreur lors du traitement du fichier ZIP {uploaded_file.name} : {str(e)}")
                    continue
            else:
                try:
                    df = pd.read_excel(uploaded_file)
                    if not all(col in df.columns for col in required_columns):
                        st.warning(f"Le fichier {uploaded_file.name} ne contient pas toutes les colonnes requises : {', '.join(required_columns)}. Il sera ignoré.")
                        continue
                    df['CATEGORIE'] = df['CATEGORIE'].astype(str).replace('nan', 'Unknown')
                    if pd.api.types.is_numeric_dtype(df['Date']):
                        df['Date'] = pd.to_datetime(df['Date'], origin='1899-12-30', unit='D')
                    elif not pd.api.types.is_datetime64_any_dtype(df['Date']):
                        df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
                    df['Montant'] = df['Montant'].astype(str).str.replace(r'[^\d.,]', '', regex=True)
                    df['Montant'] = df['Montant'].str.replace(',', '.', regex=False)
                    df['Montant'] = pd.to_numeric(df['Montant'], errors='coerce')
                    initial_rows = df.shape[0]
                    df = df.dropna(subset=['Montant'])
                    dropped_rows = initial_rows - df.shape[0]
                    df['Mois'] = df['Date'].dt.month_name()
                    months_fr = {
                        'January': 'Janvier', 'February': 'Février', 'March': 'Mars',
                        'April': 'Avril', 'May': 'Mai', 'June': 'Juin',
                        'July': 'Juillet', 'August': 'Août', 'September': 'Septembre',
                        'October': 'Octobre', 'November': 'Novembre', 'December': 'Décembre'
                    }
                    df['Mois'] = df['Mois'].map(months_fr)
                    dfs.append(df)
                except Exception as e:
                    st.warning(f"Erreur lors du chargement du fichier {uploaded_file.name} : {str(e)}. Ce fichier sera ignoré.")
                    continue

        if not dfs:
            st.error("Aucun fichier valide n'a pu être chargé. Veuillez vérifier les fichiers téléversés.")
            return pd.DataFrame()
        
        combined_df = pd.concat(dfs, ignore_index=True)
        combined_df['CATEGORIE'] = combined_df['CATEGORIE'].astype(str).replace('nan', 'Unknown')
        combined_df = combined_df[combined_df['CATEGORIE'].str.upper().isin(['DUMPER', 'FORATION', '10 TONNES'])]
        
        if combined_df.empty:
            st.error("Aucune donnée pour les catégories DUMPER, FORATION, ou 10 TONNES. Vérifiez les fichiers téléversés.")
            return pd.DataFrame()
        
        st.success(f"{len(dfs)} fichier(s) valide(s) chargé(s). Nombre total de lignes après filtrage : {combined_df.shape[0]}")
        return combined_df
    except Exception as e:
        st.error(f"Erreur générale lors du chargement des fichiers : {str(e)}")
        return pd.DataFrame()

def load_tonnage_data(uploaded_files=None):
    try:
        if uploaded_files is None or not uploaded_files:
            st.warning("Aucun fichier de tonnage téléversé. Veuillez importer un ou plusieurs fichiers Excel ou ZIP.")
            return pd.DataFrame()

        dfs = []
        required_columns = ['DATE', 'DS Sud', 'DS Nord', 'KA']
        max_file_size = 200 * 1024 * 1024  # 200 Mo en octets

        if not isinstance(uploaded_files, (list, tuple)):
            st.error(f"Erreur : uploaded_files doit être une liste ou un tuple, reçu : {type(uploaded_files)}")
            return pd.DataFrame()

        for uploaded_file in uploaded_files:
            if not hasattr(uploaded_file, 'name') or not hasattr(uploaded_file, 'read'):
                st.warning(f"Élément invalide dans uploaded_files : {type(uploaded_file)}. Cet élément sera ignoré.")
                continue

            st.write(f"Traitement du fichier de tonnage : {uploaded_file.name}, Taille : {uploaded_file.size / 1024 / 1024:.2f} Mo, Type : {'ZIP' if uploaded_file.name.endswith('.zip') else 'Excel'}")
            if uploaded_file.size > max_file_size:
                st.warning(f"Le fichier {uploaded_file.name} dépasse la limite de 200 Mo et sera ignoré.")
                continue
            try:
                uploaded_file.seek(0)
            except Exception as e:
                st.warning(f"Erreur lors de la réinitialisation du pointeur pour {uploaded_file.name} : {str(e)}. Ce fichier sera ignoré.")
                continue

            if uploaded_file.name.endswith('.zip'):
                try:
                    file_bytes = uploaded_file.read()
                    if not file_bytes:
                        st.warning(f"Le fichier ZIP {uploaded_file.name} est vide et sera ignoré.")
                        continue
                    if not isinstance(file_bytes, bytes):
                        st.warning(f"Le contenu lu de {uploaded_file.name} n'est pas un objet bytes : {type(file_bytes)}. Ce fichier sera ignoré.")
                        continue
                    file_stream = io.BytesIO(file_bytes)
                    with zipfile.ZipFile(file_stream, 'r') as z:
                        for filename in z.namelist():
                            if filename.endswith('.xlsx'):
                                with z.open(filename) as f:
                                    try:
                                        df = pd.read_excel(f)
                                        if not all(col in df.columns for col in required_columns):
                                            st.warning(f"Le fichier {filename} dans le ZIP {uploaded_file.name} ne contient pas toutes les colonnes requises : {', '.join(required_columns)}. Il sera ignoré.")
                                            continue
                                        if pd.api.types.is_numeric_dtype(df['DATE']):
                                            df['DATE'] = pd.to_datetime(df['DATE'], origin='1899-12-30', unit='D')
                                        elif not pd.api.types.is_datetime64_any_dtype(df['DATE']):
                                            df['DATE'] = pd.to_datetime(df['DATE'])
                                        df = df.dropna(subset=required_columns)
                                        for col in ['DS Sud', 'DS Nord', 'KA']:
                                            df[col] = pd.to_numeric(df[col], errors='coerce')
                                        df['CUMMULE'] = df[['DS Sud', 'DS Nord', 'KA']].sum(axis=1)
                                        dfs.append(df)
                                    except Exception as e:
                                        st.warning(f"Erreur lors du chargement du fichier {filename} dans le ZIP {uploaded_file.name} : {str(e)}")
                                        continue
                except zipfile.BadZipFile:
                    st.warning(f"Le fichier {uploaded_file.name} n'est pas un fichier ZIP valide et sera ignoré.")
                    continue
                except Exception as e:
                    st.warning(f"Erreur lors du traitement du fichier ZIP {uploaded_file.name} : {str(e)}")
                    continue
            else:
                try:
                    df = pd.read_excel(uploaded_file)
                    if not all(col in df.columns for col in required_columns):
                        st.warning(f"Le fichier {uploaded_file.name} ne contient pas toutes les colonnes requises : {', '.join(required_columns)}. Il sera ignoré.")
                        continue
                    if pd.api.types.is_numeric_dtype(df['DATE']):
                        df['DATE'] = pd.to_datetime(df['DATE'], origin='1899-12-30', unit='D')
                    elif not pd.api.types.is_datetime64_any_dtype(df['DATE']):
                        df['DATE'] = pd.to_datetime(df['DATE'])
                    df = df.dropna(subset=required_columns)
                    for col in ['DS Sud', 'DS Nord', 'KA']:
                        df[col] = pd.to_numeric(df[col], errors='coerce')
                    df['CUMMULE'] = df[['DS Sud', 'DS Nord', 'KA']].sum(axis=1)
                    dfs.append(df)
                except Exception as e:
                    st.warning(f"Erreur lors du chargement du fichier {uploaded_file.name} : {str(e)}. Ce fichier sera ignoré.")
                    continue

        if not dfs:
            st.error("Aucun fichier de tonnage valide n'a pu être chargé. Veuillez vérifier les fichiers téléversés.")
            return pd.DataFrame()
        
        combined_df = pd.concat(dfs, ignore_index=True)
        combined_df = combined_df.drop_duplicates()
        
        st.success(f"{len(dfs)} fichier(s) de tonnage valide(s) chargé(s) avec succès. Nombre total de lignes : {combined_df.shape[0]}")
        return combined_df
    except Exception as e:
        st.error(f"Erreur générale lors du chargement des fichiers de tonnage : {str(e)}")
        return pd.DataFrame()
def load_hm_data(uploaded_files=None):
    try:
        if uploaded_files is None or not uploaded_files:
            st.warning("Aucun fichier d'heures de marche téléversé.")
            return pd.DataFrame()

        dfs = []
        required_columns = ['ENGINS']
        max_file_size = 200 * 1024 * 1024  # 200 MB

        for uploaded_file in uploaded_files:
            if not hasattr(uploaded_file, 'name') or not hasattr(uploaded_file, 'read'):
                st.warning(f"Élément invalide : {type(uploaded_file)}")
                continue

            if uploaded_file.size > max_file_size:
                st.warning(f"Le fichier {uploaded_file.name} dépasse la limite de 200 MB.")
                continue
            uploaded_file.seek(0)

            if uploaded_file.name.endswith('.zip'):
                try:
                    file_bytes = uploaded_file.read()
                    if not file_bytes:
                        st.warning(f"Le fichier ZIP {uploaded_file.name} est vide.")
                        continue
                    file_stream = io.BytesIO(file_bytes)
                    with zipfile.ZipFile(file_stream, 'r') as z:
                        for filename in z.namelist():
                            if filename.endswith('.xlsx'):
                                with z.open(filename) as f:
                                    try:
                                        df = pd.read_excel(f)
                                        if not all(col in df.columns for col in required_columns):
                                            st.warning(f"{filename} manque la colonne 'ENGINS'.")
                                            continue
                                        if pd.api.types.is_numeric_dtype(df['ENGINS']):
                                            df['ENGINS'] = pd.to_datetime(df['ENGINS'], origin='1899-12-30', unit='D', errors='coerce')
                                        elif not pd.api.types.is_datetime64_any_dtype(df['ENGINS']):
                                            df['ENGINS'] = pd.to_datetime(df['ENGINS'], errors='coerce')
                                        df = df.dropna(subset=required_columns)
                                        for col in df.columns[1:]:
                                            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).round(0).astype(int)
                                        df['TOTAL_HOURS'] = df.iloc[:, 1:].sum(axis=1, skipna=True).round(0).astype(int)
                                        dfs.append(df)
                                    except Exception as e:
                                        st.warning(f"Erreur dans {filename} : {str(e)}")
                                        continue
                except zipfile.BadZipFile:
                    st.warning(f"Le fichier {uploaded_file.name} n'est pas un fichier ZIP valide.")
                    continue
                except Exception as e:
                    st.warning(f"Erreur dans le ZIP {uploaded_file.name} : {str(e)}")
                    continue
            else:
                try:
                    df = pd.read_excel(uploaded_file)
                    if not all(col in df.columns for col in required_columns):
                        st.warning(f"{uploaded_file.name} manque la colonne 'ENGINS'.")
                        continue
                    if pd.api.types.is_numeric_dtype(df['ENGINS']):
                        df['ENGINS'] = pd.to_datetime(df['ENGINS'], origin='1899-12-30', unit='D', errors='coerce')
                    elif not pd.api.types.is_datetime64_any_dtype(df['ENGINS']):
                        df['ENGINS'] = pd.to_datetime(df['ENGINS'], errors='coerce')
                    df = df.dropna(subset=required_columns)
                    for col in df.columns[1:]:
                        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).round(0).astype(int)
                    df['TOTAL_HOURS'] = df.iloc[:, 1:].sum(axis=1, skipna=True).round(0).astype(int)
                    dfs.append(df)
                except Exception as e:
                    st.warning(f"Erreur dans {uploaded_file.name} : {str(e)}")
                    continue

        if not dfs:
            st.error("Aucun fichier d'heures de marche valide chargé.")
            return pd.DataFrame()

        combined_df = pd.concat(dfs, ignore_index=True)
        combined_df = combined_df.drop_duplicates()
        st.success(f"{len(dfs)} fichier(s) chargé(s). Lignes totales : {combined_df.shape[0]}")
        return combined_df
    except Exception as e:
        st.error(f"Erreur générale : {str(e)}")
        return pd.DataFrame()
    
def compute_monthly_costs(data):
    monthly_data = data.groupby('Mois')['Montant'].sum().reset_index()
    month_order = ['Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin',
                   'Juillet', 'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre']
    monthly_data['Mois'] = pd.Categorical(monthly_data['Mois'], categories=month_order, ordered=True)
    return monthly_data.sort_values('Mois')

def compute_category_breakdown(data):
    return data.groupby('Desc_Cat')['Montant'].sum().reset_index()

def generate_word_report(filtered_data, total_cost, global_avg, category_stats, most_consumed_per_cat, 
                        pivot_engine, selected_engines, table_df, total_montant, figures, tonnage_df, tonnage_date_range, hm_df, hm_date_range):
    doc = Document()
    
    title = doc.add_heading('Rapport Complet de Consommation des Équipements Miniers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date de génération: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(f"Période couverte: du {filtered_data['Date'].min().strftime('%d/%m/%Y')} au {filtered_data['Date'].max().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Nombre d'équipements analysés: {filtered_data['Desc_CA'].nunique()}")
    
    doc.add_heading('Table des Matières', level=1)
    doc.add_paragraph('1. Indicateurs Clés\n2. Analyse par Catégorie\n3. Analyse Comparative\n4. Données Détailées\n5. Recommandations\n6. Analyse des Tonnages\n7. Analyse des Heures de Marche', style='ListBullet')
    
    doc.add_heading('1. Indicateurs Clés', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'LightShading'
    table.cell(0, 0).text = 'Indicateur'
    table.cell(0, 1).text = 'Valeur'
    table.cell(1, 0).text = 'Coût total'
    table.cell(1, 1).text = f"{total_cost:,.0f} DH"
    table.cell(2, 0).text = 'Moyenne globale par jour'
    table.cell(2, 1).text = f"{global_avg:,.0f} DH"
    
    doc.add_heading('Indicateurs par Catégorie', level=2)
    cat_table = doc.add_table(rows=category_stats.shape[0]+1, cols=4)
    cat_table.style = 'LightShading'
    cat_table.cell(0, 0).text = 'Catégorie'
    cat_table.cell(0, 1).text = 'Total (DH)'
    cat_table.cell(0, 2).text = 'Moyenne (DH)'
    cat_table.cell(0, 3).text = 'Type le plus consommé'
    
    for i, (_, row) in enumerate(category_stats.iterrows()):
        most_consumed = most_consumed_per_cat[most_consumed_per_cat['CATEGORIE'] == row['CATEGORIE']]
        most_consumed_desc = most_consumed['Desc_Cat'].iloc[0] if not most_consumed.empty else "N/A"
        
        cat_table.cell(i+1, 0).text = row['CATEGORIE']
        cat_table.cell(i+1, 1).text = f"{row['Total']:,.0f}"
        cat_table.cell(i+1, 2).text = f"{row['Moyenne']:,.0f}"
        cat_table.cell(i+1, 3).text = most_consumed_desc
    
    doc.add_heading('2. Analyse par Catégorie', level=1)
    doc.add_paragraph('Cette section présente les analyses détaillées pour chaque catégorie d\'équipement.')
    
    progress_bar = st.progress(0)
    total_figures = sum(1 for fig_name in figures if "Consommation par équipement" in fig_name) + ("Coût total par catégorie" in figures)
    
    for i, (fig_name, fig) in enumerate(figures.items()):
        if "Consommation par équipement" in fig_name:
            doc.add_heading(fig_name, level=2)
            category = fig_name.split('(')[-1].replace(')', '')
            doc.add_paragraph(f"Ce graphique montre la répartition des coûts par équipement pour la catégorie {category}. "
                            "Il permet d'identifier les équipements les plus coûteux à maintenir.")
            
            img_bytes = pio.to_image(fig, format='png', scale=1)
            doc.add_picture(BytesIO(img_bytes), width=Inches(6))
            progress_bar.progress((i + 1) / total_figures)
    
    doc.add_heading('3. Analyse Comparative', level=1)
    doc.add_paragraph('Comparaison des performances entre les différentes catégories d\'équipements.')
    
    if "Coût total par catégorie" in figures:
        doc.add_heading('Comparaison des coûts par catégorie', level=2)
        doc.add_paragraph("Ce graphique compare les coûts totaux entre les différentes catégories d'équipements. "
                        "Les catégories les plus à droite représentent les postes de dépenses les plus importants.")
        
        img_bytes = pio.to_image(figures["Coût total par catégorie"], format='png', scale=1)
        doc.add_picture(BytesIO(img_bytes), width=Inches(6))
        progress_bar.progress(1.0)
    
    doc.add_heading('4. Données Détailées', level=1)
    
    if not pivot_engine.empty:
        doc.add_heading(f'Détail des consommations pour {", ".join(selected_engines) if selected_engines else "toutes les catégories"}', level=2)
        doc.add_paragraph(f"Tableau détaillant les différents types de consommation pour chaque équipement des catégories sélectionnées.")
        
        table = doc.add_table(rows=pivot_engine.shape[0]+1, cols=pivot_engine.shape[1]+1)
        table.style = 'Table Grid'
        
        table_rows = table.rows
        table_rows[0].cells[0].text = 'Équipement'
        for j, col in enumerate(pivot_engine.columns):
            table_rows[0].cells[j+1].text = str(col)
        
        for i, (index, row) in enumerate(pivot_engine.iterrows()):
            row_cells = table_rows[i+1].cells
            row_cells[0].text = str(index)
            for j, value in enumerate(row):
                row_cells[j+1].text = f"{value:,.2f} DH"
    
    doc.add_heading('Journal complet des consommations', level=2)
    doc.add_paragraph('Liste détaillée des consommations enregistrées (limité aux 100 premières entrées).')
    
    max_rows = min(table_df.shape[0], 100)
    table = doc.add_table(rows=max_rows+2, cols=table_df.shape[1])
    table.style = 'Table Grid'
    
    table_rows = table.rows
    for j, col in enumerate(table_df.columns):
        table_rows[0].cells[j].text = col
    
    for i in range(max_rows):
        row_cells = table_rows[i+1].cells
        for j, value in enumerate(table_df.iloc[i]):
            row_cells[j].text = str(value)
    
    table_rows[max_rows+1].cells[0].text = 'Total'
    table_rows[max_rows+1].cells[table_df.shape[1]-1].text = f"{total_montant:,.2f} DH"
    
    doc.add_heading('5. Recommandations', level=1)
    
    top_categories = filtered_data.groupby('CATEGORIE')['Montant'].sum().nlargest(3).reset_index()
    doc.add_heading('Catégories prioritaires', level=2)
    for _, row in top_categories.iterrows():
        doc.add_paragraph(
            f"{row['CATEGORIE']}: {row['Montant']:,.0f} DH ({(row['Montant']/total_cost)*100:.1f}% du total)",
            style='ListBullet'
        )
    
    doc.add_heading('Actions recommandées', level=2)
    recommendations = [
        "Prioriser les analyses des équipements dans les catégories les plus coûteuses",
        "Mettre en place un suivi mensuel des consommations par catégorie",
        "Comparer les performances des équipements similaires pour identifier les anomalies",
        "Négocier avec les fournisseurs pour les pièces les plus fréquemment remplacées",
        "Étudier la possibilité de maintenance préventive pour réduire les coûts",
        "Former les opérateurs à une utilisation optimale des équipements"
    ]
    for rec in recommendations:
        doc.add_paragraph(rec, style='ListBullet')
    
    doc.add_heading('6. Analyse des Tonnages', level=1)
    doc.add_paragraph('Cette section présente les données de tonnage pour les sites DS Sud, DS Nord et KA.')
    
    if not tonnage_df.empty:
        filtered_tonnage_df = tonnage_df.copy()
        if tonnage_date_range is not None and len(tonnage_date_range) == 2:
            start_date, end_date = tonnage_date_range
            filtered_tonnage_df = filtered_tonnage_df[
                (filtered_tonnage_df['DATE'].dt.date >= start_date) & 
                (filtered_tonnage_df['DATE'].dt.date <= end_date)
            ]
        else:
            doc.add_paragraph("Plage de dates non définie pour les tonnages. Affichage de toutes les données disponibles.")
        if not filtered_tonnage_df.empty:
            doc.add_heading('Tableau des tonnages', level=2)
            max_rows = min(filtered_tonnage_df.shape[0], 100)
            table = doc.add_table(rows=max_rows+2, cols=5)
            table.style = 'Table Grid'
            
            table_rows = table.rows
            headers = ['Date', 'DS Sud (T)', 'DS Nord (T)', 'KA (T)', 'Cumulé (T)']
            for j, col in enumerate(headers):
                table_rows[0].cells[j].text = col
            
            display_tonnage_df = filtered_tonnage_df[['DATE', 'DS Sud', 'DS Nord', 'KA', 'CUMMULE']].copy()
            display_tonnage_df['DATE'] = display_tonnage_df['DATE'].dt.strftime('%d/%m/%Y')
            
            for i in range(max_rows):
                row_cells = table_rows[i+1].cells
                for j, value in enumerate(display_tonnage_df.iloc[i]):
                    row_cells[j].text = str(value) if j == 0 else f"{value:,.2f} T"
            
            total_tonnage = display_tonnage_df[['DS Sud', 'DS Nord', 'KA']].sum().to_dict()
            total_cumule = display_tonnage_df['CUMMULE'].sum()
            table_rows[max_rows+1].cells[0].text = 'Total'
            table_rows[max_rows+1].cells[1].text = f"{total_tonnage['DS Sud']:,.2f} T"
            table_rows[max_rows+1].cells[2].text = f"{total_tonnage['DS Nord']:,.2f} T"
            table_rows[max_rows+1].cells[3].text = f"{total_tonnage['KA']:,.2f} T"
            table_rows[max_rows+1].cells[4].text = f"{total_cumule:,.2f} T"
            
            if "Comparaison des tonnages par site" in figures:
                doc.add_heading('Comparaison des tonnages par site', level=2)
                img_bytes = pio.to_image(figures["Comparaison des tonnages par site"], format='png', scale=1)
                doc.add_picture(BytesIO(img_bytes), width=Inches(6))
            
            if "Tonnage total par site" in figures:
                doc.add_heading('Tonnage total par site', level=2)
                img_bytes = pio.to_image(figures["Tonnage total par site"], format='png', scale=1)
                doc.add_picture(BytesIO(img_bytes), width=Inches(6))

        doc.add_heading('7. Analyse des Heures de Marche', level=1)
        doc.add_paragraph('Cette section présente les données des heures de marche pour les équipements miniers.')
        
        if not hm_df.empty:
            filtered_hm_df = hm_df.copy()
            if hm_date_range is not None and len(hm_date_range) == 2:
                start_date, end_date = hm_date_range
                filtered_hm_df = filtered_hm_df[
                    (filtered_hm_df['ENGINS'].dt.date >= start_date) & 
                    (filtered_hm_df['ENGINS'].dt.date <= end_date)
                ]
            else:
                doc.add_paragraph("Plage de dates non définie pour les heures de marche. Affichage de toutes les données disponibles.")
            if not filtered_hm_df.empty:
                doc.add_heading('Tableau des heures de marche', level=2)
                max_rows = min(filtered_hm_df.shape[0], 100)
                equipment_columns = [col for col in filtered_hm_df.columns if col not in ['ENGINS', 'TOTAL_HOURS']]
                table = doc.add_table(rows=max_rows+2, cols=len(equipment_columns)+2)
                table.style = 'Table Grid'
                
                table_rows = table.rows
                headers = ['Date'] + equipment_columns + ['Total (h)']
                for j, col in enumerate(headers):
                    table_rows[0].cells[j].text = col
                
                display_hm_df = filtered_hm_df[['ENGINS'] + equipment_columns + ['TOTAL_HOURS']].copy()
                display_hm_df['ENGINS'] = display_hm_df['ENGINS'].dt.strftime('%d/%m/%Y')
                
                for i in range(max_rows):
                    row_cells = table_rows[i+1].cells
                    for j, value in enumerate(display_hm_df.iloc[i]):
                        row_cells[j].text = str(value) if j == 0 else f"{value:,.2f} h"
                
                total_hours = display_hm_df[equipment_columns].sum().to_dict()
                total_sum = display_hm_df['TOTAL_HOURS'].sum()
                table_rows[max_rows+1].cells[0].text = 'Total'
                for j, col in enumerate(equipment_columns, 1):
                    table_rows[max_rows+1].cells[j].text = f"{total_hours[col]:,.2f} h"
                table_rows[max_rows+1].cells[-1].text = f"{total_sum:,.2f} h"
                
                if "Comparaison des heures de marche par équipement" in figures:
                    doc.add_heading('Comparaison des heures de marche par équipement', level=2)
                    img_bytes = pio.to_image(figures["Comparaison des heures de marche par équipement"], format='png', scale=1)
                    doc.add_picture(BytesIO(img_bytes), width=Inches(6))
                
                if "Heures totales par équipement" in figures:
                    doc.add_heading('Heures totales par équipement', level=2)
                    img_bytes = pio.to_image(figures["Heures totales par équipement"], format='png', scale=1)
                    doc.add_picture(BytesIO(img_bytes), width=Inches(6))
        else:
            doc.add_paragraph("Aucune donnée d'heures de marche disponible pour la période sélectionnée.")
        doc.add_heading('7. Analyse des Heures de Marche', level=1)
    doc.add_paragraph('Cette section présente les données des heures de marche pour les équipements miniers, incluant les totaux cumulés, ainsi qu’une analyse de rentabilité basée sur le rendement moyen (tonnes par heure) et le coût par tonne (DH par tonne).')

    if not hm_df.empty:
        filtered_hm_df = hm_df.copy()
        if hm_date_range is not None and len(hm_date_range) == 2:
            start_date, end_date = hm_date_range
            filtered_hm_df = filtered_hm_df[
                (filtered_hm_df['ENGINS'].dt.date >= start_date) & 
                (filtered_hm_df['ENGINS'].dt.date <= end_date)
            ]
        else:
            doc.add_paragraph("Plage de dates non définie pour les heures de marche. Affichage de toutes les données disponibles.")
        
        if not filtered_hm_df.empty:
            
            
            doc.add_heading('Totaux pour la période sélectionnée', level=2)
            filtered_data_df = filtered_data.copy()
            if date_range is not None and len(date_range) == 2:
                start_date, end_date = date_range
                filtered_data_df = filtered_data_df[
                    (filtered_data_df['Date'].dt.date >= start_date) &
                    (filtered_data_df['Date'].dt.date <= end_date)
                ]
            
            filtered_tonnage_df = tonnage_df.copy()
            if tonnage_date_range is not None and len(tonnage_date_range) == 2:
                start_date, end_date = tonnage_date_range
                filtered_tonnage_df = filtered_tonnage_df[
                    (filtered_tonnage_df['DATE'].dt.date >= start_date) &
                    (filtered_tonnage_df['DATE'].dt.date <= end_date)
                ]
            
            total_consumption = filtered_data_df['Montant'].sum() if not filtered_data_df.empty else 0
            total_tonnage = filtered_tonnage_df['CUMMULE'].sum() if not filtered_tonnage_df.empty else 0
            total_hours = filtered_hm_df['TOTAL_HOURS'].sum() if not filtered_hm_df.empty else 0
            
            table = doc.add_table(rows=4, cols=2)
            table.style = 'LightShading'
            table.cell(0, 0).text = 'Indicateur'
            table.cell(0, 1).text = 'Valeur'
            table.cell(1, 0).text = 'Consommation totale'
            table.cell(1, 1).text = f"{total_consumption:,.2f} DH"
            table.cell(2, 0).text = 'Tonnage total'
            table.cell(2, 1).text = f"{total_tonnage:,.2f} T"
            table.cell(3, 0).text = 'Heures totales'
            table.cell(3, 1).text = f"{total_hours:d} h"
            
            doc.add_heading('Analyse de rentabilité', level=2)
            if total_tonnage == 0 or total_hours == 0 or total_consumption == 0:
                doc.add_paragraph("Données insuffisantes pour réaliser l’analyse de rentabilité. Veuillez vérifier que les fichiers de consommation, de tonnage et d’heures de marche sont chargés pour la période sélectionnée.")
            else:
                average_yield = total_tonnage / total_hours if total_hours > 0 else 0
                cost_per_tonne = total_consumption / total_tonnage if total_tonnage > 0 else float('inf')
                YIELD_THRESHOLD = 10  # Tonnes par heure minimum
                COST_PER_TONNE_THRESHOLD = 500  # Coût maximum par tonne en DH
                
                table = doc.add_table(rows=3, cols=2)
                table.style = 'LightShading'
                table.cell(0, 0).text = 'Indicateur'
                table.cell(0, 1).text = 'Valeur'
                table.cell(1, 0).text = 'Rendement moyen'
                table.cell(1, 1).text = f"{average_yield:.2f} T/h"
                table.cell(2, 0).text = 'Coût par tonne'
                table.cell(2, 1).text = f"{cost_per_tonne:.2f} DH/T"
                
                if average_yield >= YIELD_THRESHOLD and cost_per_tonne <= COST_PER_TONNE_THRESHOLD:
                    doc.add_heading('Résultat : Opération Gagnante ✅', level=3)
                    doc.add_paragraph(
                        f"Le rendement moyen ({average_yield:.2f} T/h) est supérieur au seuil de {YIELD_THRESHOLD:.2f} T/h, "
                        f"et le coût par tonne ({cost_per_tonne:.2f} DH/T) est inférieur au seuil de {COST_PER_TONNE_THRESHOLD:.2f} DH/T. "
                        "L’opération est efficace et rentable."
                    )
                else:
                    doc.add_heading('Résultat : Opération Perdante ❌', level=3)
                    reasons = []
                    if average_yield < YIELD_THRESHOLD:
                        reasons.append(f"Le rendement moyen ({average_yield:.2f} T/h) est inférieur au seuil de {YIELD_THRESHOLD:.2f} T/h.")
                    if cost_per_tonne > COST_PER_TONNE_THRESHOLD:
                        reasons.append(f"Le coût par tonne ({cost_per_tonne:.2f} DH/T) dépasse le seuil de {COST_PER_TONNE_THRESHOLD:.2f} DH/T.")
                    
                    doc.add_paragraph("L’opération présente des inefficacités :")
                    for reason in reasons:
                        doc.add_paragraph(reason, style='ListBullet')
                    
                    doc.add_heading('Recommandations', level=4)
                    recommendations = [
                        "Optimiser l’utilisation des équipements pour augmenter le rendement horaire.",
                        "Réduire les coûts d’exploitation en négociant les prix des pièces ou en améliorant la maintenance préventive.",
                        "Analyser les équipements spécifiques pour identifier les sources de surconsommation."
                    ]
                    for rec in recommendations:
                        doc.add_paragraph(rec, style='ListBullet')
    else:
        doc.add_paragraph("Aucune donnée d'heures de marche disponible pour la période sélectionnée.")

    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        "Ce rapport fournit une analyse complète des coûts de consommation des équipements miniers, des tonnages des sites, "
        "et des heures de marche des équipements. Les graphiques et tableaux présentés permettent d'identifier les principaux "
        "post_modes de dépenses, les performances des sites, et l'utilisation des équipements pour optimiser les coûts d'exploitation "
        "et la productivité."
    )
    
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = f"Généré le {datetime.now().strftime('%d/%m/%Y')} - Tableau de bord de consommation et heures de marche des équipements miniers"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Initialiser l'état de la session
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.username = None
    st.session_state.page = 'login'

if 'file_uploader_key' not in st.session_state:
    st.session_state.file_uploader_key = 0
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None

if 'tonnage_file_uploader_key' not in st.session_state:
    st.session_state.tonnage_file_uploader_key = 0
if 'uploaded_tonnage_file' not in st.session_state:
    st.session_state.uploaded_tonnage_file = None
if 'tonnage_date_range' not in st.session_state:
    st.session_state.tonnage_date_range = (
        datetime(2025, 5, 1).date(),
        datetime(2025, 5, 19).date()
    )
# Add to the session state initialization section
if 'hm_file_uploader_key' not in st.session_state:
    st.session_state.hm_file_uploader_key = 0
if 'uploaded_hm_file' not in st.session_state:
    st.session_state.uploaded_hm_file = None
if 'hm_date_range' not in st.session_state:
    st.session_state.hm_date_range = (
        datetime(2025, 5, 1).date(),
        datetime(2025, 5, 19).date()
    )
# Interface de connexion/inscription
if not st.session_state.logged_in:
    st.markdown("""
    <div class='header-container'>
        <h1 style='color: white; text-align:center; margin-top:0;'>Bienvenue</h1>
        <p style='color: white; text-align:center;'>Veuillez vous connecter ou créer un compte pour accéder au tableau de bord</p>
    <style>
    /* Style spécifique pour la page de connexion */
    .stApp { 
        background-image: url('https://img.freepik.com/premium-photo/blue-white-abstract-background-with-flowing-particles_916626-5365.jpg');
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        position: relative;
        min-height: 100vh;
    }
    
    }
    </style>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Connexion", key="show_login"):
            st.session_state.page = 'login'
    with col2:
        if st.button("Inscription", key="show_signup"):
            st.session_state.page = 'signup'

    users = load_users()

    if st.session_state.page == 'login':
        st.subheader("Connexion")
        with st.form("login_form"):
            username = st.text_input("Nom d'utilisateur")
            password = st.text_input("Mot de passe", type="password")
            submit = st.form_submit_button("Se connecter")

            if submit:
                if username in users and check_password(password, users[username]):
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.success(f"Connecté en tant que {username}")
                    st.rerun()
                else:
                    st.error("Nom d'utilisateur ou mot de passe incorrect")

    elif st.session_state.page == 'signup':
        st.subheader("Inscription")
        with st.form("signup_form"):
            new_username = st.text_input("Nouveau nom d'utilisateur")
            new_password = st.text_input("Nouveau mot de passe", type="password")
            confirm_password = st.text_input("Confirmer le mot de passe", type="password")
            submit = st.form_submit_button("S'inscrire")

            if submit:
                if new_username in users:
                    st.error("Ce nom d'utilisateur existe déjà")
                elif new_password != confirm_password:
                    st.error("Les mots de passe ne correspondent pas")
                elif not new_username or not new_password:
                    st.error("Veuillez remplir tous les champs")
                else:
                    users[new_username] = hash_password(new_password)
                    save_users(users)
                    st.success("Compte créé avec succès ! Veuillez vous connecter.")
                    st.session_state.page = 'login'
                    st.rerun()


else:
    with st.sidebar:
        if st.session_state.logged_in:
            if st.button("🚪 Déconnexion", key="logout_button"):
                st.session_state.logged_in = False
                st.session_state.username = None
                st.session_state.page = 'login'
                st.session_state.uploaded_file = None
                st.session_state.uploaded_tonnage_file = None
                st.session_state.file_uploader_key += 1
                st.session_state.tonnage_file_uploader_key += 1
                st.session_state.selected_engines = []
                st.rerun()

    # Nouvelle section d'introduction avant l'import des données
    if not st.session_state.get('uploaded_file') and not st.session_state.get('uploaded_tonnage_file'):
    
        st.markdown("""
        <div class='header-container'>
            <h1 style='color: white; text-align:center; margin-top:0;'> Tableau De Bord De La Consommation Des Engins</h1>
            <p style='color: white; text-align: center; margin-bottom:0'>Suivre et optimiser la consommation des équipements</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <div class='analysis-card'>
            <h2 style='color: #2c3e50; margin-top:0;'>Bienvenue dans le Tableau de Bord de Gestion des Équipements Miniers</h2>
            <p style='color: #7f8c8d; font-size:16px;'>Ce tableau de bord interactif vous permet d'analyser et d'optimiser la consommation, les tonnages et les heures de marche des équipements miniers. Découvrez ci-dessous les principales fonctionnalités :</p>
            <h3 style='color: #3498db; margin-top:20px;'>Fonctionnalités principales :</h3>
            <ul style='color: #2c3e50;'>
                <li><strong>Analyse des Coûts par Catégorie et Équipement :</strong> Visualisez les coûts totaux et moyens par catégorie (Dumper, Foration, 10 Tonnes) et par équipement spécifique. Identifiez rapidement les équipements ou catégories où des optimisations sont possibles.</li>
                <li><strong>Filtrage Interactif des Données :</strong> Appliquez des filtres par dates, équipements ou type de consommation pour explorer les données pertinentes. Les tableaux et graphiques se mettent à jour dynamiquement.</li>
                <li><strong>Comparaison des Performances :</strong> Comparez les coûts entre différentes catégories ou équipements pour détecter les anomalies et identifier les opportunités d'optimisation.</li>
                <li><strong>Analyse des Tonnages par Sites :</strong> Visualisez les tendances au fil du temps et les totaux par site (DS Sud, DS Nord, KA) pour évaluer les performances de production.</li>
                <li><strong>Analyse des Heures de Marche :</strong> Examinez les heures de fonctionnement des équipements avec des totaux cumulés. Évaluez la rentabilité grâce à des indicateurs comme le rendement moyen (tonnes/heure) et le coût par tonne (DH/tonne), avec une analyse indiquant si l'opération est gagnante ou perdante.</li>
                <li><strong>Génération de Rapports Détailés :</strong> Exportez un rapport Word complet incluant des tableaux, graphiques, indicateurs de rentabilité et recommandations personnalisées pour une analyse approfondie et un partage facile.</li>
                <li><strong>Recommandations Actionnables :</strong> Recevez des recommandations basées sur les données pour réduire les coûts, améliorer la maintenance préventive et optimiser l'utilisation des équipements.</li>
            </ul>
            <h3 style='color: #3498db; margin-top:20px;'>Comment commencer :</h3>
            <ol style='color: #2c3e50;'>
                <li>Utilisez le panneau latéral gauche pour importer vos fichiers de consommation, de tonnage et d'heures de marche.</li>
                <li>Les fichiers doivent être au format Excel (.xlsx) ou ZIP (.zip).</li>
                <li>Pour les consommations, les colonnes requises sont : <code>Date</code>, <code>CATEGORIE</code>, <code>Desc_Cat</code>, <code>Desc_CA</code>, <code>Montant</code>.</li>
                <li>Pour les tonnages, les colonnes requises sont : <code>DATE</code>, <code>DS Sud</code>, <code>DS Nord</code>, <code>KA</code>.</li>
                <li>Pour les heures de marche, la colonne requise est : <code>ENGINS</code> (dates), avec des colonnes supplémentaires pour chaque équipement (heures).</li>
                <li>Une fois les fichiers chargés, les analyses seront automatiquement disponibles.</li>
            </ol>
            <div style='background-color: #e3f2fd; padding: 15px; border-radius: 8px; margin-top: 20px;'>
                <h4 style='color: #1565c0; margin-top:0;'>Conseil :</h4>
                <p style='color: #1565c0;'>
                    Pour une analyse complète, importez les fichiers de consommation, de tonnage et d'heures de marche. Cela permettra de générer des rapports détaillés incluant l'analyse de rentabilité basée sur le rendement et le coût par tonne.
                </p>
            </div>
        </div>
        """, unsafe_allow_html=True)
    with st.sidebar:            
        st.subheader("Importer des fichiers de consommation")
        st.markdown("**Note** : Plusieurs fichiers Excel (.xlsx) ou ZIP (.zip) peuvent être importés (max 200 Mo par fichier).")
        st.markdown("**Fichiers importés** :")
        if st.session_state.uploaded_file:
            st.write(", ".join([f.name for f in st.session_state.uploaded_file]))
        else:
            st.write("Aucun fichier importé.")

        with st.form("file_upload_form", clear_on_submit=True):
            uploaded_files = st.file_uploader(
                "Téléverser des fichiers Excel ou ZIP (max 200 Mo par fichier)",
                type=["xlsx", "zip"],
                accept_multiple_files=True,
                key=f"file_uploader_{st.session_state.file_uploader_key}"
            )
            submit_button = st.form_submit_button("Charger les fichiers")

            if submit_button:
                if uploaded_files:
                    st.session_state.uploaded_file = uploaded_files
                    st.session_state.file_uploader_key += 1
                    df = load_data(st.session_state.uploaded_file)
                    if not df.empty:
                        st.success(f"Fichiers chargés avec succès. Nombre total de lignes : {df.shape[0]}")
                    else:
                        st.warning("Aucun fichier valide n'a pu être chargé. Veuillez vérifier les fichiers téléversés.")
                        st.session_state.uploaded_file = None
                else:
                    st.warning("Aucun fichier sélectionné. Veuillez téléverser un ou plusieurs fichiers Excel ou ZIP.")
            else:
                df = load_data(st.session_state.uploaded_file)

        if df.empty:
            st.write(' les données sont incorrect vérifier les noms des colonnes.')
            st.stop()
        
        st.subheader("Filtres")
        st.subheader("Plage de dates")
        default_start = df['Date'].min().date() if not df.empty else datetime.today().date()
        default_end = df['Date'].max().date() if not df.empty else datetime.today().date()
        date_range = st.date_input(
            "Période",
            value=(default_start, default_end),
            min_value=default_start,
            max_value=default_end,
            help="Choisir une plage de dates pour filtrer les interventions"
        )
        st.subheader("Rechercher un équipement")
        equipment_search = st.text_input("Entrer le nom de l'équipement (correspondance partielle)", "").strip()
        if equipment_search:
            available_equipment = sorted(df[df['Desc_CA'].str.contains(equipment_search, case=False, na=False)]['Desc_CA'].unique())
        else:
            available_equipment = sorted(df['Desc_CA'].unique())
        equipment_options = ["Tous les équipements"] + available_equipment
        if not available_equipment:
            st.warning("Aucun équipement ne correspond au terme de recherche.")
        selected_equipment = st.selectbox("Sélectionner l'équipement", equipment_options)
        filtered_data = df.copy()
        st.write(f"Total Montant des données brutes : {filtered_data['Montant'].sum():,.2f} DH")
        if len(date_range) == 2:
            start_date, end_date = date_range
            filtered_data = filtered_data[(filtered_data['Date'].dt.date >= start_date) & 
                                        (filtered_data['Date'].dt.date <= end_date)]
            st.write(f"Total Montant après filtre de date : {filtered_data['Montant'].sum():,.2f} DH")

        if selected_equipment != "Tous les équipements":
            filtered_data = filtered_data[filtered_data['Desc_CA'] == selected_equipment]
            st.write(f"Total Montant après filtre d'équipement : {filtered_data['Montant'].sum():,.2f} DH")

        if filtered_data.empty:
            st.warning("Aucune donnée disponible après filtrage. Veuillez ajuster les filtres.")
            st.stop()
        total_cost = filtered_data['Montant'].sum()
        global_avg = filtered_data['Montant'].mean()
        category_stats = filtered_data.groupby('CATEGORIE').agg(
            Total=('Montant', 'sum'),
            Moyenne=('Montant', 'mean')
        ).reset_index()
        most_consumed_per_cat = filtered_data.groupby(['CATEGORIE', 'Desc_Cat'])['Montant'].sum().reset_index()
        most_consumed_per_cat = most_consumed_per_cat.loc[most_consumed_per_cat.groupby('CATEGORIE')['Montant'].idxmax()]
        
        st.subheader("Exportation")
        # Update the report generation block in the sidebar
        if st.button("📄 Générer un rapport Word complet"):
            with st.spinner("Génération du rapport en cours..."):
                figures = {}
                
                fig_comp = px.bar(
                    filtered_data.groupby('CATEGORIE')['Montant'].sum().reset_index(),
                    x='CATEGORIE',
                    y='Montant',
                    title='Coût total par catégorie',
                    height=400,
                    text='Montant'
                )
                fig_comp.update_traces(
                    texttemplate='%{text:,.0f} DH',
                    textposition='auto'
                )
                fig_comp.update_layout(
                    xaxis_title="Catégorie",
                    yaxis_title="Montant total (DH)",
                    template='plotly_white'
                )
                figures["Coût total par catégorie"] = fig_comp
                
                for cat in filtered_data['CATEGORIE'].unique():
                    cat_data = filtered_data[filtered_data['CATEGORIE'] == cat]
                    equip_sum = cat_data.groupby('Desc_CA')['Montant'].sum().reset_index().sort_values('Montant', ascending=False)
                    fig_cat = px.bar(
                        equip_sum,
                        x='Desc_CA',
                        y='Montant',
                        title=f'Consommation par équipement ({cat})',
                        height=400,
                        text='Montant'
                    )
                    fig_cat.update_traces(
                        texttemplate='%{text:,.0f} DH',
                        textposition='auto'
                    )
                    fig_cat.update_layout(
                        xaxis_title="Équipement",
                        yaxis_title="Montant total (DH)",
                        template='plotly_white',
                        xaxis={'categoryorder':'total descending'}
                    )
                    figures[f"Consommation par équipement ({cat})"] = fig_cat
                
                tonnage_df = load_tonnage_data(st.session_state.uploaded_tonnage_file)
                if not tonnage_df.empty:
                    filtered_tonnage_df = tonnage_df.copy()
                    if st.session_state['tonnage_date_range'] is not None and len(st.session_state['tonnage_date_range']) == 2:
                        start_date, end_date = st.session_state['tonnage_date_range']
                        filtered_tonnage_df = filtered_tonnage_df[
                            (filtered_tonnage_df['DATE'].dt.date >= start_date) &
                            (filtered_tonnage_df['DATE'].dt.date <= end_date)
                        ]
                    if not filtered_tonnage_df.empty:
                        tonnage_melted = filtered_tonnage_df.melt(
                            id_vars=['DATE'],
                            value_vars=['DS Sud', 'DS Nord', 'KA'],
                            var_name='Site',
                            value_name='Tonnage'
                        )
                        fig_tonnage = px.line(
                            tonnage_melted,
                            x='DATE',
                            y='Tonnage',
                            color='Site',
                            title='Comparaison des tonnages par site au fil du temps',
                            height=400
                        )
                        fig_tonnage.update_layout(
                            xaxis_title="Date",
                            yaxis_title="Tonnage (T)",
                            template='plotly_white',
                            legend_title="Site"
                        )
                        figures['Comparaison des tonnages par site'] = fig_tonnage

                        total_tonnage_df = pd.DataFrame({
                            'Site': ['DS Sud', 'DS Nord', 'KA'],
                            'Tonnage Total': [
                                filtered_tonnage_df['DS Sud'].sum(),
                                filtered_tonnage_df['DS Nord'].sum(),
                                filtered_tonnage_df['KA'].sum()
                            ]
                        })
                        fig_total_tonnage = px.bar(
                            total_tonnage_df,
                            x='Site',
                            y='Tonnage Total',
                            title='Tonnage total par site',
                            height=400,
                            text='Tonnage Total'
                        )
                        fig_total_tonnage.update_traces(
                            texttemplate='%{text:,.0f} T',
                            textposition='auto'
                        )
                        fig_total_tonnage.update_layout(
                            xaxis_title="Site",
                            yaxis_title="Tonnage total (T)",
                            template='plotly_white'
                        )
                        figures['Tonnage total par site'] = fig_total_tonnage

                hm_df = load_hm_data(st.session_state.uploaded_hm_file)
                if not hm_df.empty:
                    filtered_hm_df = hm_df.copy()
                    if st.session_state['hm_date_range'] is not None and len(st.session_state['hm_date_range']) == 2:
                        start_date, end_date = st.session_state['hm_date_range']
                        filtered_hm_df = filtered_hm_df[
                            (filtered_hm_df['ENGINS'].dt.date >= start_date) &
                            (filtered_hm_df['ENGINS'].dt.date <= end_date)
                        ]
                    if not filtered_hm_df.empty:
                        equipment_columns = [col for col in filtered_hm_df.columns if col not in ['ENGINS', 'TOTAL_HOURS']]
                        hm_melted = filtered_hm_df.melt(
                            id_vars=['ENGINS'],
                            value_vars=equipment_columns,
                            var_name='Équipement',
                            value_name='Heures'
                        )
                        fig_hm = px.line(
                            hm_melted,
                            x='ENGINS',
                            y='Heures',
                            color='Équipement',
                            title='Comparaison des heures de marche par équipement au fil du temps',
                            height=400
                        )
                        fig_hm.update_layout(
                            xaxis_title="Date",
                            yaxis_title="Heures de marche (h)",
                            template='plotly_white',
                            legend_title="Équipement"
                        )
                        figures['Comparaison des heures de marche par équipement'] = fig_hm

                        total_hm_df = pd.DataFrame({
                            'Équipement': equipment_columns,
                            'Heures Totales': [filtered_hm_df[col].sum() for col in equipment_columns]
                        })
                        fig_total_hm = px.bar(
                            total_hm_df,
                            x='Équipement',
                            y='Heures Totales',
                            title='Heures totales par équipement',
                            height=400,
                            text='Heures Totales'
                        )
                        fig_total_hm.update_traces(
                            texttemplate='%{text:,.0f} h',
                            textposition='auto'
                        )
                        fig_total_hm.update_layout(
                            xaxis_title="Équipement",
                            yaxis_title="Heures totales (h)",
                            template='plotly_white'
                        )
                        figures['Heures totales par équipement'] = fig_total_hm
                
                pivot_engine = pd.DataFrame()
                selected_engines = st.session_state.get('selected_engines', [])
                if not filtered_data.empty and selected_engines and selected_engines != ["Tous les types"]:
                    pivot_engine = pd.pivot_table(
                        filtered_data[filtered_data['CATEGORIE'].isin(selected_engines)],
                        values='Montant',
                        index='Desc_CA',
                        columns='Desc_Cat',
                        aggfunc='sum',
                        fill_value=0,
                        margins=True,
                        margins_name='Total'
                    ).round(2)
                elif not filtered_data.empty:
                    pivot_data = pd.pivot_table(
                        filtered_data,
                        values='Montant',
                        index='Desc_CA',
                        columns='Desc_Cat',
                        aggfunc='sum',
                        fill_value=0,
                        margins=True,
                        margins_name='Total'
                    ).round(2)
                
                table_df = filtered_data[['Date', 'Desc_CA', 'Desc_Cat', 'Montant']].copy()
                table_df['Date'] = table_df['Date'].dt.strftime('%d/%m/%Y')
                table_df['Montant'] = table_df['Montant'].round(2)
                table_df = table_df.rename(columns={
                    'Date': 'Date',
                    'Desc_CA': 'Équipement',
                    'Desc_Cat': 'Type de consommation',
                    'Montant': 'Montant (DH)'
                })
                total_montant = table_df['Montant (DH)'].sum()
                
                report = generate_word_report(
                    filtered_data,
                    total_cost,
                    global_avg,
                    category_stats,
                    most_consumed_per_cat,
                    pivot_engine,
                    selected_engines,
                    table_df,
                    total_montant,
                    figures,
                    tonnage_df,
                    st.session_state['tonnage_date_range'],
                    hm_df,
                    st.session_state['hm_date_range']
                )
                
                st.download_button(
                    label="📥 Télécharger le rapport Word",
                    data=report,
                    file_name=f"Rapport_Consommation_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_button"
                )
                
                st.success("Rapport généré avec succès!")
    st.markdown("""
    <div class='header-container'>
        <h1 style='color: white; text-align:center; margin-top:0;'>Tableau De Bord De La Consommation Des Engins</h1>
        <p style='color: white; text-align: center; margin-bottom:0'>Suivre et optimiser la consommation des équipements</p>
    </div>
    """, unsafe_allow_html=True)
    

    kpi_container = st.container()
    with kpi_container:
        st.markdown(f"""
        <div class='analysis-card'>
            <h3 style='color: #2c3e50; margin-top:0;'>Indicateurs globaux</h3>
            <div style='display:flex; justify-content:space-between;'>
                <div class='metric-card'>
                    <p class='metric-title' style='font-size: 20px;'>Coût total</p>
                    <p class='metric-value'>{total_cost:,.0f} DH</p>
                </div>
                <div class='metric-card'>
                    <p class='metric-title' style='font-size: 20px;'>Moyenne globale des engins par jour</p>
                    <p class='metric-value'>{global_avg:,.0f} DH</p>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        categories = category_stats['CATEGORIE'].unique()
        cols = st.columns(len(categories))

        for idx, (col, (_, row)) in enumerate(zip(cols, category_stats.iterrows())):
            with col:
                most_consumed = most_consumed_per_cat[most_consumed_per_cat['CATEGORIE'] == row['CATEGORIE']]
                most_consumed_desc = most_consumed['Desc_Cat'].iloc[0] if not most_consumed.empty else "Aucune"
                
                st.markdown(f"""
                <div class='metric-card'>
                    <h4 style='color: #2c3e50; margin-top:0; text-align:center;'>{row['CATEGORIE']}</h4>
                    <div style='display:flex; justify-content:space-between; margin-bottom:5px;'>
                        <span class='metric-title'>Total:</span>
                        <span class='metric-value'>{row['Total']:,.0f} DH</span>
                    </div>
                    <div style='display:flex; justify-content:space-between; margin-bottom:5px;'>
                        <span class='metric-title'>Moyenne des engins par jour:</span>
                        <span class='metric-value'>{row['Moyenne']:,.0f} DH</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("<div class='analysis-card'><h3 style='color: #2c3e50;'>Consommation des catégories par type de consommation</h3></div>", unsafe_allow_html=True)
        hist_data = filtered_data.groupby(['CATEGORIE', 'Desc_Cat'])['Montant'].sum().reset_index()
        fig_hist = px.bar(
            hist_data,
            x='CATEGORIE',
            y='Montant',
            color='Desc_Cat',
            barmode='group',
            title='Consommation par catégorie et type de consommation',
            height=500,
            text='Desc_Cat'
        )
        fig_hist.update_traces(
            texttemplate='%{text}',
            textposition='inside',
            textfont=dict(
                size=30,
                color='#000000',
                family='Gravitas One, sans-serif'
            )
        )
        fig_hist.update_layout(
            xaxis_title="Catégorie",
            yaxis_title="Montant total (DH)",
            template='plotly_white',
            legend_title="Type de consommation",
            xaxis={'tickangle': 45},
            showlegend=False
        )
        st.plotly_chart(fig_hist, use_container_width=True, key="category_consumption")
        
        st.markdown("<div class='analysis-card'><h3 style='color: #2c3e50;'>Consommation totale par type d'engin et catégorie de consommation</h3></div>", unsafe_allow_html=True)
        pivot_table = pd.pivot_table(
            filtered_data,
            values='Montant',
            index='CATEGORIE',
            columns='Desc_Cat',
            aggfunc='sum',
            fill_value=0,
            margins=True,
            margins_name='Total'
        ).round(2)
        st.dataframe(
            pivot_table.style.format("{:,.2f} DH").set_properties(**{
                'background-color': 'white',
                'border': '1px solid #dfe6e9',
                'text-align': 'center',
                'color': '#2c3e50'
            }).set_table_styles([
                {'selector': 'th', 'props': [('background-color', 'white'), ('color', '#3498db'), ('font-weight', 'bold')]}
            ]),
            use_container_width=True
        )

        st.markdown("<div class='analysis-card'><h3 style='color: #2c3e50;'>Consommation par équipement pour les types d'engin sélectionnés</h3></div>", unsafe_allow_html=True)
        engine_data = filtered_data.copy()
        if not engine_data.empty:
            st.markdown("<h4 style='color: #2c3e50;'>Filtrer par type d'engin</h4>", unsafe_allow_html=True)
            engine_types = ["Tous les types", "DUMPER", "FORATION", "10 TONNES"]
            selected_engines = st.multiselect(
                "Sélectionner les types d'engin",
                engine_types,
                default=["Tous les types"],
                key="engine_type_multiselect"
            )
            st.session_state['selected_engines'] = selected_engines

            selected_engines = [str(engine) for engine in selected_engines]
            
            if "Tous les types" not in selected_engines and selected_engines:
                try:
                    engine_data = engine_data[engine_data['CATEGORIE'].isin(selected_engines)]
                    st.write(f"Total Montant après filtre par selected_engines : {engine_data['Montant'].sum():,.2f} DH")
                except TypeError as e:
                    st.error(f"Erreur lors du filtrage par catégorie : {str(e)}")
                    st.write(f"Valeurs problématiques dans CATEGORIE : {engine_data['CATEGORIE'].unique()}")
                    st.stop()

            if engine_data.empty:
                st.warning("Aucune donnée disponible pour les types d'engin sélectionnés.")
            else:
                pivot_engine = pd.pivot_table(
                    engine_data,
                    values='Montant',
                    index='Desc_CA',
                    columns='Desc_Cat',
                    aggfunc='sum',
                    fill_value=0,
                    margins=True,
                    margins_name='Total'
                ).round(2)
                st.dataframe(
                    pivot_engine.style.format("{:,.2f} DH").set_properties(**{
                        'background-color': 'white',
                        'border': '1px solid #dfe6e9',
                        'text-align': 'center',
                        'color': '#2c3e50'
                    }).set_table_styles([
                        {'selector': 'th', 'props': [('background-color', 'white'), ('color', '#3498db'), ('font-weight', 'bold')]}
                    ]),
                    use_container_width=True
                )
        else:
            st.warning("Aucune donnée disponible pour les critères sélectionnés.")

    tabs = st.tabs(
        [f"📋 {cat}" for cat in sorted(filtered_data['CATEGORIE'].unique())] + 
        ["📊 Analyse comparative", "💡 Recommandations", "📋 Tableau des équipements", "📈 Tonnage des Sites", "⏰ Heures de Marche"]
    )

    for i, cat in enumerate(sorted(filtered_data['CATEGORIE'].unique())):
        with tabs[i]:
            cat_data = filtered_data[filtered_data['CATEGORIE'] == cat]
            st.markdown(f"""
            <div class='analysis-card'>
                <h2 style='color: #2c3e50; margin-top:0;'>Analyse pour la catégorie {cat}</h2>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("<h3 style='color: #2c3e50;'>Consommation par équipement</h3>", unsafe_allow_html=True)
            equip_sum = cat_data.groupby('Desc_CA')['Montant'].sum().reset_index().sort_values('Montant', ascending=False)
            fig2 = px.bar(
                equip_sum,
                x='Desc_CA',
                y='Montant',
                title=f'Consommation totale par équipement ({cat})',
                height=400,
                text='Montant'
            )
            fig2.update_traces(
                texttemplate='%{text:,.0f} DH',
                textposition='auto'
            )
            fig2.update_layout(
                xaxis_title="Équipement",
                yaxis_title="Montant total (DH)",
                template='plotly_white',
                xaxis={'categoryorder':'total descending'}
            )
            st.plotly_chart(fig2, use_container_width=True, key=f"equip_sum_{cat}")

    with tabs[-5]:
        st.markdown("""
        <div class='analysis-card'>
            <h2 style='color: #2c3e50; margin-top:0;'>Analyse comparative</h2>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<h3 style='color: #2c3e50;'>Comparaison des catégories</h3>", unsafe_allow_html=True)
        fig_comp = px.bar(
            filtered_data.groupby('CATEGORIE')['Montant'].sum().reset_index(),
            x='CATEGORIE',
            y='Montant',
            title='Coût total par catégorie',
            height=400,
            text='Montant'
        )
        fig_comp.update_traces(
            texttemplate='%{text:,.0f} DH',
            textposition='auto'
        )
        fig_comp.update_layout(
            xaxis_title="Catégorie",
            yaxis_title="Montant total (DH)",
            template='plotly_white'
        )
        st.plotly_chart(fig_comp, use_container_width=True, key="category_comparison")
    with tabs[-4]:
        st.markdown("""
        <div class='analysis-card'>
            <h2 style='color: #2c3e50; margin-top:0;'>Recommandations</h2>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<h3 style='color: #2c3e50;'>Catégories prioritaires</h3>", unsafe_allow_html=True)
        top_categories = filtered_data.groupby('CATEGORIE')['Montant'].sum().nlargest(3).reset_index()
        cols = st.columns(3)
        for i, (col, (_, row)) in enumerate(zip(cols, top_categories.iterrows())):
            with col:
                st.markdown(f"""
                <div class='metric-card'>
                    <h4 style='color: #2c3e50; text-align:center;'>{row['CATEGORIE']}</h4>
                    <p style='color: #2c3e50; text-align:center; font-size:24px; font-weight:bold;'>{row['Montant']:,.0f} DH</p>
                    <p style='color: #7f8c8d; text-align:center;'>{(row['Montant']/total_cost)*100:.1f}% du total</p>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("""
        <div class='analysis-card'>
            <h3 style='color: #2c3e50;'>Actions recommandées</h3>
            <ul style='color: #2c3e50;'>
                <li>Prioriser les analyses des équipements dans les catégories les plus coûteuses</li>
                <li>Mettre en place un suivi mensuel des consommations par catégorie</li>
                <li>Comparer les performances des équipements similaires pour identifier les anomalies</li>
                <li>Négocier avec les fournisseurs pour les pièces les plus fréquemment remplacées</li>
                <li>Étudier la possibilité de maintenance préventive pour réduire les coûts</li>
                <li>Former les opérateurs à une utilisation optimale des équipements</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

    with tabs[-3]:
        st.markdown("""
        <div class='analysis-card'>
            <h2 style='color: #2c3e50; margin-top:0;'>Tableau de la consommation des équipements</h2>
            <p style='color: #7f8c8d;'>Consommation détaillée par équipement pour la catégorie sélectionnée</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<h3 style='color: #2c3e50;'>Filtrer par type de consommation</h3>", unsafe_allow_html=True)
        consumption_types = ["Tous les types"] + sorted(filtered_data['Desc_Cat'].unique())
        selected_consumption_types = st.multiselect(
            "Sélectionner les types de consommation",
            consumption_types,
            default=["Tous les types"],
            key="consumption_type_multiselect"
        )
        
        table_df = filtered_data[['Date', 'Desc_CA', 'Desc_Cat', 'Montant']].copy()
        
        if "Tous les types" not in selected_consumption_types and selected_consumption_types:
            table_df = table_df[table_df['Desc_Cat'].isin(selected_consumption_types)]
            # Debug: Total après filtre de type de consommation
            st.write(f"Total Montant après filtre de type de consommation : {table_df['Montant'].sum():,.2f} DH")
        
        if table_df.empty:
            st.warning("Aucune donnée disponible pour les critères sélectionnés.")
        else:
            table_df['Date'] = table_df['Date'].dt.strftime('%d/%m/%Y')
            table_df['Montant'] = table_df['Montant'].round(2)
            table_df = table_df.rename(columns={
                'Date': 'Date',
                'Desc_CA': 'Équipement',
                'Desc_Cat': 'Type de consommation',
                'Montant': 'Montant (DH)'
            })
            
            total_montant = table_df['Montant (DH)'].sum()
            # Debug: Comparer les totaux
            st.write(f"Total Montant dans le Tableau des équipements : {total_montant:,.2f} DH")
            st.write(f"Total Montant dans les données brutes (avant filtres) : {df['Montant'].sum():,.2f} DH")
            
            # Ajout du bouton pour télécharger table_df en CSV
            csv = table_df.to_csv(index=False)
            st.download_button(
                label="📥 Télécharger les données du tableau (CSV)",
                data=csv,
                file_name=f"tableau_equipements_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
                key="download_table_csv"
            )
            
            st.dataframe(
                table_df.style.format({
                    'Montant (DH)': '{:,.2f} DH',
                    'Date': lambda x: x if x else ''
                }).set_properties(**{
                    'background-color': 'white',
                    'border': '1px solid #dfe6e9',
                    'text-align': 'center',
                    'color': '#2c3e50'
                }).set_table_styles([
                    {'selector': 'th', 'props': [('background-color', 'white'), ('color', '#3498db'), ('font-weight', 'bold')]}
                ]),
                height=600,
                use_container_width=True
            )
            
            st.markdown(f"""
            <div style='background-color: white; padding:10px; border-radius:10px; margin-top:10px; border: 1px solid #dfe6e9;'>
                <p style='color: #2c3e50; font-size:16px; font-weight:bold; text-align:right;'>Total : {total_montant:,.2f} DH</p>
            </div>
            """, unsafe_allow_html=True)

    with tabs[-2]:
        st.markdown("""
        <div class='analysis-card'>
            <h2 style='color: #2c3e50; margin-top:0;'>Tonnage des Sites</h2>
            <p style='color: #7f8c8d;'>Comparaison des tonnages pour DS Sud, DS Nord et KA</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<h3 style='color: #2c3e50;'>Importer des fichiers de tonnage</h3>", unsafe_allow_html=True)
        st.markdown("**Note** : Plusieurs fichiers Excel (.xlsx) ou ZIP (.zip) peuvent être importés (max 200 Mo par fichier).")
        st.markdown("**Fichiers importés** :")
        if st.session_state.uploaded_tonnage_file:
            try:
                st.write(", ".join([f.name for f in st.session_state.uploaded_tonnage_file]))
            except AttributeError:
                st.error("Erreur : Les fichiers de tonnage stockés sont invalides. Veuillez réimporter les fichiers.")
                st.session_state.uploaded_tonnage_file = None
        else:
            st.write("Aucun fichier de tonnage importé.")

        tonnage_df = pd.DataFrame()

        with st.form("tonnage_file_upload_form", clear_on_submit=True):
            uploaded_tonnage_files = st.file_uploader(
                "Téléverser des fichiers Excel ou ZIP pour les tonnages (max 200 Mo par fichier)",
                type=["xlsx", "zip"],
                accept_multiple_files=True,
                key=f"tonnage_file_uploader_{st.session_state.tonnage_file_uploader_key}"
            )
            submit_tonnage_button = st.form_submit_button("Charger les fichiers de tonnage")

            if submit_tonnage_button:
                if uploaded_tonnage_files:
                    st.session_state.uploaded_tonnage_file = uploaded_tonnage_files
                    st.session_state.tonnage_file_uploader_key += 1
                    tonnage_df = load_tonnage_data(st.session_state.uploaded_tonnage_file)
                    if not tonnage_df.empty:
                        st.success(f"Fichiers de tonnage chargés avec succès. Nombre total de lignes : {tonnage_df.shape[0]}")
                    else:
                        st.warning("Aucun fichier de tonnage valide n'a pu être chargé. Veuillez vérifier les fichiers téléversés.")
                        st.session_state.uploaded_tonnage_file = None
                else:
                    st.warning("Aucun fichier de tonnage sélectionné. Veuillez téléverser un ou plusieurs fichiers Excel ou ZIP.")
            else:
                if st.session_state.uploaded_tonnage_file:
                    tonnage_df = load_tonnage_data(st.session_state.uploaded_tonnage_file)

        if tonnage_df.empty:
            st.warning("Aucune donnée de tonnage disponible. Veuillez téléverser un fichier Excel ou ZIP valide.")
        else:
            st.markdown("<h3 style='color: #2c3e50;'>Filtrer par plage de dates</h3>", unsafe_allow_html=True)
            default_tonnage_start = tonnage_df['DATE'].min().date() if not tonnage_df.empty else datetime.today().date()
            default_tonnage_end = tonnage_df['DATE'].max().date() if not tonnage_df.empty else datetime.today().date()
            tonnage_date_range = st.date_input(
                "Période pour les tonnages",
                value=(default_tonnage_start, default_tonnage_end),
                min_value=default_tonnage_start,
                max_value=default_tonnage_end,
                help="Choisir une plage de dates pour filtrer les données de tonnage",
                key="tonnage_date_range"
            )

            filtered_tonnage_df = tonnage_df.copy()
            if len(tonnage_date_range) == 2:
                start_date, end_date = tonnage_date_range
                filtered_tonnage_df = filtered_tonnage_df[
                    (filtered_tonnage_df['DATE'].dt.date >= start_date) & 
                    (filtered_tonnage_df['DATE'].dt.date <= end_date)
                ]

            if filtered_tonnage_df.empty:
                st.warning("Aucune donnée de tonnage disponible après filtrage. Veuillez ajuster les filtres.")
            else:
                st.markdown("<h3 style='color: #2c3e50;'>Tableau des tonnages</h3>", unsafe_allow_html=True)
                display_tonnage_df = filtered_tonnage_df[['DATE', 'DS Sud', 'DS Nord', 'KA', 'CUMMULE']].copy()
                display_tonnage_df['DATE'] = display_tonnage_df['DATE'].dt.strftime('%d/%m/%Y')
                display_tonnage_df = display_tonnage_df.rename(columns={
                    'DATE': 'Date',
                    'DS Sud': 'DS Sud (T)',
                    'DS Nord': 'DS Nord (T)',
                    'KA': 'KA (T)',
                    'CUMMULE': 'Cumulé (T)'
                })
                
                total_tonnage = display_tonnage_df[['DS Sud (T)', 'DS Nord (T)', 'KA (T)']].sum().to_dict()
                total_cumule = display_tonnage_df['Cumulé (T)'].sum()
                
                st.dataframe(
                    display_tonnage_df.style.format({
                        'DS Sud (T)': '{:,.2f} T',
                        'DS Nord (T)': '{:,.2f} T',
                        'KA (T)': '{:,.2f} T',
                        'Cumulé (T)': '{:,.2f} T',
                        'Date': lambda x: x if x else ''
                    }).set_properties(**{
                        'background-color': 'white',
                        'border': '1px solid #dfe6e9',
                        'text-align': 'center',
                        'color': '#2c3e50'
                    }).set_table_styles([
                        {'selector': 'th', 'props': [('background-color', 'white'), ('color', '#3498db'), ('font-weight', 'bold')]}
                    ]),
                    height=600,
                    use_container_width=True
                )
                
                st.markdown(f"""
                <div style='background-color: white; padding:10px; border-radius:10px; margin-top:10px; border: 1px solid #dfe6e9;'>
                    <p style='color: #2c3e50; font-size:16px; font-weight:bold; text-align:right;'>
                        Total DS Sud: {total_tonnage['DS Sud (T)']:,.2f} T | 
                        Total DS Nord: {total_tonnage['DS Nord (T)']:,.2f} T | 
                        Total KA: {total_tonnage['KA (T)']:,.2f} T | 
                        Total Cumulé: {total_cumule:,.2f} T
                    </p>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("<h3 style='color: #2c3e50;'>Comparaison des tonnages par site</h3>", unsafe_allow_html=True)
                tonnage_melted = filtered_tonnage_df.melt(
                    id_vars=['DATE'],
                    value_vars=['DS Sud', 'DS Nord', 'KA'],
                    var_name='Site',
                    value_name='Tonnage'
                )
                fig_tonnage = px.line(
                    tonnage_melted,
                    x='DATE',
                    y='Tonnage',
                    color='Site',
                    title='Comparaison des tonnages par site au fil du temps',
                    height=500,
                    markers=True
                )
                fig_tonnage.update_layout(
                    xaxis_title="Date",
                    yaxis_title="Tonnage (T)",
                    template='plotly_white',
                    legend_title="Site",
                    xaxis={'tickangle': 45}
                )
                st.plotly_chart(fig_tonnage, use_container_width=True, key="tonnage_comparison")

                total_tonnage_df = pd.DataFrame({
                    'Site': ['DS Sud', 'DS Nord', 'KA'],
                    'Tonnage Total': [
                        filtered_tonnage_df['DS Sud'].sum(),
                        filtered_tonnage_df['DS Nord'].sum(),
                        filtered_tonnage_df['KA'].sum()
                    ]
                })
                fig_total_tonnage = px.bar(
                    total_tonnage_df,
                    x='Site',
                    y='Tonnage Total',
                    title='Tonnage total par site',
                    height=400,
                    text='Tonnage Total'
                )
                fig_total_tonnage.update_traces(
                    texttemplate='%{text:,.0f} T',
                    textposition='auto'
                )
                fig_total_tonnage.update_layout(
                    xaxis_title="Site",
                    yaxis_title="Tonnage total (T)",
                    template='plotly_white'
                )
                st.plotly_chart(fig_total_tonnage, use_container_width=True, key="total_tonnage_comparison")
    with tabs[-1]:
        st.markdown("""
        <div class='analysis-card'>
            <h2 style='color: #2c3e50; margin-top:0;'>Heures de Marche des Équipements</h2>
            <p style='color: #7f8c8d;'>Analyse des heures de marche, tonnages et consommations pour la période sélectionnée</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<h3 style='color: #2c3e50;'>Importer des fichiers d'heures de marche</h3>", unsafe_allow_html=True)
        st.markdown("**Note** : Plusieurs fichiers Excel (.xlsx) ou ZIP (.zip) peuvent être importés (max 200 Mo par fichier).")
        st.markdown("**Fichiers importés** :")
        if st.session_state.uploaded_hm_file:
            try:
                st.write(", ".join([f.name for f in st.session_state.uploaded_hm_file]))
            except AttributeError:
                st.error("Erreur : Les fichiers d'heures de marche stockés sont invalides. Veuillez réimporter les fichiers.")
                st.session_state.uploaded_hm_file = None
        else:
            st.write("Aucun fichier d'heures de marche importé.")

        hm_df = pd.DataFrame()

        with st.form("hm_file_upload_form", clear_on_submit=True):
            uploaded_hm_files = st.file_uploader(
                "Téléverser des fichiers Excel ou ZIP pour les heures de marche (max 200 Mo par fichier)",
                type=["xlsx", "zip"],
                accept_multiple_files=True,
                key=f"hm_file_uploader_{st.session_state.hm_file_uploader_key}"
            )
            submit_hm_button = st.form_submit_button("Charger les fichiers d'heures de marche")

            if submit_hm_button:
                if uploaded_hm_files:
                    st.session_state.uploaded_hm_file = uploaded_hm_files
                    st.session_state.hm_file_uploader_key += 1
                    hm_df = load_hm_data(st.session_state.uploaded_hm_file)
                    if not hm_df.empty:
                        st.success(f"Fichiers d'heures de marche chargés avec succès. Nombre total de lignes : {hm_df.shape[0]}")
                    else:
                        st.warning("Aucun fichier d'heures de marche valide n'a pu être chargé. Veuillez vérifier les fichiers téléversés.")
                        st.session_state.uploaded_hm_file = None
                else:
                    st.warning("Aucun fichier d'heures de marche sélectionné. Veuillez téléverser un ou plusieurs fichiers Excel ou ZIP.")
            else:
                if st.session_state.uploaded_hm_file:
                    hm_df = load_hm_data(st.session_state.uploaded_hm_file)

        if hm_df.empty:
            st.warning("Aucune donnée d'heures de marche disponible.")
        else:
            st.markdown("<h3 style='color: #2c3e50;'>Filtrer par plage de dates</h3>", unsafe_allow_html=True)
            default_hm_start = hm_df['ENGINS'].min().date() if not hm_df.empty else datetime.today().date()
            default_hm_end = hm_df['ENGINS'].max().date() if not hm_df.empty else datetime.today().date()
            hm_date_range = st.date_input(
                "Période pour les heures de marche",
                value=(default_hm_start, default_hm_end),
                min_value=default_hm_start,
                max_value=default_hm_end,
                help="Choisir une plage de dates pour filtrer les données d'heures de marche",
                key="hm_date_range"
            )

            filtered_hm_df = hm_df.copy()
            if len(hm_date_range) == 2:
                start_date, end_date = hm_date_range
                filtered_hm_df = filtered_hm_df[
                    (filtered_hm_df['ENGINS'].dt.date >= start_date) &
                    (filtered_hm_df['ENGINS'].dt.date <= end_date)
                ]

            if filtered_hm_df.empty:
                st.warning("Aucune donnée d'heures de marche disponible après filtrage.")
            else:
                st.markdown("<h3 style='color: #2c3e50;'>Tableau des heures de marche</h3>", unsafe_allow_html=True)
                equipment_columns = [col for col in filtered_hm_df.columns if col not in ['ENGINS', 'TOTAL_HOURS']]
                display_hm_df = filtered_hm_df[['ENGINS'] + equipment_columns + ['TOTAL_HOURS']].copy()
                display_hm_df['ENGINS'] = display_hm_df['ENGINS'].dt.strftime('%d/%m/%Y')
                display_hm_df = display_hm_df.rename(columns={
                    'ENGINS': 'Date',
                    'TOTAL_HOURS': 'Total (h)'
                })
                for col in equipment_columns:
                    display_hm_df = display_hm_df.rename(columns={col: f"{col} (h)"})

                # Ensure hours are integers
                for col in equipment_columns:
                    display_hm_df[f"{col} (h)"] = display_hm_df[f"{col} (h)"].fillna(0).astype(int)
                display_hm_df['Total (h)'] = display_hm_df['Total (h)'].fillna(0).astype(int)

                # Calculate totals for each equipment and overall total
                total_hours = display_hm_df[[f"{col} (h)" for col in equipment_columns]].sum().to_dict()
                total_sum = display_hm_df['Total (h)'].sum()

                # Create a totals row
                totals_row = {'Date': 'Total'}
                for col in equipment_columns:
                    totals_row[f"{col} (h)"] = int(total_hours[f"{col} (h)"])
                totals_row['Total (h)'] = int(total_sum)
                
                # Append totals row to the DataFrame
                totals_df = pd.DataFrame([totals_row])
                display_hm_df = pd.concat([display_hm_df, totals_df], ignore_index=True)

                st.dataframe(
                    display_hm_df.style.format({
                        **{f"{col} (h)": '{:d} h' for col in equipment_columns},  # Integer format
                        'Total (h)': '{:d} h',  # Integer format
                        'Date': lambda x: x if x else ''
                    }).set_properties(**{
                        'background-color': 'white',
                        'border': '1px solid #dfe6e9',
                        'text-align': 'center',
                        'color': '#2c3e50'
                    }).set_table_styles([
                        {'selector': 'th', 'props': [('background-color', 'white'), ('color', '#3498db'), ('font-weight', 'bold')]},
                        {'selector': 'tr:last-child', 'props': [('background-color', '#f0f2f6'), ('font-weight', 'bold')]}  # Highlight totals row
                    ]),
                    height=600,
                    use_container_width=True
                )

                st.markdown(f"""
                <div style='background-color: white; padding:10px; border-radius:10px; margin-top:10px; border: 1px solid #dfe6e9;'>
                    <p style='color: #2c3e50; font-size:16px; font-weight:bold; text-align:right;'>
                        Total Cumulé: {total_sum:d} h
                    </p>
                </div>
                """, unsafe_allow_html=True)

                # Calcul des totaux pour la période sélectionnée
                st.markdown("<h3 style='color: #2c3e50;'>Totaux pour la période sélectionnée</h3>", unsafe_allow_html=True)
                
                # Filtrer les données de consommation
                filtered_data_df = filtered_data.copy()
                if len(date_range) == 2:
                    start_date, end_date = date_range
                    filtered_data_df = filtered_data_df[
                        (filtered_data_df['Date'].dt.date >= start_date) &
                        (filtered_data_df['Date'].dt.date <= end_date)
                    ]
                    consumption_period = f"du {start_date.strftime('%d/%m/%Y')} au {end_date.strftime('%d/%m/%Y')}"
                else:
                    consumption_period = "Toutes les dates disponibles"
                
                # Filtrer les données de tonnage
                filtered_tonnage_df = tonnage_df.copy()
                if len(tonnage_date_range) == 2:
                    start_date, end_date = tonnage_date_range
                    filtered_tonnage_df = filtered_tonnage_df[
                        (filtered_tonnage_df['DATE'].dt.date >= start_date) &
                        (filtered_tonnage_df['DATE'].dt.date <= end_date)
                    ]
                    tonnage_period = f"du {start_date.strftime('%d/%m/%Y')} au {end_date.strftime('%d/%m/%Y')}"
                else:
                    tonnage_period = "Toutes les dates disponibles"
                
                # Filtrer les données d'heures de marche
                filtered_hm_df = hm_df.copy()
                if len(hm_date_range) == 2:
                    start_date, end_date = hm_date_range
                    filtered_hm_df = filtered_hm_df[
                        (filtered_hm_df['ENGINS'].dt.date >= start_date) &
                        (filtered_hm_df['ENGINS'].dt.date <= end_date)
                    ]
                    hm_period = f"du {start_date.strftime('%d/%m/%Y')} au {end_date.strftime('%d/%m/%Y')}"
                else:
                    hm_period = "Toutes les dates disponibles"
                
                # Calcul des totaux
                total_consumption = filtered_data_df['Montant'].sum() if not filtered_data_df.empty else 0
                total_tonnage = filtered_tonnage_df['CUMMULE'].sum() if not filtered_tonnage_df.empty else 0
                total_hours = filtered_hm_df['TOTAL_HOURS'].sum() if not filtered_hm_df.empty else 0
                # Affichage des totaux
                st.markdown(f"""
                <div class='analysis-card'>
                    <h4 style='color: #2c3e50;'>Résumé des totaux</h4>
                    <div style='display:flex; justify-content:space-between;'>
                        <div class='metric-card'>
                            <p class='metric-title'>Consommation Totale</p>
                            <p class='metric-value'>{total_consumption:,.2f} DH</p>
                        </div>
                        <div class='metric-card'>
                            <p class='metric-title'>Tonnage Total</p>
                            <p class='metric-value'>{total_tonnage:,.2f} T</p>
                        </div>
                        <div class='metric-card'>
                            <p class='metric-title'>Heures Totales</p>
                            <p class='metric-value'>{total_hours:d} h</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # Analyse de rentabilité (Gagnant ou Perte)
                st.markdown("<h3 style='color: #2c3e50;'>Analyse de Rentabilité</h3>", unsafe_allow_html=True)
                
                if total_tonnage == 0 or total_hours == 0 or total_consumption == 0:
                    st.warning("Données insuffisantes pour l'analyse de rentabilité. Veuillez vérifier que les fichiers de consommation, de tonnage et d'heures de marche sont chargés.")
                else:
                    # Calcul du rendement moyen (tonnes par heure)
                    average_yield = total_tonnage / total_hours if total_hours > 0 else 0
                    
                    # Calcul du coût par tonne
                    cost_per_tonne = total_consumption / total_tonnage if total_tonnage > 0 else float('inf')

                    # Seuils pour déterminer si l'opération est gagnante ou perdante
                    # À ajuster selon votre contexte
                    YIELD_THRESHOLD = 10  # Tonnes par heure minimum
                    COST_PER_TONNE_THRESHOLD = 500  # Coût maximum par tonne en DH

                    st.markdown(f"""
                    <div class='analysis-card'>
                        <h4 style='color: #2c3e50;'>Indicateurs de performance</h4>
                        <p><strong>Rendement moyen :</strong> {average_yield:.2f} T/h</p>
                        <p><strong>Coût par tonne :</strong> {cost_per_tonne:.2f} DH/T</p>
                    </div>
                    """, unsafe_allow_html=True)

                    # Déterminer si l'opération est gagnante ou perdante
                    if average_yield >= YIELD_THRESHOLD and cost_per_tonne <= COST_PER_TONNE_THRESHOLD:
                        st.markdown("""
                        <div style='background-color: #e8f5e9; padding:15px; border-radius:10px; border: 1px solid #4caf50;'>
                            <h4 style='color: #2c3e50;'>Résultat : Opération Gagnante ✅</h4>
                            <p style='color: #2c3e50;'>Le rendement moyen ({:.2f} T/h) est supérieur au seuil de {:.2f} T/h, et le coût par tonne ({:.2f} DH/T) est inférieur au seuil de {:.2f} DH/T. L'opération est efficace et rentable.</p>
                        </div>
                        """.format(average_yield, YIELD_THRESHOLD, cost_per_tonne, COST_PER_TONNE_THRESHOLD), unsafe_allow_html=True)
                    else:
                        reasons = []
                        if average_yield < YIELD_THRESHOLD:
                            reasons.append(f"Le rendement moyen ({average_yield:.2f} T/h) est inférieur au seuil de {YIELD_THRESHOLD:.2f} T/h.")
                        if cost_per_tonne > COST_PER_TONNE_THRESHOLD:
                            reasons.append(f"Le coût par tonne ({cost_per_tonne:.2f} DH/T) dépasse le seuil de {COST_PER_TONNE_THRESHOLD:.2f} DH/T.")
                        
                        st.markdown("""
                        <div style='background-color: #ffebee; padding:15px; border-radius:10px; border: 1px solid #e57373;'>
                            <h4 style='color: #2c3e50;'>Résultat : Opération Perdante ❌</h4>
                            <p style='color: #2c3e50;'>L'opération présente des inefficacités :</p>
                            <ul style='color: #2c3e50;'>
                                {}
                            </ul>
                            <p style='color: #2c3e50; font-weight:bold;'>Recommandations :</p>
                            <ul style='color: #2c3e50;'>
                                <li>Optimiser l'utilisation des équipements pour augmenter le rendement horaire.</li>
                                <li>Réduire les coûts d'exploitation en négociant les prix des pièces ou en améliorant la maintenance préventive.</li>
                                <li>Analyser les équipements spécifiques pour identifier les sources de surconsommation.</li>
                            </ul>
                        </div>
                        """.format("".join(f"<li>{reason}</li>" for reason in reasons)), unsafe_allow_html=True)
