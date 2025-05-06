import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
import plotly.io as pio

# Configurer l'option Pandas pour augmenter le nombre maximum d'éléments pour Styler
pd.set_option("styler.render.max_elements", 600000)

# Fonction pour formater les grands nombres
def format_number(num):
    if pd.isna(num) or num == 0:
        return "0"
    if num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"
    if num >= 1_000:
        return f"{num / 1_000:.1f}K"
    return f"{num:.2f}"

# Fonction pour convertir le DataFrame en CSV
def convert_df_to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

# Fonction pour sauvegarder un graphique Plotly en tant qu'image et retourner le chemin
def save_plotly_fig_as_image(fig, filename):
    pio.write_image(fig, filename, format="png")
    return filename

# Fonction pour nettoyer les chaînes de caractères pour éviter les problèmes dans le document Word
def sanitize_text(text):
    if pd.isna(text):
        return "N/A"
    return str(text).replace("\n", " ").replace("\r", " ").strip()

# Fonction pour créer un document Word
def create_word_doc(df, df_aggregated, filtered_df, filtered_df_aggregated, total_items, total_quantity, total_cost, unique_groups, most_expensive, turnover_combined, cost_trend, top_expensive, abc_summary, top_quantity, top_cost, cost_by_group, cost_by_time, filtered_df_box, paged_df, alerts_df=None):
    doc = Document()
    doc.add_heading("Rapport d'Analyse des Stocks", 0)

    # Métriques Globales
    doc.add_heading("Métriques Globales", level=1)
    doc.add_paragraph(f"Nombre Total d'Articles: {format_number(total_items)}")
    doc.add_paragraph(f"Quantité Totale: {format_number(total_quantity)}")
    doc.add_paragraph(f"Coût Total: {format_number(total_cost)}")
    doc.add_paragraph(f"Groupes Uniques: {unique_groups}")

    # Aperçu Intéressant
    doc.add_heading("Aperçu Intéressant", level=1)
    if not filtered_df_aggregated.empty and not pd.isna(most_expensive["unit_price"]):
        doc.add_paragraph(
            f"L'article le plus cher par unité est {sanitize_text(most_expensive['DES_ARTICLE'])} "
            f"à {format_number(most_expensive['unit_price'])} par unité."
        )
    else:
        doc.add_paragraph("Aucun article trouvé pour l'aperçu intéressant.")

    # Analyses et Visualisations Avancées
    doc.add_heading("Analyses et Visualisations Avancées", level=1)

    # 1. Taux de Rotation des Stocks
    doc.add_heading("Taux de Rotation des Stocks (Articles à Rotation Rapide vs Lente)", level=2)
    doc.add_paragraph(
        "Ce graphique montre les 5 articles à rotation rapide et les 5 articles à rotation lente en fonction de leur taux de rotation (quantité par période). "
        "Il aide à identifier les articles qui sont utilisés ou vendus rapidement par rapport à ceux qui stagnent, facilitant l'optimisation des stocks."
    )
    fig_turnover = px.bar(
        turnover_combined,
        x="Taux de Rotation",
        y="DES_ARTICLE",
        orientation="h",
        title="Top 5 Articles à Rotation Rapide et Lente (Taux de Rotation)",
        color="Taux de Rotation",
        color_continuous_scale="Viridis",
        text_auto=True
    )
    fig_turnover.update_layout(xaxis_title="Taux de Rotation (Quantité par Période)", yaxis_title="Article")
    turnover_img = "turnover.png"
    save_plotly_fig_as_image(fig_turnover, turnover_img)
    doc.add_picture(turnover_img, width=Inches(6))

    # 2. Tendance des Coûts au Fil du Temps
    doc.add_heading("Tendance des Coûts au Fil du Temps", level=2)
    doc.add_paragraph(
        "Ce graphique en ligne montre le coût total au fil du temps (par mois). Il permet de suivre les fluctuations des coûts, d'identifier les tendances saisonnières, "
        "et de repérer les périodes de dépenses inhabituellement élevées ou faibles, utile pour la planification budgétaire et financière."
    )
    fig_cost_trend = px.line(
        cost_trend,
        x="Mois",
        y="MONTANT",
        title="Tendance du Coût Total au Fil du Temps",
        markers=True,
        color_discrete_sequence=["#FF4D4F"]
    )
    fig_cost_trend.update_layout(xaxis_title="Mois", yaxis_title="Coût Total", xaxis_tickangle=45)
    fig_cost_trend.update_traces(
        hovertemplate="<b>Mois:</b> %{x}<br><b>Coût:</b> %{y:,.2f}",
        line=dict(width=3)
    )
    cost_trend_img = "cost_trend.png"
    save_plotly_fig_as_image(fig_cost_trend, cost_trend_img)
    doc.add_picture(cost_trend_img, width=Inches(6))

        # 3. Top 10 Articles les Plus Chers par Prix Unitaire
        # 3. Top 10 Articles les Plus Chers par Prix Unitaire
    doc.add_heading("Top 10 Articles les Plus Chers par Prix Unitaire", level=2)
    doc.add_paragraph(
        "Ce tableau liste les 10 articles ayant les prix unitaires les plus élevés, agrégés pour éviter les doublons. Il montre la quantité totale et le coût "
        "pour chaque article, aidant à identifier les articles de grande valeur qui peuvent nécessiter une attention particulière dans les stratégies de tarification ou d'approvisionnement."
    )
    if not top_expensive.empty:
        # Add table with header
        table = doc.add_table(rows=1, cols=len(top_expensive.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(top_expensive.columns):
            hdr_cells[i].text = str(col)
        
        # Add data rows
        for _, row in top_expensive.iterrows():
            row_cells = table.add_row().cells
            for col_idx, value in enumerate(row):
                try:
                    row_cells[col_idx].text = sanitize_text(value)
                except Exception as e:
                    row_cells[col_idx].text = "Erreur"
    else:
        doc.add_paragraph("Aucun article trouvé pour ce tableau.")

    # 4. Analyse ABC
    doc.add_heading("Analyse ABC des Articles", level=2)
    doc.add_paragraph(
        "Ce graphique montre la répartition des articles selon l'analyse ABC : les articles de catégorie A (20% des articles, 80% du coût), B (30% des articles, 15% du coût), "
        "et C (50% des articles, 5% du coût). Cela aide à prioriser la gestion des stocks en se concentrant sur les articles les plus coûteux."
    )
    fig_abc = px.pie(
        abc_summary,
        names="Catégorie",
        values="MONTANT",
        title="Répartition des Coûts par Catégorie ABC",
        color_discrete_sequence=["#FF4D4F", "#FFA500", "#3B82F6"]
    )
    fig_abc.update_traces(
        textinfo="percent+label",
        hovertemplate="<b>Catégorie:</b> %{label}<br><b>Coût:</b> %{value:,.2f}<br><b>Pourcentage:</b> %{percent}"
    )
    abc_img = "abc.png"
    save_plotly_fig_as_image(fig_abc, abc_img)
    doc.add_picture(abc_img, width=Inches(6))

    # Répartition de l'Inventaire
    doc.add_heading("Répartition de l'Inventaire", level=1)

    # Top 5 Articles par Quantité
    doc.add_heading("Top 5 Articles par Quantité", level=2)
    doc.add_paragraph(
        "Ce graphique en barres met en évidence les 5 articles ayant les quantités totales les plus élevées en stock. Il aide à identifier les articles les plus stockés, "
        "ce qui peut indiquer une forte demande ou un éventuel surstockage."
    )
    fig_quantity = px.bar(
        top_quantity,
        x="DES_ARTICLE",
        y="QUANTITE",
        title="Top 5 Articles par Quantité",
        color_discrete_sequence=["#3B82F6"],
        text_auto=True
    )
    fig_quantity.update_layout(xaxis_title="Article", yaxis_title="Quantité", xaxis_tickangle=45)
    fig_quantity.update_traces(
        texttemplate="%{y:,.2f}",
        textposition="auto",
        hovertemplate="<b>Article:</b> %{x}<br><b>Quantité:</b> %{y:,.2f}"
    )
    quantity_img = "quantity.png"
    save_plotly_fig_as_image(fig_quantity, quantity_img)
    doc.add_picture(quantity_img, width=Inches(6))

    # Top 5 Articles par Coût
    doc.add_heading("Top 5 Articles par Coût", level=2)
    doc.add_paragraph(
        "Ce graphique en barres montre les 5 articles ayant les coûts totaux les plus élevés. Il aide à repérer les articles les plus coûteux de l'inventaire, "
        "qui peuvent nécessiter une gestion des coûts ou une renégociation avec les fournisseurs."
    )
    fig_cost = px.bar(
        top_cost,
        x="DES_ARTICLE",
        y="MONTANT",
        title="Top 5 Articles par Coût",
        color_discrete_sequence=["#10B981"],
        text_auto=True
    )
    fig_cost.update_layout(xaxis_title="Article", yaxis_title="Coût", xaxis_tickangle=45)
    fig_cost.update_traces(
        texttemplate="%{y:,.2f}",
        textposition="auto",
        hovertemplate="<b>Article:</b> %{x}<br><b>Coût:</b> %{y:,.2f}"
    )
    cost_img = "cost.png"
    save_plotly_fig_as_image(fig_cost, cost_img)
    doc.add_picture(cost_img, width=Inches(6))

    # Répartition des Coûts par Groupe
    doc.add_heading("Répartition des Coûts par Groupe", level=2)
    doc.add_paragraph(
        "Ce graphique en secteurs illustre la proportion du coût total attribuée à chaque groupe. Il aide à comprendre quels groupes contribuent le plus aux dépenses, "
        "guidant l'allocation des ressources et les efforts d'optimisation des coûts."
    )
    fig_pie = px.pie(
        cost_by_group,
        names="GROUPE",
        values="MONTANT",
        title="Répartition des Coûts par Groupe",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )
    fig_pie.update_traces(
        textinfo="percent+label",
        hovertemplate="<b>Groupe:</b> %{label}<br><b>Coût:</b> %{value:,.2f}<br><b>Pourcentage:</b> %{percent}"
    )
    pie_img = "pie.png"
    save_plotly_fig_as_image(fig_pie, pie_img)
    doc.add_picture(pie_img, width=Inches(6))

    # Coût Total par Groupe au Fil du Temps
    doc.add_heading("Coût Total par Groupe au Fil du Temps", level=2)
    doc.add_paragraph(
        "Ce graphique en barres montre le coût total pour chaque groupe au fil du temps (par mois). Il aide à identifier les tendances des dépenses par groupe, "
        "révélant quels groupes ont des coûts croissants ou décroissants au fil du temps."
    )
    fig_time = px.bar(
        cost_by_time,
        x="Mois",
        y="MONTANT",
        color="GROUPE",
        title="Coût Total par Groupe au Fil du Temps",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )
    fig_time.update_layout(xaxis_title="Mois", yaxis_title="Coût Total", xaxis_tickangle=45)
    fig_time.update_traces(
        hovertemplate="<b>Mois:</b> %{x}<br><b>Groupe:</b> %{fullData.name}<br><b>Coût:</b> %{y:,.2f}"
    )
    time_img = "time.png"
    save_plotly_fig_as_image(fig_time, time_img)
    doc.add_picture(time_img, width=Inches(6))

    # Répartition de la Quantité par Groupe
    doc.add_heading("Répartition de la Quantité par Groupe", level=2)
    doc.add_paragraph(
        "Ce graphique en boîte affiche la répartition des quantités pour chaque groupe, mettant en évidence les médianes, les quartiles et les valeurs aberrantes. "
        "Il aide à identifier la variabilité des niveaux de stock au sein des groupes, utile pour détecter les catégories surstockées ou sous-stockées."
    )
    fig_box = px.box(
        filtered_df_box,
        x="GROUPE",
        y="QUANTITE",
        title="Répartition de la Quantité par Groupe",
        color="GROUPE",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )
    fig_box.update_layout(xaxis_title="Groupe", yaxis_title="Quantité", xaxis_tickangle=45, showlegend=False)
    fig_box.update_traces(
        hovertemplate="<b>Groupe:</b> %{x}<br><b>Quantité:</b> %{y:,.2f}"
    )
    box_img = "box.png"
    save_plotly_fig_as_image(fig_box, box_img)
    doc.add_picture(box_img, width=Inches(6))

    # Données de Stock
    doc.add_heading("Données de Stock", level=1)
    doc.add_paragraph("Tableau des données de stock (première page) :")
    if not paged_df.empty:
        num_cols = len(paged_df.columns)
        num_rows = len(paged_df) + 1  # +1 for header
        table = doc.add_table(rows=num_rows, cols=num_cols)
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(paged_df.columns):
            hdr_cells[i].text = col
        for row_idx, row in paged_df.iterrows():
            row_cells = table.rows[row_idx + 1].cells
            if len(row_cells) != num_cols:
                st.error(f"Erreur: Nombre de cellules ({len(row_cells)}) ne correspond pas au nombre de colonnes ({num_cols}) pour la ligne {row_idx + 1} dans 'Paged Data'")
                continue
            for col_idx in range(num_cols):
                try:
                    value = row.iloc[col_idx]
                    row_cells[col_idx].text = sanitize_text(value)
                except Exception as e:
                    st.error(f"Erreur lors du remplissage de la table 'Paged Data' à la ligne {row_idx + 1}, colonne {col_idx}: {e}")
                    row_cells[col_idx].text = "Erreur"
    else:
        doc.add_paragraph("Aucune donnée de stock disponible.")

    # Alertes (si affichées)
    if alerts_df is not None and not alerts_df.empty:
        doc.add_heading("Alertes", level=1)
        doc.add_paragraph(f"{len(alerts_df)} alertes trouvées.")
        num_cols = len(alerts_df.columns)
        num_rows = len(alerts_df) + 1  # +1 for header
        table = doc.add_table(rows=num_rows, cols=num_cols)
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(alerts_df.columns):
            hdr_cells[i].text = col
        for row_idx, row in alerts_df.iterrows():
            row_cells = table.rows[row_idx + 1].cells
            if len(row_cells) != num_cols:
                st.error(f"Erreur: Nombre de cellules ({len(row_cells)}) ne correspond pas au nombre de colonnes ({num_cols}) pour la ligne {row_idx + 1} dans 'Alerts'")
                continue
            for col_idx in range(num_cols):
                try:
                    value = row.iloc[col_idx]
                    row_cells[col_idx].text = sanitize_text(value)
                except Exception as e:
                    st.error(f"Erreur lors du remplissage de la table 'Alerts' à la ligne {row_idx + 1}, colonne {col_idx}: {e}")
                    row_cells[col_idx].text = "Erreur"
    elif alerts_df is not None:
        doc.add_paragraph("Aucune alerte trouvée.")

    # Sauvegarder le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Charger et nettoyer les données
@st.cache_data
def load_data():
    file_path = "stock.xlsx"
    if not os.path.exists(file_path):
        st.error(f"Le fichier {file_path} n'a pas été trouvé.")
        return pd.DataFrame(), pd.DataFrame()
    
    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        st.error(f"Erreur lors du chargement du fichier : {str(e)}.")
        return pd.DataFrame(), pd.DataFrame()
    
    # Nettoyer les données
    required_columns = ['article', 'DES_ARTICLE', 'GROUPE', 'QUANTITE', 'MONTANT', 'Mois']
    if not all(col in df.columns for col in required_columns):
        st.error("Colonnes requises manquantes dans l'ensemble de données.")
        return pd.DataFrame(), pd.DataFrame()
    
    df["QUANTITE"] = pd.to_numeric(df["QUANTITE"], errors="coerce").fillna(0)
    df["MONTANT"] = pd.to_numeric(df["MONTANT"], errors="coerce").fillna(0)
    df["DES_ARTICLE"] = df["DES_ARTICLE"].astype(str).str.strip()
    df["GROUPE"] = df["GROUPE"].astype(str).str.strip()
    df["Mois"] = df["Mois"].astype(str).str.strip()
    df["unit_price"] = df.apply(lambda row: row["MONTANT"] / row["QUANTITE"] if row["QUANTITE"] > 0 else 0, axis=1)
    
    # Créer une version agrégée pour éviter les doublons dans les visualisations par produit
    df_aggregated = df.groupby("article").agg({
        "QUANTITE": "sum",
        "MONTANT": "sum",
        "DES_ARTICLE": "first",
        "GROUPE": "first",
        "Mois": "first"
    }).reset_index()
    # Recalculer unit_price après agrégation
    df_aggregated["unit_price"] = df_aggregated.apply(
        lambda row: row["MONTANT"] / row["QUANTITE"] if row["QUANTITE"] > 0 else 0, axis=1
    )
    
    return df, df_aggregated

# Charger les données
df, df_aggregated = load_data()

# Vérifier si les données sont chargées
if df.empty or df_aggregated.empty:
    st.stop()

# Configuration de la page Streamlit
st.set_page_config(page_title="Tableau de Bord des Stocks", layout="wide")

# Titre
st.title("Tableau de Bord des Stocks")

# Filtres interactifs dans la barre latérale
st.sidebar.header("Filtres")
groupe_options = ["Tous"] + sorted(df["GROUPE"].unique().tolist())

selected_groupe = st.sidebar.selectbox("Sélectionner un Groupe", groupe_options)

# Appliquer les filtres
filtered_df = df.copy()
filtered_df_aggregated = df_aggregated.copy()

if selected_groupe != "Tous":
    filtered_df = filtered_df[filtered_df["GROUPE"] == selected_groupe]
    filtered_df_aggregated = filtered_df_aggregated[filtered_df_aggregated["GROUPE"] == selected_groupe]


# Métriques récapitulatives (globales, non filtrées)
st.header("📊 Métriques Globales")
total_items = len(df_aggregated)
total_quantity = df_aggregated["QUANTITE"].sum()
total_cost = df_aggregated["MONTANT"].sum()
unique_groups = df_aggregated["GROUPE"].nunique()

col1, col2, col3, col4 = st.columns(4)
col1.metric("Nombre Total d'Articles", format_number(total_items))
col2.metric("Quantité Totale", format_number(total_quantity))
col3.metric("Coût Total", format_number(total_cost))
col4.metric("Groupes Uniques", unique_groups)

# Aperçu intéressant (basé sur les données agrégées filtrées)
st.header("🔍 Aperçu Intéressant")
if not filtered_df_aggregated.empty:
    most_expensive = filtered_df_aggregated.loc[filtered_df_aggregated["unit_price"].idxmax()]
    st.info(
        f"L'article le plus cher par unité est **{most_expensive['DES_ARTICLE']}** "
        f"à **{format_number(most_expensive['unit_price'])}** par unité."
    )
else:
    st.warning("Aucune donnée disponible pour l'aperçu intéressant.")

# Exporter les données
st.header("💾 Exporter les Données")
col_csv, _ = st.columns(2)

with col_csv:
    st.download_button(
        label="Télécharger les Données en CSV",
        data=convert_df_to_csv(filtered_df_aggregated),
        file_name="donnees_stocks.csv",
        mime="text/csv",
    )

# Analyses et Visualisations Avancées
st.header("📈 Analyses et Visualisations Avancées")

# 1. Analyse du Taux de Rotation des Stocks
with st.container():
    st.subheader("Taux de Rotation des Stocks (Articles à Rotation Rapide vs Lente)")
    st.markdown(
        "Ce graphique montre les 5 articles à rotation rapide et les 5 articles à rotation lente en fonction de leur taux de rotation (quantité par période). "
        "Il aide à identifier les articles qui sont utilisés ou vendus rapidement par rapport à ceux qui stagnent, facilitant l'optimisation des stocks."
    )
    turnover_df = filtered_df_aggregated[["DES_ARTICLE", "QUANTITE"]].copy()
    total_periods = df["Mois"].nunique()
    turnover_df["Taux de Rotation"] = turnover_df["QUANTITE"] / total_periods
    turnover_df = turnover_df.sort_values("Taux de Rotation", ascending=False)
    top_turnover = turnover_df.head(5)
    bottom_turnover = turnover_df.tail(5)
    turnover_combined = pd.concat([top_turnover, bottom_turnover])
    fig_turnover = px.bar(
        turnover_combined,
        x="Taux de Rotation",
        y="DES_ARTICLE",
        orientation="h",
        title="Top 5 Articles à Rotation Rapide et Lente (Taux de Rotation)",
        color="Taux de Rotation",
        color_continuous_scale="Viridis",
        text_auto=True
    )
    fig_turnover.update_layout(xaxis_title="Taux de Rotation (Quantité par Période)", yaxis_title="Article")
    st.plotly_chart(fig_turnover, use_container_width=True)

# 2. Tendance des Coûts au Fil du Temps
with st.container():
    st.subheader("Tendance des Coûts au Fil du Temps")
    st.markdown(
        "Ce graphique en ligne montre le coût total au fil du temps (par mois). Il permet de suivre les fluctuations des coûts, d'identifier les tendances saisonnières, "
        "et de repérer les périodes de dépenses inhabituellement élevées ou faibles, utile pour la planification budgétaire et financière."
    )
    cost_trend = df.groupby("Mois")["MONTANT"].sum().reset_index()
    st.write("Debug - Cost Trend Data:", cost_trend)  # Ajout pour débogage
    fig_cost_trend = px.line(
        cost_trend,
        x="Mois",
        y="MONTANT",
        title="Tendance du Coût Total au Fil du Temps",
        markers=True,
        color_discrete_sequence=["#FF4D4F"]
    )
    fig_cost_trend.update_layout(xaxis_title="Mois", yaxis_title="Coût Total", xaxis_tickangle=45)
    fig_cost_trend.update_traces(
        hovertemplate="<b>Mois:</b> %{x}<br><b>Coût:</b> %{y:,.2f}",
        line=dict(width=3)
    )
    st.plotly_chart(fig_cost_trend, use_container_width=True)

# 3. Top 10 Articles les Plus Chers par Prix Unitaire
with st.container():
    st.subheader("Top 10 Articles les Plus Chers par Prix Unitaire")
    st.markdown(
        "Ce tableau liste les 10 articles ayant les prix unitaires les plus élevés, agrégés pour éviter les doublons. Il montre la quantité totale et le coût "
        "pour chaque article, aidant à identifier les articles de grande valeur qui peuvent nécessiter une attention particulière dans les stratégies de tarification ou d'approvisionnement."
    )
    top_expensive = filtered_df_aggregated.nlargest(10, "unit_price")[["article", "DES_ARTICLE", "GROUPE", "QUANTITE", "MONTANT", "unit_price"]]
    st.dataframe(
        top_expensive.style.format({"QUANTITE": "{:.2f}", "MONTANT": "{:.2f}", "unit_price": "{:.2f}"}),
        use_container_width=True
    )

# 4. Analyse ABC
with st.container():
    st.subheader("Analyse ABC des Articles")
    st.markdown(
        "Ce graphique montre la répartition des articles selon l'analyse ABC : les articles de catégorie A (20% des articles, 80% du coût), B (30% des articles, 15% du coût), "
        "et C (50% des articles, 5% du coût). Cela aide à prioriser la gestion des stocks en se concentrant sur les articles les plus coûteux."
    )
    abc_df = filtered_df_aggregated[["DES_ARTICLE", "MONTANT"]].copy()
    abc_df = abc_df.sort_values("MONTANT", ascending=False)
    abc_df["Pourcentage Cumulatif"] = abc_df["MONTANT"].cumsum() / abc_df["MONTANT"].sum() * 100
    abc_df["Catégorie"] = "C"
    abc_df.loc[abc_df["Pourcentage Cumulatif"] <= 80, "Catégorie"] = "A"
    abc_df.loc[(abc_df["Pourcentage Cumulatif"] > 80) & (abc_df["Pourcentage Cumulatif"] <= 95), "Catégorie"] = "B"
    abc_summary = abc_df.groupby("Catégorie")["MONTANT"].sum().reset_index()
    fig_abc = px.pie(
        abc_summary,
        names="Catégorie",
        values="MONTANT",
        title="Répartition des Coûts par Catégorie ABC",
        color_discrete_sequence=["#FF4D4F", "#FFA500", "#3B82F6"]
    )
    fig_abc.update_traces(
        textinfo="percent+label",
        hovertemplate="<b>Catégorie:</b> %{label}<br><b>Coût:</b> %{value:,.2f}<br><b>Pourcentage:</b> %{percent}"
    )
    st.plotly_chart(fig_abc, use_container_width=True)

# Répartition de l'Inventaire
st.header("📉 Répartition de l'Inventaire")

# Top 5 Articles par Quantité et par Coût
with st.container():
    col_chart1, col2 = st.columns(2)
    
    with col_chart1:
        st.subheader("Top 5 Articles par Quantité")
        st.markdown(
            "Ce graphique en barres met en évidence les 5 articles ayant les quantités totales les plus élevées en stock. Il aide à identifier les articles les plus stockés, "
            "ce qui peut indiquer une forte demande ou un éventuel surstockage."
        )
        top_quantity = filtered_df_aggregated.nlargest(5, "QUANTITE")[["DES_ARTICLE", "QUANTITE"]]
        fig_quantity = px.bar(
            top_quantity,
            x="DES_ARTICLE",
            y="QUANTITE",
            title="Top 5 Articles par Quantité",
            color_discrete_sequence=["#3B82F6"],
            text_auto=True
        )
        fig_quantity.update_layout(xaxis_title="Article", yaxis_title="Quantité", xaxis_tickangle=45)
        fig_quantity.update_traces(
            texttemplate="%{y:,.2f}",
            textposition="auto",
            hovertemplate="<b>Article:</b> %{x}<br><b>Quantité:</b> %{y:,.2f}"
        )
        st.plotly_chart(fig_quantity, use_container_width=True)

    with col2:
        st.subheader("Top 5 Articles par Coût")
        st.markdown(
            "Ce graphique en barres montre les 5 articles ayant les coûts totaux les plus élevés. Il aide à repérer les articles les plus coûteux de l'inventaire, "
            "qui peuvent nécessiter une gestion des coûts ou une renégociation avec les fournisseurs."
        )
        top_cost = filtered_df_aggregated.nlargest(5, "MONTANT")[["DES_ARTICLE", "MONTANT"]]
        fig_cost = px.bar(
            top_cost,
            x="DES_ARTICLE",
            y="MONTANT",
            title="Top 5 Articles par Coût",
            color_discrete_sequence=["#10B981"],
            text_auto=True
        )
        fig_cost.update_layout(xaxis_title="Article", yaxis_title="Coût", xaxis_tickangle=45)
        fig_cost.update_traces(
            texttemplate="%{y:,.2f}",
            textposition="auto",
            hovertemplate="<b>Article:</b> %{x}<br><b>Coût:</b> %{y:,.2f}"
        )
        st.plotly_chart(fig_cost, use_container_width=True)

# Répartition des Coûts par Groupe
with st.container():
    st.subheader("Répartition des Coûts par Groupe")
    st.markdown(
        "Ce graphique en secteurs illustre la proportion du coût total attribuée à chaque groupe. Il aide à comprendre quels groupes contribuent le plus aux dépenses, "
        "guidant l'allocation des ressources et les efforts d'optimisation des coûts."
    )
    cost_by_group = filtered_df.groupby("GROUPE")["MONTANT"].sum().reset_index()
    cost_by_group = cost_by_group[cost_by_group["MONTANT"] > 0]
    fig_pie = px.pie(
        cost_by_group,
        names="GROUPE",
        values="MONTANT",
        title="Répartition des Coûts par Groupe",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )
    fig_pie.update_traces(
        textinfo="percent+label",
        hovertemplate="<b>Groupe:</b> %{label}<br><b>Coût:</b> %{value:,.2f}<br><b>Pourcentage:</b> %{percent}"
    )
    st.plotly_chart(fig_pie, use_container_width=True)

# Coût Total par Groupe au Fil du Temps
with st.container():
    st.subheader("Coût Total par Groupe au Fil du Temps")
    st.markdown(
        "Ce graphique en barres montre le coût total pour chaque groupe au fil du temps (par mois). Il aide à identifier les tendances des dépenses par groupe, "
        "révélant quels groupes ont des coûts croissants ou décroissants au fil du temps."
    )
    cost_by_time = filtered_df.groupby(["Mois", "GROUPE"])["MONTANT"].sum().reset_index()
    fig_time = px.bar(
        cost_by_time,
        x="Mois",
        y="MONTANT",
        color="GROUPE",
        title="Coût Total par Groupe au Fil du Temps",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )
    fig_time.update_layout(xaxis_title="Mois", yaxis_title="Coût Total", xaxis_tickangle=45)
    fig_time.update_traces(
        hovertemplate="<b>Mois:</b> %{x}<br><b>Groupe:</b> %{fullData.name}<br><b>Coût:</b> %{y:,.2f}"
    )
    st.plotly_chart(fig_time, use_container_width=True)

# Répartition de la Quantité par Groupe
with st.container():
    st.subheader("Répartition de la Quantité par Groupe")
    st.markdown(
        "Ce graphique en boîte affiche la répartition des quantités pour chaque groupe, mettant en évidence les médianes, les quartiles et les valeurs aberrantes. "
        "Il aide à identifier la variabilité des niveaux de stock au sein des groupes, utile pour détecter les catégories surstockées ou sous-stockées."
    )
    fig_box = px.box(
        filtered_df,
        x="GROUPE",
        y="QUANTITE",
        title="Répartition de la Quantité par Groupe",
        color="GROUPE",
        color_discrete_sequence=px.colors.qualitative.Plotly
    )
    fig_box.update_layout(xaxis_title="Groupe", yaxis_title="Quantité", xaxis_tickangle=45, showlegend=False)
    fig_box.update_traces(
        hovertemplate="<b>Groupe:</b> %{x}<br><b>Quantité:</b> %{y:,.2f}"
    )
    st.plotly_chart(fig_box, use_container_width=True)

# Tableau des Données de Stock avec Pagination
st.header("📋 Données de Stock")
items_per_page = 10
total_pages = (len(filtered_df_aggregated) + items_per_page - 1) // items_per_page
page = st.slider("Page", 1, max(1, total_pages), 1)
start_idx = (page - 1) * items_per_page
end_idx = start_idx + items_per_page
paged_df = filtered_df_aggregated.iloc[start_idx:end_idx]
st.dataframe(
    paged_df[["article", "DES_ARTICLE", "GROUPE", "QUANTITE", "MONTANT", "unit_price"]].style.format(
        {"QUANTITE": "{:.2f}", "MONTANT": "{:.2f}", "unit_price": "{:.2f}"}
    ),
    use_container_width=True
)

# Section Alertes
st.header("⚠️ Alertes")
show_alerts = st.checkbox("Afficher les Alertes (Cela peut prendre du temps à charger)", value=False)
alerts_df = None
if show_alerts:
    # Conditions d'alerte vectorisées
    low_stock = filtered_df_aggregated[filtered_df_aggregated["QUANTITE"] < 10].copy()
    low_stock["Raison de l'Alerte"] = low_stock["QUANTITE"].apply(lambda x: f"Stock Faible : Quantité = {x:.2f}")

    high_cost = filtered_df_aggregated[filtered_df_aggregated["MONTANT"] > 100000].copy()
    high_cost["Raison de l'Alerte"] = high_cost["MONTANT"].apply(lambda x: f"Coût Élevé : Coût = {format_number(x)}")

    high_unit_price = filtered_df_aggregated[filtered_df_aggregated["unit_price"] > 1000].copy()
    high_unit_price["Raison de l'Alerte"] = high_unit_price["unit_price"].apply(lambda x: f"Prix Unitaire Élevé : {format_number(x)}")

    data_issues = filtered_df_aggregated[(filtered_df_aggregated["QUANTITE"] <= 0) | (filtered_df_aggregated["MONTANT"] <= 0)].copy()
    data_issues["Raison de l'Alerte"] = "Problème de Données : Quantité ou Coût Invalide"

    # Concaténer les alertes
    alerts_df = pd.concat([low_stock, high_cost, high_unit_price, data_issues], ignore_index=True)

    # Limiter à 100 alertes pour un rendu plus rapide
    if len(alerts_df) > 100:
        alerts_df = alerts_df.head(100)
        st.warning(f"Affichage des 100 premières alertes sur {len(alerts_df)} en raison de contraintes de performance.")
    else:
        st.warning(f"{len(alerts_df)} alertes trouvées.")

    if not alerts_df.empty:
        with st.expander("Voir les Alertes", expanded=False):
            st.dataframe(
                alerts_df[["article", "DES_ARTICLE", "GROUPE", "QUANTITE", "MONTANT", "unit_price", "Raison de l'Alerte"]],
                use_container_width=True
            )
    else:
        st.success("Aucune alerte trouvée.")

# Informations sur la pagination
st.write(f"Page {page} sur {total_pages}")

# Téléchargement du Rapport Word à la fin
st.header("📥 Téléchargement du Rapport")
col_docx, _ = st.columns(2)

with col_docx:
    # Préparer les données pour le document Word
    turnover_df = filtered_df_aggregated[["DES_ARTICLE", "QUANTITE"]].copy()
    total_periods = df["Mois"].nunique()
    turnover_df["Taux de Rotation"] = turnover_df["QUANTITE"] / total_periods
    turnover_df = turnover_df.sort_values("Taux de Rotation", ascending=False)
    top_turnover = turnover_df.head(5)
    bottom_turnover = turnover_df.tail(5)
    turnover_combined = pd.concat([top_turnover, bottom_turnover])

    cost_trend = df.groupby("Mois")["MONTANT"].sum().reset_index()
    top_expensive = filtered_df_aggregated.nlargest(10, "unit_price")[["article", "DES_ARTICLE", "GROUPE", "QUANTITE", "MONTANT", "unit_price"]]

    abc_df = filtered_df_aggregated[["DES_ARTICLE", "MONTANT"]].copy()
    abc_df = abc_df.sort_values("MONTANT", ascending=False)
    abc_df["Pourcentage Cumulatif"] = abc_df["MONTANT"].cumsum() / abc_df["MONTANT"].sum() * 100
    abc_df["Catégorie"] = "C"
    abc_df.loc[abc_df["Pourcentage Cumulatif"] <= 80, "Catégorie"] = "A"
    abc_df.loc[(abc_df["Pourcentage Cumulatif"] > 80) & (abc_df["Pourcentage Cumulatif"] <= 95), "Catégorie"] = "B"
    abc_summary = abc_df.groupby("Catégorie")["MONTANT"].sum().reset_index()

    top_quantity = filtered_df_aggregated.nlargest(5, "QUANTITE")[["DES_ARTICLE", "QUANTITE"]]
    top_cost = filtered_df_aggregated.nlargest(5, "MONTANT")[["DES_ARTICLE", "MONTANT"]]
    cost_by_group = filtered_df.groupby("GROUPE")["MONTANT"].sum().reset_index()
    cost_by_group = cost_by_group[cost_by_group["MONTANT"] > 0]
    cost_by_time = filtered_df.groupby(["Mois", "GROUPE"])["MONTANT"].sum().reset_index()

    doc_buffer = create_word_doc(
        df, df_aggregated, filtered_df, filtered_df_aggregated,
        total_items, total_quantity, total_cost, unique_groups, most_expensive,
        turnover_combined, cost_trend, top_expensive, abc_summary,
        top_quantity, top_cost, cost_by_group, cost_by_time, filtered_df, paged_df, alerts_df
    )
    st.download_button(
        label="Télécharger le Rapport en Word",
        data=doc_buffer,
        file_name="rapport_stocks.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )