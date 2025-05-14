import streamlit as st
import pandas as pd
import plotly.express as px
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import plotly.io as pio
from PIL import Image

# Setting page configuration
st.set_page_config(page_title="Mining Equipment Consumption Dashboard", layout="wide")
st.markdown("""
    <style>
    .stApp { 
        background-image: url("https://img.freepik.com/premium-photo/underground-mining-truck_873668-11862.jpg"); 
        background-size: cover; 
        background-repeat: no-repeat; 
    }
    .stApp > div { 
        padding: 20px; 
        border-radius: 10px; 
    }
    h1, h2, h3 { 
        color: #003087; 
        font-family: Arial, sans-serif; 
    }
    .stMetric { 
        background-color: #Ff7f00; 
        border-left: 5px solid #FFC107; 
        padding: 10px; 
        border-radius: 5px; 
    }
    .stButton>button { 
        background-color: #003087; 
        color: white; 
        border-radius: 5px; 
    }
    .stButton>button:hover { 
        background-color: #FFC107; 
        color: black; 
    }
    </style>
""", unsafe_allow_html=True)

# Adding title and KPIs
st.markdown("""
<div style='background-color:#e3f2fd; padding:20px; border-radius:10px; border-left:5px solid #1976d2; margin-bottom:20px;'>
    <h1 style='color:#F28C38; text-align:center; margin-top:0;'>📊 Mining Equipment Consumption Dashboard</h1>
    <p style='color:#424242; text-align:center;'>Track and optimize equipment consumption</p>
</div>
""", unsafe_allow_html=True)



# Loading data
@st.cache_data
def load_data():
    df = pd.read_excel("engins2.xlsx")
    if pd.api.types.is_numeric_dtype(df['Date']):
        df['Date'] = pd.to_datetime(df['Date'], origin='1899-12-30', unit='D')
    elif not pd.api.types.is_datetime64_any_dtype(df['Date']):
        df['Date'] = pd.to_datetime(df['Date'])
    df = df.dropna(subset=['CATEGORIE', 'Desc_Cat', 'Desc_CA', 'Montant'])
    df['Montant'] = pd.to_numeric(df['Montant'], errors='coerce')
    df['Mois'] = df['Date'].dt.month_name()
    months_fr = {
        'January': 'Janvier', 'February': 'Février', 'March': 'Mars',
        'April': 'Avril', 'May': 'Mai', 'June': 'Juin',
        'July': 'Juillet', 'August': 'Août', 'September': 'Septembre',
        'October': 'Octobre', 'November': 'Novembre', 'December': 'Décembre'
    }
    df['Mois'] = df['Mois'].map(months_fr)
    
    return df

df = load_data()

# Cache computations
@st.cache_data
def compute_monthly_costs(data):
    monthly_data = data.groupby('Mois')['Montant'].sum().reset_index()
    month_order = ['Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin',
                   'Juillet', 'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre']
    monthly_data['Mois'] = pd.Categorical(monthly_data['Mois'], categories=month_order, ordered=True)
    return monthly_data.sort_values('Mois')

@st.cache_data
def compute_category_breakdown(data):
    return data.groupby('Desc_Cat')['Montant'].sum().reset_index()

# Function to save Plotly figure as image
def save_plotly_fig_as_image(fig):
    img_bytes = pio.to_image(fig, format='png', width=800, height=350)
    img = Image.open(BytesIO(img_bytes))
    img_buffer = BytesIO()
    img.save(img_buffer, format='PNG')
    return img_buffer.getvalue()

# Function to generate Word report
def generate_word_report(engin_data, selected_category, figs, descriptions, metrics, predictions, budget_threshold):
    doc = Document()
    
    title = doc.add_heading(f'Consumption Report for {selected_category}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Report Date: {datetime.now().strftime("%d/%m/%Y")}')
    
    doc.add_heading('Table of Contents', level=2)
    doc.add_paragraph('1. Introduction\n2. Summary\n3. Visualizations\n4. Key Metrics\n5. Predictions\n6. Recommendations', style='List Bullet')
    
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'This report provides a detailed analysis of equipment consumption, including visualizations, metrics, predictions, and recommendations for optimization.'
    )
    
    doc.add_heading('2. Summary', level=2)
    total_cost = engin_data['Montant'].sum()
    top_category = engin_data.groupby('Desc_Cat')['Montant'].sum().idxmax()
    trend = predictions.get('trend', 'Stable')
    doc.add_paragraph(
        f'Total Cost: {total_cost:,.0f} DH\n'
        f'Main Category: {top_category}\n'
        f'Recent Trend: {trend}\n'
        f'Number of Interventions: {len(engin_data)}\n'
        f'Median Cost: {engin_data["Montant"].median():,.0f} DH'
    )
    
    doc.add_heading('3. Visualizations', level=2)
    for i, (fig, desc, title) in enumerate(zip(figs, descriptions, [
        'Consumption Trend with Projection', 'Cost Distribution',
        'Consumption by Type', 'Monthly Costs'
    ])):
        doc.add_heading(f'3.{i+1} {title}', level=3)
        img_stream = BytesIO(save_plotly_fig_as_image(fig))
        doc.add_picture(img_stream, width=Inches(6))
        doc.add_paragraph(desc)
    
    doc.add_heading('4. Key Metrics', level=2)
    table = doc.add_table(rows=len(metrics) + 1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    for i, (metric, value) in enumerate(metrics.items()):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)
    
    doc.add_heading('5. Predictions', level=2)
    doc.add_paragraph(
        f'Average Estimate: {predictions["avg"]:,.0f} DH/month\n'
        f'Confidence Interval (95%): {predictions["ci_lower"]:,.0f} - {predictions["ci_upper"]:,.0f} DH\n'
        f'Reliability: {predictions["reliability"]:,.0f}%\n'
        f'Trend: {predictions["trend"]}'
    )
    
    if budget_threshold > 0:
        high_costs = engin_data[engin_data['Montant'] > budget_threshold]
        doc.add_heading('Budget Threshold Analysis', level=3)
        doc.add_paragraph(
            f'Threshold: {budget_threshold:,.0f} DH\n'
            f'Number of Interventions Exceeding Threshold: {len(high_costs)}\n'
            f'Total Cost of High Interventions: {high_costs["Montant"].sum():,.0f} DH'
        )
    
    doc.add_heading('6. Recommendations', level=2)
    doc.add_paragraph(
        '- Prioritize preventive maintenance to reduce costs in the main category.\n'
        f'- Review high-cost interventions in {top_category} for cost-saving opportunities.\n'
        '- Implement monthly monitoring to detect rising trends early.\n'
        '- Negotiate with suppliers for frequently used parts.'
    )
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Sidebar for filters
with st.sidebar:
    st.subheader("Filters")
    
    selected_category = st.radio(
        'Select Category',
        ['All'] + sorted(df['CATEGORIE'].unique().tolist()),
        help="Filter by specific category"
    )
    
    st.subheader("Date Range")
    default_start = df['Date'].min().date()
    default_end = df['Date'].max().date()
    date_range = st.date_input(
        "Period",
        value=(default_start, default_end),
        min_value=default_start,
        max_value=default_end,
        help="Choose a date range to filter interventions"
    )
    
    filtered_data = df.copy()
    if selected_category != 'All':
        filtered_data = filtered_data[filtered_data['CATEGORIE'] == selected_category]
    if len(date_range) == 2:
        start_date, end_date = date_range
        filtered_data = filtered_data[(filtered_data['Date'].dt.date >= start_date) & 
                                    (filtered_data['Date'].dt.date <= end_date)]
    
    st.subheader("Search Equipment")
    equipment_search = st.text_input("Enter Equipment Name (partial match)", "").strip()
    category_df = filtered_data
    if equipment_search:
        available_equipment = sorted(category_df[category_df['Desc_CA'].str.contains(equipment_search, case=False, na=False)]['Desc_CA'].unique())
    else:
        available_equipment = sorted(category_df['Desc_CA'].unique())
    equipment_options = ["All Equipment"] + available_equipment
    if not available_equipment:
        st.warning("No equipment matches the search term.")
    selected_equipment = st.selectbox("Select Equipment", equipment_options)
    
    
# Apply filters
filtered_data = df.copy()
if selected_category != 'All':
    filtered_data = filtered_data[filtered_data['CATEGORIE'] == selected_category]
if len(date_range) == 2:
    start_date, end_date = date_range
    filtered_data = filtered_data[(filtered_data['Date'].dt.date >= start_date) & 
                                (filtered_data['Date'].dt.date <= end_date)]

if selected_equipment != "All Equipment":
    filtered_data = filtered_data[filtered_data['Desc_CA'] == selected_equipment]

if filtered_data.empty:
    st.warning("No data available after filtering. Please adjust the filters.")
    st.stop()

kpi_container = st.container()
with kpi_container:
    kpi1, kpi2, kpi3, kpi4 = st.columns(4)
    with kpi1:
        total_cost = filtered_data['Montant'].sum()
        st.markdown(f"""
        <div style='background-color:#fff8e1; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Total Cost</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{total_cost:,.0f} DH</p>
            <p style='color:#424242; font-size:12px;'>Total expenditure</p>
        </div>
        """, unsafe_allow_html=True)
    with kpi2:
        top_equipment = filtered_data.groupby('Desc_CA')['Montant'].sum().idxmax() if not filtered_data.empty else "N/A"
        st.markdown(f"""
        <div style='background-color:#e8f5e9; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Most Expensive Equipment</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{top_equipment}</p>
            <p style='color:#424242; font-size:12px;'>Highest cost equipment</p>
        </div>
        """, unsafe_allow_html=True)
    with kpi3:
        top_category = filtered_data.groupby('Desc_Cat')['Montant'].sum().idxmax() if not filtered_data.empty else "N/A"
        st.markdown(f"""
        <div style='background-color:#f3e5f5; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Main Category</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{top_category}</p>
            <p style='color:#424242; font-size:12px;'>Dominant consumption type</p>
        </div>
        """, unsafe_allow_html=True)
    with kpi4:
        avg_cost = filtered_data['Montant'].mean() if not filtered_data.empty else 0
        st.markdown(f"""
        <div style='background-color:#ffebee; padding:15px; border-radius:10px; text-align:center;'>
            <h6 style='color:#F28C38; margin-bottom:10px;'>Average Cost</h6>
            <p style='color:#424242; font-size:18px; font-weight:bold;'>{avg_cost:,.0f} DH</p>
            <p style='color:#424242; font-size:12px;'>Per intervention</p>
        </div>
        """, unsafe_allow_html=True)
# Tabs for organization
tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Data", "📊 Analysis", "🔄 Comparisons", "💡 Recommendations", "📋 Equipment Consumption Table"])

with tab1:
    st.markdown("""
    <div style='background-color:#e8f5e9; padding:20px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>Detailed Consumption Data</h2>
        <p style='color:#424242;'>View and export recorded interventions</p>
    </div>
    """, unsafe_allow_html=True)
    
    height = 600 if st.checkbox("Show all data (scrollable)") else 300
    display_df = filtered_data[['Date', 'Desc_CA', 'Desc_Cat', 'Montant']].copy()
    display_df['Date'] = display_df['Date'].dt.strftime('%d/%m/%Y')
    display_df['Montant'] = display_df['Montant'].round(2)
    display_df = display_df.sort_values('Date')
    display_df = display_df.rename(columns={
        'Date': 'Date',
        'Desc_CA': 'Equipment',
        'Desc_Cat': 'Consumption Type',
        'Montant': 'Amount (DH)'
    })
    st.dataframe(display_df, height=height, use_container_width=True)
    
    export_col1, export_col2 = st.columns([1, 3])
    with export_col1:
        export_format = st.radio("Format", ['CSV', 'Excel'])
    with export_col2:
        if export_format == 'CSV':
            csv = display_df.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Download CSV", csv, "consumption_data.csv", "text/csv")
        else:
            output = BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                display_df.to_excel(writer, index=False)
            st.download_button("📥 Download Excel", output.getvalue(), "consumption_data.xlsx", "application/vnd.ms-excel")

with tab2:
    if selected_category != 'All':
        engin_data = filtered_data.copy()
    else:
        engin_data = filtered_data.groupby(['Date', 'Desc_CA', 'Desc_Cat', 'CATEGORIE', 'Mois'])['Montant'].sum().reset_index()
    
    col1, col2 = st.columns([7, 3])
    
    with col1:
        st.markdown("""
        <div style='background-color:#fff8e1; padding:15px; border-radius:10px; border-left:5px solid #ffa000; margin-bottom:20px;'>
            <h3 style='color:#F28C38; margin-top:0;'>📈 Consumption Visualizations</h3>
        </div>
        """, unsafe_allow_html=True)
        
        figs = []
        descriptions = []
        
        st.markdown("#### Consumption Trend with Projection")
        fig1 = px.line(
            engin_data.groupby('Date')['Montant'].sum().reset_index(),
            x='Date', y='Montant',
            title='Consumption Trend with Projection',
            height=350,
            template='plotly_white'
        )
        if len(engin_data) >= 3:
            dates = engin_data.groupby('Date')['Montant'].sum().index
            x = np.arange(len(dates))
            y = engin_data.groupby('Date')['Montant'].sum().values
            coeff = np.polyfit(x, y, 1)
            future_dates = [dates[-1] + pd.DateOffset(months=i) for i in range(1, 4)]
            projection = np.polyval(coeff, [x[-1]+1, x[-1]+2, x[-1]+3])
            fig1.add_scatter(x=future_dates, y=projection, mode='lines+markers', name='Projection', line=dict(color='red', dash='dot'))
            fig1.add_scatter(x=future_dates, y=projection * 1.3, mode='lines', name='High Range (+30%)', line=dict(color='orange', dash='dash'))
            fig1.add_scatter(x=future_dates, y=projection * 0.7, mode='lines', name='Low Range (-30%)', line=dict(color='green', dash='dash'))
        fig1.update_layout(xaxis_title="Date", yaxis_title="Amount (DH)", legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5))
        st.plotly_chart(fig1, use_container_width=True)
        desc1 = "This line chart shows consumption trends with a 3-month projection and ±30% confidence intervals."
        st.markdown(f"<p style='color:#FFFFFF; font-size:14px;'>{desc1}</p>", unsafe_allow_html=True)
        figs.append(fig1)
        descriptions.append(desc1)
        
        st.markdown("#### Cost Distribution")
        fig2 = px.histogram(engin_data, x='Montant', title='Cost Distribution', height=350, template='plotly_white', nbins=20)
        fig2.update_traces(marker=dict(color='#1976d2'))
        fig2.update_layout(xaxis_title="Amount (DH)", yaxis_title="Frequency")
        st.plotly_chart(fig2, use_container_width=True)
        desc2 = "This histogram shows the frequency of cost values to identify outliers."
        st.markdown(f"<p style='color:#FFFFFF; font-size:14px;'>{desc2}</p>", unsafe_allow_html=True)
        figs.append(fig2)
        descriptions.append(desc2)
        
        st.markdown("#### Consumption by Type")
        fig3 = px.pie(compute_category_breakdown(engin_data), values='Montant', names='Desc_Cat', title='Consumption by Type', height=350, template='plotly_white')
        fig3.update_traces(textinfo='percent+label')
        fig3.update_layout(margin=dict(t=50, b=50))
        st.plotly_chart(fig3, use_container_width=True)
        desc3 = "This pie chart shows the distribution of consumption by type."
        st.markdown(f"<p style='color:#FFFFFF; font-size:14px;'>{desc3}</p>", unsafe_allow_html=True)
        figs.append(fig3)
        descriptions.append(desc3)
        
        st.markdown("#### Monthly Costs")
        monthly_data = compute_monthly_costs(engin_data)
        fig4 = px.bar(monthly_data, x='Mois', y='Montant', title='Monthly Costs', height=350, template='plotly_white')
        fig4.update_traces(marker=dict(color='#388e3c'))
        fig4.update_layout(xaxis_title="Month", yaxis_title="Amount (DH)")
        st.plotly_chart(fig4, use_container_width=True)
        desc4 = "This bar chart displays monthly consumption trends."
        st.markdown(f"<p style='color:#FFFFFF; font-size:14px;'>{desc4}</p>", unsafe_allow_html=True)
        figs.append(fig4)
        descriptions.append(desc4)
    
    with col2:
        st.markdown("""
        <div style='background-color:#f3e5f5; padding:15px; border-radius:10px; border-left:5px solid #8e24aa; margin-bottom:20px;'>
            <h3 style='color:#F28C38; margin-top:0;'>📊 Key Metrics</h3>
        </div>
        """, unsafe_allow_html=True)
        if not engin_data.empty:
            last_month = engin_data.groupby('Mois')['Montant'].sum().iloc[-1] if not engin_data.groupby('Mois')['Montant'].sum().empty else 0
            avg_3m = engin_data.groupby('Mois')['Montant'].sum().tail(3).mean() if len(engin_data.groupby('Mois')['Montant'].sum()) >= 3 else 0
            max_cost = engin_data['Montant'].max()
            total_cost = engin_data['Montant'].sum()
            num_interventions = len(engin_data)
            median_cost = engin_data['Montant'].median()
            cost_variance = engin_data['Montant'].var() if len(engin_data) > 1 else 0
            metrics = {
                'Last Month': f"{last_month:,.0f} DH",
                '3-Month Average': f"{avg_3m:,.0f} DH",
                'Max Cost': f"{max_cost:,.0f} DH",
                'Total Cost': f"{total_cost:,.0f} DH",
                'Interventions': num_interventions,
                'Median Cost': f"{median_cost:,.0f} DH",
                'Cost Variance': f"{cost_variance:,.0f} DH²"
            }
            st.markdown(f"""
            <div style='color:#FFFFFF; font-size:14px;'>
                <p><strong>Last Month:</strong> {last_month:,.0f} DH</p>
                <p><strong>3-Month Average:</strong> {avg_3m:,.0f} DH</p>
                <p><strong>Max Cost:</strong> {max_cost:,.0f} DH</p>
                <p><strong>Total Cost:</strong> {total_cost:,.0f} DH</p>
                <p><strong>Interventions:</strong> {num_interventions}</p>
                <p><strong>Median Cost:</strong> {median_cost:,.0f} DH</p>
                <p><strong>Cost Variance:</strong> {cost_variance:,.0f} DH²</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style='background-color:#e8f5e9; padding:15px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:20px;'>
            <h3 style='color:#F28C38; margin-top:0;'>🔮 Predictions</h3>
            <p style='color:#424242;'>Estimates based on historical data</p>
        </div>
        """, unsafe_allow_html=True)
        predictions = {}
        if len(engin_data) >= 3:
            last_3 = engin_data.groupby('Mois')['Montant'].sum().tail(3)
            avg = last_3.mean()
            std = last_3.std() if len(last_3) > 1 else 0
            num_months = len(engin_data.groupby('Mois'))
            reliability = min(90, 50 + 5 * num_months - 10 * (std / avg if avg > 0 else 0))
            reliability = max(50, reliability)
            ci_lower = avg - 1.96 * std / np.sqrt(len(last_3)) if std > 0 else avg * 0.7
            ci_upper = avg + 1.96 * std / np.sqrt(len(last_3)) if std > 0 else avg * 1.3
            trend = "Stable"
            if len(last_3) >= 2:
                diff = last_3.iloc[-1] - last_3.iloc[-2]
                if diff > 0.1 * avg:
                    trend = "Rising"
                elif diff < -0.1 * avg:
                    trend = "Falling"
            predictions = {'avg': avg, 'ci_lower': ci_lower, 'ci_upper': ci_upper, 'reliability': reliability, 'trend': trend}
            st.markdown(f"""
            <div style='color:#FFFFFF; font-size:14px;'>
                <p><strong>Average Estimate:</strong> {avg:,.0f} DH/month</p>
                <p><strong>Confidence Interval (95%):</strong> {ci_lower:,.0f} - {ci_upper:,.0f} DH</p>
                <p><strong>Reliability:</strong> {reliability:.0f}%</p>
                <p><strong>Trend:</strong> {trend}</p>
            </div>
            """, unsafe_allow_html=True)
            st.progress(int(reliability), "Prediction Reliability")
            if st.button("💾 Export Predictions"):
                future_dates = pd.date_range(start=engin_data['Date'].max() + pd.DateOffset(months=1), periods=3, freq='M')
                projections = pd.DataFrame({
                    'Month': future_dates.strftime('%Y-%m'),
                    'Average Estimate (DH)': [avg]*3,
                    'Lower CI (DH)': [ci_lower]*3,
                    'Upper CI (DH)': [ci_upper]*3,
                    'Reliability (%)': [reliability]*3
                })
                csv = projections.to_csv(index=False).encode('utf-8')
                st.download_button("Download CSV", csv, f'predictions_{selected_category}.csv', "text/csv")
        else:
            st.markdown("""
            <div style='background-color:#ffebee; padding:10px; border-radius:5px;'>
                <p style='color:#d32f2f;'>⚠ Insufficient data (minimum 3 months required)</p>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown("""
    <div style='background-color:#e3f2fd; padding:15px; border-radius:10px; border-left:5px solid #1976d2; margin-top:20px;'>
        <h3 style='color:#F28C38; margin-top:0;'>📄 Export Analyses</h3>
    </div>
    """, unsafe_allow_html=True)
    if not engin_data.empty and st.button("📝 Export Word Report"):
        word_buffer = generate_word_report(engin_data, selected_category, figs, descriptions, metrics, predictions, 10000)
        st.download_button("Download Word Report", word_buffer, f'report_{selected_category}.docx', "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab3:
    st.markdown("""
    <div style='background-color:#e3f2fd; padding:20px; border-radius:10px; border-left:5px solid #1976d2; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>Comparisons</h2>
        <p style='color:#424242;'>Analyze costs by category and equipment</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.plotly_chart(
        px.imshow(
            filtered_data.pivot_table(index='Desc_CA', columns='Desc_Cat', values='Montant', aggfunc='sum'),
            labels=dict(x="Consumption Type", y="Equipment", color="Amount (DH)"),
            title='Heatmap of Costs by Equipment and Type',
            aspect="auto",
            height=400
        ),
        use_container_width=True
    )
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        <div style='background-color:#fff8e1; padding:15px; border-radius:10px; border-left:5px solid #ffa000; margin-bottom:10px;'>
            <h3 style='color:#F28C38; margin-top:0;'>Total Cost by Equipment</h3>
        </div>
        """, unsafe_allow_html=True)
        st.plotly_chart(
            px.bar(
                filtered_data.groupby('Desc_CA')['Montant'].sum().reset_index().sort_values('Montant'),
                x='Montant', y='Desc_CA',
                title='',
                height=400
            ),
            use_container_width=True
        )
    with col2:
        st.markdown("""
        <div style='background-color:#e8f5e9; padding:15px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:10px;'>
            <h3 style='color:#F28C38; margin-top:0;'>Cost Distribution by Equipment</h3>
        </div>
        """, unsafe_allow_html=True)
        st.plotly_chart(
            px.box(
                filtered_data, x='Desc_CA', y='Montant',
                title='',
                height=400
            ),
            use_container_width=True
        )

with tab4:
    st.markdown("""
    <div style='background-color:#e3f2fd; padding:20px; border-radius:10px; border-left:5px solid #1976d2; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>Recommendations</h2>
        <p style='color:#424242;'>Practical steps to reduce costs</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        top_categories = filtered_data.groupby('Desc_Cat')['Montant'].sum().nlargest(3)
        st.markdown(f"""
        <div style='background-color:#e3f2fd; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #1976d2;'>
            <h3 style='color:#F28C38; margin-top:0;'>🔍 Top 3 Expenses</h3>
            <ol style='color:#424242;'>
                <li style='margin-bottom:10px;'><b>{top_categories.index[0]}</b> <span style='color:#d32f2f; font-weight:bold;'>{top_categories.values[0]:,.0f} DH</span></li>
                <li style='margin-bottom:10px;'><b>{top_categories.index[1]}</b> <span style='color:#d32f2f; font-weight:bold;'>{top_categories.values[1]:,.0f} DH</span></li>
                <li><b>{top_categories.index[2]}</b> <span style='color:#d32f2f; font-weight:bold;'>{top_categories.values[2]:,.0f} DH</span></li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        problem_equipment = filtered_data.groupby('Desc_CA')['Montant'].sum().nlargest(3)
        st.markdown(f"""
        <div style='background-color:#fff8e1; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #ffa000;'>
            <h3 style='color:#F28C38; margin-top:0;'>🚜 Priority Equipment</h3>
            <div style='display:flex; align-items:center; margin-bottom:10px; color:#424242;'>
                <div style='background-color:#d32f2f; width:20px; height:20px; border-radius:50%; margin-right:10px;'></div>
                <div>{problem_equipment.index[0]} <span style='color:#d32f2f; font-weight:bold;'>- {problem_equipment.values[0]:,.0f} DH</span></div>
            </div>
            <div style='display:flex; align-items:center; margin-bottom:10px; color:#424242;'>
                <div style='background-color:#ffa000; width:20px; height:20px; border-radius:50%; margin-right:10px;'></div>
                <div>{problem_equipment.index[1]} <span style='color:#d32f2f; font-weight:bold;'>- {problem_equipment.values[1]:,.0f} DH</span></div>
            </div>
            <div style='display:flex; align-items:center; color:#424242;'>
                <div style='background-color:#fbc02d; width:20px; height:20px; border-radius:50%; margin-right:10px;'></div>
                <div>{problem_equipment.index[2]} <span style='color:#d32f2f; font-weight:bold;'>- {problem_equipment.values[2]:,.0f} DH</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("""
    <div style='background-color:#e8f5e9; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #388e3c;'>
        <h3 style='color:#F28C38; margin-top:0;'>📅 3-Month Action Plan</h3>
        <div style='border-left:4px solid #388e3c; padding-left:20px;'>
            <div style='margin-bottom:15px; color:#424242;'>
                <h4 style='margin-bottom:5px; color:#2e7d32;'>Month 1: Initial Audit</h4>
                <p>• Review the top 3 most expensive equipment<br>• Analyze frequently replaced parts</p>
            </div>
            <div style='margin-bottom:15px; color:#424242;'>
                <h4 style='margin-bottom:5px; color:#2e7d32;'>Month 2: Negotiations</h4>
                <p>• Contact suppliers for volume discounts<br>• Standardize common parts</p>
            </div>
            <div style='color:#424242;'>
                <h4 style='margin-bottom:5px; color:#2e7d32;'>Month 3: Optimization</h4>
                <p>• Implement preventive maintenance<br>• Train operators on best practices</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div style='background-color:#f3e5f5; padding:20px; border-radius:10px; margin-bottom:20px; border-left:5px solid #8e24aa;'>
        <h3 style='color:#F28C38; margin-top:0;'>✅ Key Action Checklist</h3>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Identify top 5 parts replaced</div>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Compare costs with market standards</div>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Diagnose priority equipment</div>
        <div style='margin-bottom:10px; color:#424242;'><input type='checkbox' style='margin-right:10px;'> Schedule supplier meeting</div>
        <div style='color:#424242;'><input type='checkbox' style='margin-right:10px;'> Set up monthly cost tracking</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown(f"""
    <div style='background-color:#ffebee; padding:20px; border-radius:10px; border-left:5px solid #d32f2f;'>
        <h3 style='color:#F28C38; text-align:center; margin-top:0;'>💡 3 Tips to Reduce Costs</h3>
        <div style='display:flex; justify-content:space-between; text-align:center; margin-top:20px;'>
            <div style='width:30%;'>
                <div style='background-color:#ffcdd2; padding:15px; border-radius:10px; height:120px;'>
                    <h4 style='color:#c62828;'>🛠️ Maintenance</h4>
                    <p style='color:#424242;'>Potential Savings:<br><b>{df['Montant'].mean()*0.25:,.0f} DH/month</b></p>
                </div>
            </div>
            <div style='width:30%;'>
                <div style='background-color:#c8e6c9; padding:15px; border-radius:10px; height:120px;'>
                    <h4 style='color:#2e7d32;'>🔄 Standard Parts</h4>
                    <p style='color:#424242;'>Potential Savings:<br><b>{df['Montant'].mean()*0.40:,.0f} DH/month</b></p>
                </div>
            </div>
            <div style='width:30%;'>
                <div style='background-color:#bbdefb; padding:15px; border-radius:10px; height:120px;'>
                    <h4 style='color:#1565c0;'>📊 Rigorous Tracking</h4>
                    <p style='color:#424242;'>Potential Savings:<br><b>{df['Montant'].mean()*0.15:,.0f} DH/month</b></p>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
with tab5:
    st.markdown("""
    <div style='background-color:#e8f5e9; padding:20px; border-radius:10px; border-left:5px solid #388e3c; margin-bottom:20px;'>
        <h2 style='color:#F28C38; margin-top:0;'>Equipment Consumption Table</h2>
        <p style='color:#424242;'>Detailed consumption by equipment for the selected category</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Prepare the table data without the total row
    table_df = filtered_data[['Date', 'Desc_CA', 'Desc_Cat', 'Montant']].copy()
    table_df['Date'] = table_df['Date'].dt.strftime('%d/%m/%Y')
    table_df['Montant'] = table_df['Montant'].round(2)
    table_df = table_df.rename(columns={
        'Date': 'Date',
        'Desc_CA': 'Equipment',
        'Desc_Cat': 'Consumption Type',
        'Montant': 'Amount (DH)'
    })
    
    # Calculate the total for the 'Amount (DH)' column
    total_montant = table_df['Amount (DH)'].sum()
    
    # Display the table without the total row
    st.dataframe(
        table_df.style.format({
            'Amount (DH)': '{:,.2f} DH',
            'Date': lambda x: x if x else ''
        }).set_properties(**{
            'background-color': '#000000',
            'border': '1px solid #ddd',
            'text-align': 'center'
        }).set_table_styles([
            {'selector': 'th', 'props': [('background-color', '#000000'), ('color', '#424242'), ('font-weight', 'bold')]}
        ]),
        height=600,
        use_container_width=True
    )
    
    # Display the total separately below the table
    st.markdown(f"""
    <div style='background-color:#ffebee; padding:10px; border-radius:10px; text-align:right; margin-top:10px;'>
        <p style='color:#424242; font-size:16px; font-weight:bold;'>Total: {total_montant:,.2f} DH</p>
    </div>
    """, unsafe_allow_html=True)